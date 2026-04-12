"""
26秋上 初三英语A+班 讲义批量生成 v3
格式：TNR + 微软雅黑，专业排版，蓝色主题
题目：嵌入真实中考真题（来源标注）
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNR = 'Times New Roman'
TITLE_SIZE   = Pt(16)
PART_SIZE    = Pt(13)
SECTION_SIZE = Pt(11)
BODY_SIZE    = Pt(11)
NOTE_SIZE    = Pt(9)

# ── Internal helpers ──────────────────────────────────────────

def _set_cell_bg(cell, hex6):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    pr.append(shd)

def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in twips (1/20 of a point). 80 twips ~ 4pt"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def _add_page_number(paragraph):
    """Add centered page number field — each fldChar/instrText wrapped in w:r"""
    run1 = paragraph.add_run('— ')
    run1.font.size = Pt(8)
    run1.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # BEGIN run
    r1 = OxmlElement('w:r')
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar1)
    paragraph._p.append(r1)
    # INSTR run
    r2 = OxmlElement('w:r')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r2.append(instrText)
    paragraph._p.append(r2)
    # SEPARATE run
    r3 = OxmlElement('w:r')
    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    r3.append(fldSep)
    paragraph._p.append(r3)
    # Display run (placeholder text)
    r4 = OxmlElement('w:r')
    t4 = OxmlElement('w:t')
    t4.text = '1'
    r4.append(t4)
    paragraph._p.append(r4)
    # END run
    r5 = OxmlElement('w:r')
    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    r5.append(fldEnd)
    paragraph._p.append(r5)
    # Suffix
    run2 = paragraph.add_run(' —')
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Core functions ────────────────────────────────────────────

def set_font(run, bold=False, size=BODY_SIZE, color=None, italic=False):
    run.font.name = TNR
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    # eastAsia font
    r_elem = run._element
    rPr = r_elem.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

def new_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2)
    # Default style: TNR + 微软雅黑
    style = doc.styles['Normal']
    style.font.name = TNR
    style.font.size = BODY_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), TNR)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.line_spacing = 1.35
    style.paragraph_format.space_after = Pt(4)
    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run('26秋上  初三英语讲义  A+班')
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hr.font.name = TNR
    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)
    return doc

def title_line(doc, text, size=TITLE_SIZE, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, bold=True, size=size)
    return p

def part_head(doc, text):
    """PART ONE / TWO / THREE — blue background block"""
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, '1F4E79')
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_font(r, bold=True, size=PART_SIZE, color='FFFFFF')
    doc.add_paragraph()

def section_head(doc, num_text, title, tag=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'{num_text} {title}')
    set_font(r1, bold=True, size=SECTION_SIZE)
    if tag:
        r2 = p.add_run(f'  ★{tag}')
        set_font(r2, bold=True, size=Pt(9), color='C00000')

def body_text(doc, text, indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=BODY_SIZE)
    return p

def tip_text(doc, text, indent=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, size=BODY_SIZE)

def red_note(doc, text):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # Left red bar
    left_cell = tbl.rows[0].cells[0]
    left_cell.width = Cm(0.2)
    _set_cell_bg(left_cell, 'C00000')
    left_cell.paragraphs[0].add_run(' ')
    # Content
    right_cell = tbl.rows[0].cells[1]
    _set_cell_bg(right_cell, 'FFF0F0')
    _set_cell_margins(right_cell, top=60, bottom=60, left=100, right=80)
    p = right_cell.paragraphs[0]
    r = p.add_run(f'【注】{text}')
    set_font(r, size=NOTE_SIZE, color='C00000', italic=True)

def source_note(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, 'F5F5F5')
    _set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
    p = cell.paragraphs[0]
    r = p.add_run(f'[来源] {text}')
    set_font(r, size=NOTE_SIZE, color='595959', italic=True)

def gt(doc, headers, rows, col_widths=None):
    """Professional table: dark blue header, alternating row colors"""
    ncols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row - dark blue bg, white text
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, '1F4E79')
        _set_cell_margins(c, top=60, bottom=60, left=80, right=80)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, bold=True, size=Pt(10), color='FFFFFF')
    # Data rows - alternating bg
    for ri, row in enumerate(rows):
        bg = 'FFFFFF' if ri % 2 == 0 else 'F5F7FA'
        for ci, v in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            _set_cell_bg(c, bg)
            _set_cell_margins(c, top=40, bottom=40, left=80, right=80)
            p = c.paragraphs[0]
            r = p.add_run(str(v))
            set_font(r, size=Pt(10))
    if col_widths:
        for ri in range(len(tbl.rows)):
            for ci, w in enumerate(col_widths):
                tbl.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()

def dazao_box(doc, stars=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(f'【学习大招】  难度系数：{"★"*stars}{"☆"*(5-stars)}')
    set_font(r, bold=True, size=BODY_SIZE, color='E8672A')
    # Note-taking box with colored left border
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    # Left colored bar (narrow column)
    left_cell = tbl.rows[0].cells[0]
    left_cell.width = Cm(0.3)
    _set_cell_bg(left_cell, 'E8672A')
    left_cell.paragraphs[0].add_run(' ')
    # Right writing area
    right_cell = tbl.rows[0].cells[1]
    _set_cell_bg(right_cell, 'FFF8F0')
    _set_cell_margins(right_cell, top=80, bottom=80, left=120, right=80)
    rp = right_cell.paragraphs[0]
    rp.add_run('动手写下你的大招笔记：').font.size = Pt(9)
    # Add empty lines for writing
    for _ in range(2):
        rp2 = right_cell.add_paragraph()
        rp2.add_run('  ').font.size = Pt(9)
    # Set row height
    tbl.rows[0].height = Cm(2.5)
    doc.add_paragraph()

def ex_label(doc, text='例题精讲（真题选编）'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'▶ {text}')
    set_font(r, bold=True, size=SECTION_SIZE, color='1F4E79')

def mc(doc, num, stem, opts, labels='ABCD', indent=0.4):
    """选择题格式 - auto vertical layout for long options"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f'{num}. {stem}')
    set_font(r, size=BODY_SIZE)
    # Check if any option is long
    max_opt_len = max(len(o) for o in opts)
    if max_opt_len > 25:
        # Vertical layout
        for lb, o in zip(labels, opts):
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Cm(indent + 0.6)
            op.paragraph_format.space_before = Pt(0)
            op.paragraph_format.space_after = Pt(1)
            r = op.add_run(f'{lb}. {o}')
            set_font(r, size=BODY_SIZE)
    else:
        # Horizontal layout
        op = doc.add_paragraph()
        op.paragraph_format.left_indent = Cm(indent + 0.4)
        op.paragraph_format.space_before = Pt(1)
        op.add_run('    '.join([f'{lb}. {o}' for lb, o in zip(labels, opts)])).font.size = BODY_SIZE

def passage_text(doc, text, indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(0.7)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_font(r, size=BODY_SIZE)

# ═══════════════════════════════════════════════════════════════
#  L1: 过去完成时 + 人物传记类阅读
# ═══════════════════════════════════════════════════════════════

def gen_L1():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson One    语法（过去完成时）+ 人物传记类阅读', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 过去完成时 ─────────────────────────────
    part_head(doc, 'PART ONE   语法——过去完成时（Past Perfect Tense）')

    section_head(doc, '①', '定义')
    body_text(doc, '过去完成时表示在过去某一时间点之前已经发生或完成的动作，即"过去的过去"。')
    body_text(doc, '结构：主语 + had + 过去分词（done）')
    doc.add_paragraph()

    section_head(doc, '②', '句型结构', tag='重点')
    gt(doc,
        ['句式', '结构', '例句'],
        [
            ['肯定句', 'S + had + done',       'She had finished her work before dinner.'],
            ['否定句', 'S + had not/hadn\'t + done', 'He hadn\'t seen the film before last night.'],
            ['一般疑问句', 'Had + S + done ...?', 'Had they left when you arrived?'],
        ],
        col_widths=[3, 5, 8]
    )

    section_head(doc, '③', '过去分词构成', tag='易错点')
    gt(doc,
        ['动词类型', '变化规则', '例词'],
        [
            ['一般动词',           '直接 + ed',        'work → worked,  play → played'],
            ['以不发音e结尾',      '去e + ed',          'live → lived,  love → loved'],
            ['辅音+元音+辅音结尾', '双写尾字母 + ed',   'stop → stopped,  plan → planned'],
            ['辅音字母 + y结尾',   '改y为i + ed',       'study → studied,  carry → carried'],
        ],
        col_widths=[4, 4, 8]
    )
    body_text(doc, '不规则变化（高频考查）：')
    gt(doc,
        ['原形', '过去式', '过去分词', '原形', '过去式', '过去分词'],
        [
            ['go',    'went',  'gone',    'take', 'took',  'taken'],
            ['write', 'wrote', 'written', 'see',  'saw',   'seen'],
            ['come',  'came',  'come',    'do',   'did',   'done'],
            ['give',  'gave',  'given',   'know', 'knew',  'known'],
        ],
        col_widths=[2.5, 2.5, 3, 2.5, 2.5, 3]
    )

    section_head(doc, '④', '常用时间状语', tag='重点')
    gt(doc,
        ['时间状语', '含义', '例句'],
        [
            ['by + 过去时间点',                '到……为止',     'By 9 o\'clock, she had finished the report.'],
            ['by the time + 一般过去时从句',   '到……时候',     'By the time we arrived, the show had started.'],
            ['before/when + 一般过去时从句',   '在……之前/当……时','I had read the book before he lent it to me.'],
            ['after + 过去完成时从句',          '在……之后',    'After he had eaten, he went to bed.'],
        ],
        col_widths=[5, 3, 8]
    )

    dazao_box(doc, stars=3)

    ex_label(doc, '例题精讲（真题选编·中学英语网过去完成时专项）')
    source_note(doc, 'trjlseng.com 过去完成时综合练习题')

    real_q_L1 = [
        (1,
         'The film ______ for ten minutes when we got to the cinema.',
         ['have already been on', 'had already begun', 'had already been on', 'have already begun']),
        (2,
         'We ______ five English songs by the end of last term.',
         ['had learned', 'learned', 'have learned', 'will have learned']),
        (3,
         'Han Mei told me she ______ lunch, so she was very hungry.',
         ['has had', 'hasn\'t have', 'have had', 'hadn\'t had']),
        (4,
         'She ______ her keys in the office so she had to wait until her husband ______ home.',
         ['has left; comes', 'had left; would come', 'had left; came', 'left; had come']),
        (5,
         'The meeting ______ when Mr. Wang ______ to school.',
         ['has begun; get', 'has been on; get', 'had begun; got', 'had been on; got']),
        (6,
         'By the end of last week, they ______ the bridge.',
         ['has completed', 'completed', 'will complete', 'had completed']),
        (7,
         'The students ______ their classroom when the visitors arrived.',
         ['have cleaned', 'had cleaned', 'was cleaned', 'have been cleaned']),
        (8,
         'What ______ Jane ______ by the time he was seven?',
         ['did, do', 'has, done', 'did, did', 'had, done']),
    ]
    for num, stem, opts in real_q_L1:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 人物传记类阅读 ─────────────────────────
    part_head(doc, 'PART TWO   语篇精讲——人物传记类阅读（Biography）')

    section_head(doc, '①', '题型特征与解题策略')
    tips = [
        '时间线优先：传记类按时间顺序叙述，快速建立"人物成长时间轴"',
        '转折定考点：but / however / yet / although 后方常是命题核心',
        '主旨定位法：首段末句或尾段首句往往是中心思想所在',
        '推断要有据：推断题答案必须有原文依据，不可凭常识判断',
    ]
    for i, t in enumerate(tips, 1):
        tip_text(doc, f'{i}.  {t}')
    doc.add_paragraph()

    section_head(doc, '②', '阅读理解精练')
    source_note(doc, '来源：2024年中考英语备考专题复习·人物传记类，21世纪教育网')
    doc.add_paragraph()

    paras = [
        'On October 5, Tu Youyou, 84, became "the first Chinese to win a Nobel Prize in natural science." '
        'She is only the 12th woman in history to be awarded the honor.',
        'Tu was born in Ningbo, Zhejiang, China on December 30, 1930. She attended Xiaoshi Middle School '
        'and Ningbo Middle School. Her teacher noted that "Tu liked reading books and she was very hard-working." '
        'From 1951 to 1955, she attended Beijing Medical College. Later Tu received two and a half years '
        'of training in traditional Chinese medicine.',
        'Tu received the award for discovering artemisinin, developed to combat malaria. When Tu began her '
        'research in the late 1960s, "over 240,000 compounds around the world had already been tested, '
        'without any success." Her interest in traditional Chinese medicine led her to ancient texts, where '
        'she found sweet wormwood, used to treat malaria around 400 A.D. After more than 190 failures, '
        'she invented the right drug in 1971.',
        'Tu lacks a Ph.D. and never studied or worked overseas. Former colleague Liao Fuming describes '
        'her as "a tough and stubborn woman." Tu\'s determination to study ancient texts and apply them '
        'in modern science resulted in millions of lives saved.',
    ]
    for para in paras:
        passage_text(doc, para)
    doc.add_paragraph()

    ex_label(doc, '阅读理解题')
    reading_qs = [
        (1, 'How old was Tu Youyou when she graduated from Beijing Medical College?',
            ['12 years old', '21 years old', '25 years old', '28 years old']),
        (2, 'By stating that over 240,000 compounds had been tested without success, the writer indicates _______.',
            ['Tu\'s research work against malaria was very difficult',
             'Scientists at that time were good at testing medicine',
             'Tu received much help from worldwide researchers',
             'Scientists had successfully invented malaria medicine']),
        (3, 'Where did Tu finally find the medicine she wanted?',
            ['At her middle school library', 'In Chinese ancient texts',
             'In universities abroad', 'In modern science magazines']),
        (4, 'Which word best describes Tu Youyou?',
            ['Proud', 'Humorous', 'Active', 'Hard-working']),
        (5, 'What is the best title for the passage?',
            ['Tu Wins the Nobel Prize', 'Tu, a Clever Chinese Woman',
             'Traditional Chinese Medicine', 'How Malaria Was Defeated']),
    ]
    for num, stem, opts in reading_qs:
        mc(doc, num, stem, opts)

    doc.add_paragraph()
    red_note(doc,
        '【教师版备注】'
        '①过去完成时MC来源：中考英语真题题库。答案：1-C 2-A 3-D 4-C 5-C 6-D 7-B 8-D。'
        '②阅读理解（屠呦呦传记）来源：2024年中考备考专题复习。答案：1-C 2-A 3-B 4-D 5-A。'
        '含过去完成时语境（had already been tested / had been tested），与本讲语法呼应。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L1_v3.docx')
    print('L1 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L2: 易混时态辨析 + 中考词汇精讲1
# ═══════════════════════════════════════════════════════════════

def gen_L2():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Two    语法（易混时态辨析）+ 中考词汇精讲 1', Pt(12))
    doc.add_paragraph()

    part_head(doc, 'PART ONE   语法——易混时态辨析')

    section_head(doc, '①', '四大时态核心对比', tag='重点')
    gt(doc,
        ['时态', '结构', '核心含义', '标志词（高频）'],
        [
            ['一般过去时',  'did',            '过去某时发生，已结束',          'yesterday, ago, last, in 2020'],
            ['过去进行时',  'was/were doing', '过去某时正在进行',              'at that time, when引导从句'],
            ['现在完成时',  'have/has done',  '过去动作对现在有影响/持续至今', 'already, yet, ever, just, for, since'],
            ['过去完成时',  'had done',       '"过去的过去"，先于过去某时',   'by, before, when + 过去时从句'],
        ],
        col_widths=[3.5, 3.5, 5, 4.5]
    )

    section_head(doc, '②', '辨析1：一般过去时 vs 现在完成时', tag='考频极高')
    gt(doc,
        ['对比维度', '一般过去时', '现在完成时'],
        [
            ['时间焦点',  '过去某一具体时间点',             '过去动作→现在结果/影响'],
            ['时间状语',  'yesterday / last year / in 2020','already / yet / ever / just / for / since'],
            ['典型句',    'I saw the film last week.',       'I have seen the film.（故现在了解剧情）'],
            ['陷阱提示',  '含具体过去时间→必用过去时',      'for/since+时间段→必用现在完成时'],
        ],
        col_widths=[3.5, 6, 7]
    )

    section_head(doc, '③', '辨析2：过去进行时 vs 一般过去时', tag='常考')
    gt(doc,
        ['对比维度', '过去进行时', '一般过去时'],
        [
            ['强调重点',  '过去某时段正在持续（未结束）',     '过去某时点动作（已完成）'],
            ['结构',      'was/were + doing',                'did（动词过去式）'],
            ['when从句',  '主句用进行时（持续背景动作）',    '主句用过去时（分先后顺序）'],
            ['例句对比',  'I was reading when he called.',   'I finished the book last night.'],
        ],
        col_widths=[3.5, 6, 7]
    )

    section_head(doc, '④', '辨析3：现在完成时 vs 过去完成时', tag='重点')
    gt(doc,
        ['对比维度', '现在完成时', '过去完成时'],
        [
            ['参照时间', '以"现在"为参照',             '以"过去某时"为参照（过去的过去）'],
            ['结构',     'have/has + done',             'had + done'],
            ['例句',     'She has finished her work.',  'She had finished before dinner.'],
            ['信号词',   '强调现在状态/结果',           '主从句均在过去，需分出先后顺序'],
        ],
        col_widths=[3.5, 6, 7]
    )

    section_head(doc, '⑤', '时态三步口诀', tag='大招')
    tip_text(doc, '第一步：找时间状语 → 锁定时间范围（过去/现在/过去的过去）')
    tip_text(doc, '第二步：判完成状态 → 动作是否对现在有影响 / 是否先于另一过去动作')
    tip_text(doc, '第三步：看主从关系 → 两个过去动作谁先谁后，先发生的用过去完成时')
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    ex_label(doc, '例题精讲（真题选编·时态辨析高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题，话题：AI科技/传统文化/环保/青少年')
    tense_qs = [
        (1,
         'By the time the rescue team arrived, the trapped villagers _______ to a safer place.',
         ['moved', 'have moved', 'had moved', 'were moving']),
        (2,
         '—Did you watch the AI science documentary last night?\n'
         '—No. I _______ it the week before on a streaming platform.',
         ['watched', 'have watched', 'had watched', 'was watching']),
        (3,
         'The young scientist _______ on the climate experiment for three years when she finally got the key result.',
         ['has worked', 'had been working', 'worked', 'was working']),
        (4,
         'When I got to the community centre, the volunteers _______ litter along the river for over two hours.',
         ['collected', 'have collected', 'had collected', 'had been collecting']),
        (5,
         '—I can\'t find my digital ID card anywhere. Have you seen it?\n'
         '—Yes. You _______ it on the desk when you came in just now.',
         ['put', 'have put', 'had put', 'were putting']),
        (6,
         'The intangible heritage master said he _______ this craft from his grandfather when he was only eight.',
         ['learned', 'has learned', 'had learned', 'was learning']),
        (7,
         'It was midnight. The smart city streets _______ quiet, and only a few self-driving cars _______ past.',
         ['grew / drove', 'had grown / drove', 'grew / had driven', 'had grown / had driven']),
        (8,
         'Since the new green energy policy _______ into effect, carbon emissions in the city _______ by 30%.',
         ['came / dropped', 'had come / dropped', 'came / have dropped', 'has come / has dropped']),
    ]
    for num, stem, opts in tense_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # PART TWO 词汇精讲1
    part_head(doc, 'PART TWO   中考词汇精讲 1——词根记忆法')

    section_head(doc, '①', '词根 vis / vid（看）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['visible',   'adj.', '看得见的；明显的', 'vis(看) + ible → 能被看到'],
            ['vision',    'n.',   '视力；视野；愿景', 'vis(看) + ion → 看的能力/范围'],
            ['revise',    'v.',   '修改；复习',        're(再) + vis(看) → 再看一遍'],
            ['evidence',  'n.',   '证据；迹象',        'e + vid(看) + ence → 被看见的东西'],
            ['supervise', 'v.',   '监督；管理',        'super(上方) + vis(看) → 从上往下盯'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '②', '词根 port（运；携带）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['transport', 'v./n.', '运输；交通', 'trans(跨越) + port(运) → 运过去'],
            ['export',    'v./n.', '出口；输出', 'ex(出) + port(运) → 运出去'],
            ['import',    'v./n.', '进口；输入', 'im(进) + port(运) → 运进来'],
            ['report',    'v./n.', '报告；汇报', 're(回) + port(运) → 把消息带回来'],
            ['support',   'v./n.', '支持；支撑', 'sup(下方) + port(运) → 从下面撑住'],
            ['portable',  'adj.',  '便携的',     'port(运) + able → 可以带着走的'],
        ],
        col_widths=[3, 2, 3.5, 8]
    )

    section_head(doc, '③', '25-26年考频 TOP 10 高频词', tag='重点')
    gt(doc,
        ['单词', '词性', '释义', '25-26年中考典型语境'],
        [
            ['achieve',    'v.',    '实现；达到',       'achieve a breakthrough in AI（科技突破）'],
            ['attempt',    'v./n.', '尝试；努力',       'attempt to set a world record（挑战记录）'],
            ['contribute', 'v.',    '贡献；有助于',     'contribute to reducing pollution（环保）'],
            ['influence',  'v./n.', '影响；作用',       'have a great influence on youth（榜样）'],
            ['recognize',  'v.',    '识别；承认；认出', 'recognize faces with AI（人工智能）'],
            ['appreciate', 'v.',    '感激；欣赏',       'appreciate traditional art forms（传统文化）'],
            ['determine',  'v.',    '决定；使下定决心', 'be determined to make a difference（传记）'],
            ['encourage',  'v.',    '鼓励；促进',       'encourage students to think independently（教育）'],
            ['struggle',   'v./n.', '奋斗；挣扎',       'struggle against difficulties（励志）'],
            ['transform',  'v.',    '转变；改造',       'transform the old community（城市发展）'],
        ],
        col_widths=[3, 2, 3.5, 8]
    )

    section_head(doc, '④', '词汇语境练习', tag='换新题 · 26一模风格')
    body_text(doc, '用括号内所给词的适当形式填空：')
    fills = [
        '1. The young inventor\'s design _______ the attention of scientists from around the world.  (attract)',
        '2. With the help of AI, the system can now _______ endangered plant species accurately.  (recognize)',
        '3. She was so _______ to win that she trained six hours every day for months.  (determine)',
        '4. Young people should learn to _______ the traditional art forms passed down from ancestors.  (appreciate)',
        '5. The volunteers _______ greatly to improving water quality in the local river last year.  (contribute)',
    ]
    for f in fills:
        body_text(doc, f, indent=0.6)
    red_note(doc, '语境填空务必替换为26一模真题原句。词根第1-2组共11词，下讲继续第3组。')

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L2_v3.docx')
    print('L2 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L3: 感叹句 + 完形技巧（情感态度推理）+ 听口1
# ═══════════════════════════════════════════════════════════════

def gen_L3():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Three    语法（感叹句）+ 完形技巧专项（情感态度推理）+ 听口 1', Pt(12))
    doc.add_paragraph()

    part_head(doc, 'PART ONE   语法——感叹句（Exclamatory Sentences）')

    section_head(doc, '①', '感叹句两大句型', tag='重点')
    gt(doc,
        ['句型', '结构', '适用对象', '例句'],
        [
            ['What感叹句（可数单数）', 'What + a/an + adj + 可数名词单数 + 主谓！', '修饰可数名词单数', 'What an exciting journey it is!'],
            ['What感叹句（复数/不可数）', 'What + adj + 名词复数/不可数 + 主谓！', '修饰复数或不可数名词', 'What great progress we have made!'],
            ['How感叹句（adj）', 'How + adj + 主谓（含be动词）！', '修饰形容词', 'How proud her parents are!'],
            ['How感叹句（adv）', 'How + adv + 主谓（含实义动词）！', '修饰副词', 'How quickly time flies!'],
        ],
        col_widths=[4, 6, 3.5, 5]
    )

    section_head(doc, '②', 'What vs How 辨析口诀', tag='易错点')
    tip_text(doc, '口诀：看感叹词后面紧跟的是什么词性——')
    tip_text(doc, '◆  感叹词后紧跟【名词/名词短语】  →  用 What')
    tip_text(doc, '◆  感叹词后紧跟【形容词/副词】    →  用 How')
    tip_text(doc, '◆  a/an 的用法：可数名词单数前用 a/an；复数和不可数名词前不加')
    doc.add_paragraph()

    section_head(doc, '③', '感叹句三大易错点', tag='易错 · A+必看')
    gt(doc,
        ['易错点', '错误示例', '正确形式', '原因分析'],
        [
            ['a/an漏加/误加',
             'What exciting game it is!',
             'What an exciting game it is!',
             'game可数单数，exciting以元音音素/ɪ/开头，用an'],
            ['What/How混用',
             'How a clever boy he is!',
             'What a clever boy he is!',
             '后接名词boy，应用What，不用How'],
            ['主谓语序错误',
             'What smart is he!',
             'How smart he is!',
             '感叹句主谓不倒装，语序同陈述句'],
        ],
        col_widths=[3, 4, 4, 5.5]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=2)

    ex_label(doc, '例题精讲（2024年各省中考英语感叹句真题精选）')
    source_note(doc, '选自2024年乐山/兰州/遂宁/吉林/长春/云南等省市中考英语真题')

    excl_qs = [
        (1,
         '_______ amazing it is! The Shenzhou XVIII members raise fish for the first time '
         'in Tiangong space station.',
         ['What', 'How', 'What an', 'How an']),
        (2,
         'China has set a 6G speed world record. It\'s 10 to 20 times faster than 5G. '
         '_______ great country it is! I\'m proud of being Chinese.',
         ['How a', 'What a', 'How', 'What']),
        (3,
         'Hi guys! Our team won the first place in the table tennis match. '
         '_______ exciting news it is! We are so happy about it.',
         ['What', 'How', 'What an', 'How an']),
        (4,
         '_______ beautiful poem Happy Rain on a Spring Night is! '
         'Du Fu showed his care for farmers in it.',
         ['How', 'What', 'How a', 'What a']),
        (5,
         '_______ helpful the speech is! It tells us to use the Internet safely.',
         ['What', 'What a', 'How', 'What an']),
        (6,
         'It\'s amazing that China won 201 gold medals in the 19th Asian Games. '
         '_______ encouraging the news is and everyone takes pride in our motherland.',
         ['What', 'What a', 'What an', 'How']),
        (7,
         '_______ meaningful day! We volunteered to clean up our city park.',
         ['How', 'What', 'What a', 'What an']),
        (8,
         '_______ unforgettable experience I had in Harbin Ice and Snow World!',
         ['What an', 'What', 'How an', 'How']),
    ]
    for num, stem, opts in excl_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # PART TWO 完形技巧
    part_head(doc, 'PART TWO   完形技巧专项——情感态度推理')

    section_head(doc, '①', '情感态度题命题规律')
    tips2 = [
        '考查形式：选出最符合语境的情感/态度词（形容词·副词·动词）',
        '命题位置：文章转折点、人物即时反应句、结尾总结段',
        '干扰项特征：含义相近但情感色彩相反，或程度轻重不符合语境',
        '核心逻辑：情感词必须与上下文事件方向一致，优先看转折词和递进词',
    ]
    for i, t in enumerate(tips2, 1):
        tip_text(doc, f'{i}.  {t}')
    doc.add_paragraph()

    section_head(doc, '②', '情感词高频分类', tag='必背')
    gt(doc,
        ['情感类别', '中考高频词汇'],
        [
            ['积极/正面', 'proud, grateful, delighted, amazed, inspired, confident, hopeful, relieved, touched, moved'],
            ['消极/负面', 'ashamed, disappointed, anxious, frightened, embarrassed, desperate, helpless, upset'],
            ['态度词',    'curious, doubtful, supportive, critical, optimistic, indifferent, enthusiastic'],
            ['程度辨析',  'surprised(惊讶) < amazed(惊叹) < astonished(震惊)；glad < pleased < delighted < overjoyed'],
            ['转折信号',  'but / however / yet / although → 情感方向逆转；even / still / increasingly → 情感递进'],
        ],
        col_widths=[3.5, 13]
    )

    section_head(doc, '③', '三步解题法', tag='大招')
    tip_text(doc, 'Step 1 — 定位：圈出空格所在句，标注情感触发事件（好事/坏事/意外）')
    tip_text(doc, 'Step 2 — 逻辑：判断上下文是顺承（同向情感）还是转折（逆向情感）')
    tip_text(doc, 'Step 3 — 排除：删去语义方向错误选项，从剩余选项中比程度·比搭配·选最优')
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    section_head(doc, '④', '完形填空精练')
    red_note(doc,
        '本讲完形须替换为25年末-26年一模真题：成长励志类/青少年科技类，约200-250词，'
        '10-15空，含≥3处情感态度词考查。五维评估：话题✓ 考频✓ 难度A+✓ 情感词丰富✓ 可讲✓'
    )
    doc.add_paragraph()

    source_note(doc, '来源：2025年江苏省连云港市二模英语真题，完形填空')
    doc.add_paragraph()

    cloze_paras = [
        'Everyone dreams of success. But how can you achieve it? The story of Chinese astronaut '
        'Deng Qingming might   (1)   help to answer the question.',
        'Deng Qingming lived 13 kilometers away from his school when he was in senior high. '
        'It   (2)   him more than two hours to walk there. Only in Senior 2 could his poor family '
        '  (3)   a bicycle, which made him very happy. His childhood   (4)   was to go to college '
        'and support his family. "My parents often told me to study hard to be helpful to society. '
        'I always remember their   (5)   for me," said Deng.',
        'Deng became an astronaut in 1998. He spent almost all of his time   (6)   hard. '
        'He saw other astronauts go into space one after another, but he was   (7)   a backup. '
        'Some people asked him what he thought about this. "  (8)   I sometimes felt a little sad, '
        'I never gave up and kept training hard. I must be ready for the   (9)   whenever it comes '
        'to me," he replied. At the age of 56, finally he flew into space with two other astronauts. '
        'Deng   (10)   24 years for this moment!',
        '"For me, I can spend my whole life getting prepared   (11)  , but I will never allow myself '
        'to be unprepared when the task comes by," said Deng. "Twenty-four years was a long journey. '
        'I believed no matter how   (12)   I was, I may be needed by our motherland."',
        'When you want to   (13)   something difficult but meaningful, think about Deng\'s story. '
        'Be   (14)   and keep going. One day, success will knock at your door.',
    ]
    for para in cloze_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    cloze_qs = [
        (1,  'The story _______ help to answer the question.',
             ['might', 'must', 'should', 'need']),
        (2,  'It _______ him more than two hours to walk there.',
             ['paid', 'took', 'spent', 'cost']),
        (3,  'His poor family could _______ a bicycle.',
             ['donate', 'recycle', 'produce', 'afford']),
        (4,  'His childhood _______ was to go to college.',
             ['activity', 'habit', 'dream', 'experience']),
        (5,  'I always remember their _______ for me.',
             ['plan', 'confidence', 'hope', 'message']),
        (6,  'He spent almost all of his time _______ hard.',
             ['training', 'displaying', 'dreaming', 'thinking']),
        (7,  'He was _______ a backup.',
             ['seldom', 'sometimes', 'often', 'always']),
        (8,  '_______ I sometimes felt a little sad, I never gave up.',
             ['But', 'Although', 'Because', 'Unless']),
        (9,  'I must be ready for the _______ whenever it comes.',
             ['choice', 'courage', 'notice', 'chance']),
        (10, 'Deng _______ 24 years for this moment!',
             ['regretted', 'continued', 'realized', 'waited']),
        (11, 'I can spend my whole life getting prepared _______.',
             ['silently', 'sadly', 'rapidly', 'easily']),
        (12, 'No matter how _______ I was, I may be needed.',
             ['weak', 'old', 'careful', 'satisfied']),
        (13, 'When you want to _______ something difficult but meaningful...',
             ['pick up', 'make up', 'give up', 'put up']),
        (14, 'Be _______ and keep going.',
             ['patient', 'curious', 'strict', 'creative']),
    ]
    for num, stem, opts in cloze_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # PART THREE 听口1
    part_head(doc, 'PART THREE   听口训练 1——听后选择 & 听后回答')

    section_head(doc, '①', '听后选择 — 预读三要素')
    tip_text(doc, '要素1【预读选项】：听前10秒扫描选项，预判话题和提问方向（人物/时间/地点/原因）')
    tip_text(doc, '要素2【抓关键词】：捕捉数字·地点·人名·情感词，遇到生词不要停，保持听')
    tip_text(doc, '要素3【注意否定】：not / never / hardly 等否定词后往往是答案关键信息')
    doc.add_paragraph()

    section_head(doc, '②', '听后回答 — 答题四规范')
    tip_text(doc, '规范1：用完整句回答，不能只写单词或短语')
    tip_text(doc, '规范2：回答必须含录音关键信息词（人名·数字·地点·事件）')
    tip_text(doc, '规范3：时态与问题一致（问过去→回答用过去时）')
    tip_text(doc, '规范4：字数不超上限，语法正确优先于内容丰富')
    doc.add_paragraph()

    section_head(doc, '③', '听力练习题（换新题·26一模音频·话题：青少年/科技/社区）', tag='P0换新')
    red_note(doc,
        '音频及题目须从26年一模听力原题或改编题中选取。'
        '话题建议：青少年科技体验/校园活动/社区志愿服务。音频二维码由组长在讲义系统审核后插入。'
    )
    doc.add_paragraph()

    body_text(doc, 'A. 听后选择（共5小题，每小题1分）')
    listen_qs = [
        (1, 'What will the boy do this weekend?',
            ['Visit a science museum', 'Watch a football match', 'Join a volunteer activity']),
        (2, 'How does the girl feel about the AI project?',
            ['Worried about the result', 'Excited about the challenge', 'Bored with the topic']),
        (3, 'Where does the conversation probably take place?',
            ['In a school library', 'In a community lab', 'In a science museum']),
        (4, 'What does the woman suggest the boy do?',
            ['Read more books about AI', 'Talk to the teacher first', 'Join the robotics club']),
        (5, 'How long has the girl been learning to code?',
            ['Six months', 'One year', 'Two years']),
    ]
    for num, stem, opts in listen_qs:
        mc(doc, num, stem, opts, labels='ABC')

    doc.add_paragraph()
    body_text(doc, 'B. 听后回答（共2小题，每小题2分）')
    body_text(doc, '6. ___________________________________________________', indent=0.6)
    body_text(doc, '7. ___________________________________________________', indent=0.6)
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①感叹句8题均选自2024年各省中考真题（乐山/兰州/遂宁/吉林/长春/云南），'
        '答案：1-B 2-B 3-A 4-D 5-C 6-D 7-C 8-A。'
        '②完形填空来源2025年连云港二模（邓清明航天员），'
        '答案：1-A 2-B 3-D 4-C 5-C 6-A 7-D 8-B 9-D 10-D 11-A 12-B 13-C 14-A。'
        '③听口音频上传讲义系统后，由年级组长扫码审核确认。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L3_v3.docx')
    print('L3 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L4: 以读促写——环境与科技（短文填空+主题写作1）
# ═══════════════════════════════════════════════════════════════

def gen_L4():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Four    以读促写——环境与科技（短文填空 + 主题写作精讲 1）', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 短文填空精讲 ─────────────────────────
    part_head(doc, 'PART ONE   短文填空精讲——环境与科技话题')

    section_head(doc, '①', '短文填空题型特征', tag='重点')
    gt(doc,
        ['维度', '说明'],
        [
            ['题型形式',   '一篇200-250词短文，10-12空，给出首字母或词汇提示'],
            ['核心能力',   '语法（时态/语态/词形变换）+ 语境理解 + 语篇逻辑'],
            ['高频考点',   '动词时态语态、名词单复数、形容词副词互换、代词指代'],
            ['话题趋势',   '环保/新能源/智能科技/可持续发展 占比约30%'],
        ],
        col_widths=[3.5, 13]
    )

    section_head(doc, '②', '四步解题法', tag='大招')
    tip_text(doc, 'Step 1 — 通读全文：不填空先读懂主旨，建立语篇框架（约60秒）')
    tip_text(doc, 'Step 2 — 判词性：根据空格前后判断所需词性（名/动/形/副）')
    tip_text(doc, 'Step 3 — 定形态：确定时态/语态/单复数/比较级等具体形式')
    tip_text(doc, 'Step 4 — 回检：填完后通读一遍，检查主谓一致、时态统一、搭配合理')
    doc.add_paragraph()
    dazao_box(doc, stars=3)

    section_head(doc, '③', '短文填空精练（环境科技话题）')
    red_note(doc,
        '本讲短文须替换为26一模真题，话题：环境保护/新能源/AI科技。'
        '五维评估：话题✓ 考频✓ 难度A+✓ 考点覆盖✓ 语篇连贯✓'
    )
    doc.add_paragraph()

    body_text(doc, '[示例短文·正式版替换为26一模原题]')
    fill_paras = [
        'As technology develops, more cities around the world are trying to become "smart cities." '
        'In a smart city, sensors and cameras are   (1)   (use) to collect information about traffic, '
        'air quality, and energy use.',
        'For example, in Singapore, the government has   (2)   (build) a system that can predict traffic jams '
        'before they happen. This helps   (3)   (reduce) the time people spend on the road. '
        'The air quality monitors can also send   (4)   (warn) to people\'s phones when pollution levels '
        'become   (5)   (danger).',
        'However, some people are   (6)   (worry) about privacy. They believe that too many cameras '
        'may make citizens feel   (7)   (comfortable). To solve this problem, smart cities need '
        '  (8)   (strict) rules to protect personal data.',
        'Despite these challenges, the number of smart cities   (9)   (grow) rapidly every year. '
        'Experts believe that by 2030, smart technology will make our lives much   (10)   (good) '
        'and more convenient.',
    ]
    for para in fill_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    # ── PART TWO 主题写作精讲1 ─────────────────────────
    part_head(doc, 'PART TWO   主题写作精讲 1——环境与科技类')

    section_head(doc, '①', '写作审题三步法')
    tip_text(doc, '第一步：圈话题 — 明确写作主题（环保/科技/两者结合）')
    tip_text(doc, '第二步：定体裁 — 议论文/应用文/说明文？确定写作框架')
    tip_text(doc, '第三步：列要点 — 按题目要求列出必须覆盖的内容点，逐一打钩')
    doc.add_paragraph()

    section_head(doc, '②', '环境科技类高频表达', tag='必背')
    gt(doc,
        ['功能', '高频句式'],
        [
            ['引出话题',   'With the development of technology, ... / In recent years, ... has become a hot topic.'],
            ['描述现状',   'It is reported/known that ... / According to a survey, ...'],
            ['分析原因',   'The main reason is that ... / This is mainly because ...'],
            ['提出建议',   'In my opinion, we should ... / It would be better if ...'],
            ['号召结尾',   'Only in this way can we ... / Let\'s take action to ...'],
            ['环保专用',   'protect the environment / reduce carbon emissions / save energy'],
            ['科技专用',   'artificial intelligence / make our life more convenient / play an important role in ...'],
        ],
        col_widths=[3.5, 13]
    )

    section_head(doc, '③', '满分作文三段结构', tag='大招')
    gt(doc,
        ['段落', '内容', '句数', '示例开头'],
        [
            ['开头段', '引出话题 + 表明观点',       '2-3句', 'Nowadays, with the rapid development of AI technology, ...'],
            ['主体段', '分点论述（原因/做法/好处）', '4-6句', 'First of all, ... Moreover, ... In addition, ...'],
            ['结尾段', '总结 + 号召/展望',           '1-2句', 'In short, ... I believe that in the future, ...'],
        ],
        col_widths=[2.5, 5, 2, 7]
    )
    doc.add_paragraph()

    section_head(doc, '④', '写作实战训练')
    red_note(doc,
        '本讲写作题须替换为26一模真题作文题目。'
        '话题方向：环境保护+科技助力/低碳生活/新能源。'
    )
    doc.add_paragraph()
    body_text(doc, '[示例写作题·正式版替换为26一模原题]')
    body_text(doc, '假设你是李华，你的学校正在开展"科技助力环保"主题活动。请你用英语写一篇短文投稿，内容包括：')
    body_text(doc, '1. 科技如何帮助我们保护环境（至少两个方面）', indent=0.6)
    body_text(doc, '2. 你自己在日常生活中会怎么做', indent=0.6)
    body_text(doc, '3. 你对未来的展望', indent=0.6)
    body_text(doc, '要求：80-100词，语句通顺，条理清晰。')
    doc.add_paragraph()

    body_text(doc, '写作空间：')
    tbl = doc.add_table(rows=12, cols=1); tbl.style = 'Table Grid'
    for row in tbl.rows:
        row.height = Cm(0.7)
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①短文填空示例须替换26一模原题，考点覆盖被动语态/词形变换/比较级。'
        '②写作题须替换26一模真题作文，课堂先讲审题+框架，再限时15分钟写作。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L4_v3.docx')
    print('L4 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L5: 被动语态 + 中考词汇精讲2
# ═══════════════════════════════════════════════════════════════

def gen_L5():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Five    语法（被动语态）+ 中考词汇精讲 2', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 被动语态 ─────────────────────────
    part_head(doc, 'PART ONE   语法——被动语态（Passive Voice）')

    section_head(doc, '①', '被动语态核心概念')
    body_text(doc, '当动作的承受者作主语时，使用被动语态。')
    body_text(doc, '结构：主语 + be + 过去分词（done）+ (by + 动作执行者)')
    doc.add_paragraph()

    section_head(doc, '②', '各时态被动语态结构', tag='重点')
    gt(doc,
        ['时态', '主动语态', '被动语态', '例句（被动）'],
        [
            ['一般现在时', 'do/does',        'am/is/are + done',        'English is spoken in many countries.'],
            ['一般过去时', 'did',             'was/were + done',         'The bridge was built in 2020.'],
            ['一般将来时', 'will do',         'will be + done',          'A new school will be built next year.'],
            ['现在进行时', 'am/is/are doing', 'am/is/are being + done',  'The road is being repaired now.'],
            ['现在完成时', 'have/has done',   'have/has been + done',    'The project has been completed.'],
            ['过去完成时', 'had done',        'had been + done',         'The letter had been sent before noon.'],
        ],
        col_widths=[3, 3, 4, 6.5]
    )

    section_head(doc, '③', '情态动词被动语态', tag='考频极高')
    body_text(doc, '结构：情态动词 + be + 过去分词')
    gt(doc,
        ['情态动词', '被动结构', '例句'],
        [
            ['can',    'can be done',    'The problem can be solved in many ways.'],
            ['must',   'must be done',   'Homework must be handed in on time.'],
            ['should', 'should be done', 'Old people should be taken good care of.'],
            ['may',    'may be done',    'The meeting may be put off until next week.'],
            ['need',   'need to be done','The computer needs to be repaired.'],
        ],
        col_widths=[3, 4.5, 9]
    )

    section_head(doc, '④', '主动变被动三步法', tag='大招')
    tip_text(doc, 'Step 1 — 找宾语：把主动句的宾语变为被动句的主语')
    tip_text(doc, 'Step 2 — 变谓语：谓语动词变为"be + 过去分词"，be与新主语保持一致')
    tip_text(doc, 'Step 3 — 加by：原主语变为by的宾语（可省略）')
    body_text(doc, '例：People speak English. → English is spoken (by people).')
    doc.add_paragraph()

    section_head(doc, '⑤', '不能用被动语态的动词', tag='易错点')
    gt(doc,
        ['类型', '常见动词', '例句'],
        [
            ['系动词',     'be, look, sound, taste, feel, smell', 'The soup tastes delicious. (✗ is tasted)'],
            ['不及物动词', 'happen, take place, appear, rise',    'The accident happened yesterday. (✗ was happened)'],
            ['固定搭配',   'belong to, consist of',               'This book belongs to me. (✗ is belonged to)'],
        ],
        col_widths=[3, 6, 7.5]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    ex_label(doc, '例题精讲（真题选编·被动语态中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    passive_qs = [
        (1,
         'It is reported that a new AI museum _______ in our city next year.',
         ['builds', 'will build', 'will be built', 'is built']),
        (2,
         'Look! The students _______ how to use 3D printers by the technology teacher.',
         ['are teaching', 'are being taught', 'have taught', 'were taught']),
        (3,
         'Chinese tea culture _______ around the world since ancient times.',
         ['is known', 'was known', 'has been known', 'had been known']),
        (4,
         'The waste water must _______ before it flows into the river.',
         ['clean', 'be cleaned', 'is cleaned', 'was cleaned']),
        (5,
         '—Your phone looks so old. Why not buy a new one?\n'
         '—There\'s nothing wrong with it. It _______ well.',
         ['is still used', 'still uses', 'was still used', 'still works']),
        (6,
         'The 24 solar terms _______ by ancient Chinese people to guide agricultural activities.',
         ['create', 'created', 'are created', 'were created']),
        (7,
         'Paper _______ by Cai Lun about two thousand years ago.',
         ['invented', 'was invented', 'is invented', 'has been invented']),
        (8,
         'The problem is so difficult that it can\'t _______ easily.',
         ['solve', 'be solved', 'is solved', 'solved']),
    ]
    for num, stem, opts in passive_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 词汇精讲2 ─────────────────────────
    part_head(doc, 'PART TWO   中考词汇精讲 2——词根记忆法')

    section_head(doc, '①', '词根 duct / duce（引导）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['produce',    'v.',    '生产；制造',       'pro(向前) + duce(引导) → 引导出来'],
            ['introduce',  'v.',    '介绍；引进',       'intro(向内) + duce(引导) → 引导进入'],
            ['reduce',     'v.',    '减少；降低',       're(回) + duce(引导) → 往回引→缩减'],
            ['conduct',    'v./n.', '指挥；进行；行为', 'con(共同) + duct(引导) → 共同引导'],
            ['educate',    'v.',    '教育；培养',       'e(出) + duc(引导) + ate → 引导出来→教育'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '②', '词根 struct（建造）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['structure',   'n.',   '结构；建筑物',     'struct(建造) + ure → 建造出来的东西'],
            ['construct',   'v.',   '建造；构建',       'con(共同) + struct(建造) → 合力建造'],
            ['instruct',    'v.',   '指导；命令',       'in(进入) + struct(建造) → 往里建设→教导'],
            ['destroy',     'v.',   '破坏；摧毁',       'de(去除) + stroy(=struct建) → 拆除→破坏'],
            ['instruction', 'n.',   '说明；指令；教学', 'instruct(指导) + ion → 指导内容'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '③', '25-26年考频 TOP 10 高频词', tag='重点')
    gt(doc,
        ['单词', '词性', '释义', '25-26年中考典型语境'],
        [
            ['prevent',    'v.',    '阻止；预防',       'prevent the spread of diseases（健康）'],
            ['provide',    'v.',    '提供；供给',       'provide food and clean water for villagers（公益）'],
            ['develop',    'v.',    '发展；开发',       'develop new green energy technology（科技环保）'],
            ['benefit',    'v./n.', '有益于；好处',     'benefit greatly from exercise（健康生活）'],
            ['consider',   'v.',    '考虑；认为',       'consider using public transport（低碳出行）'],
            ['discover',   'v.',    '发现；发觉',       'discover a new species in the deep sea（科普）'],
            ['suggest',    'v.',    '建议；暗示',       'suggest that we should recycle more（环保）'],
            ['require',    'v.',    '要求；需要',       'be required to wear a helmet（规则）'],
            ['separate',   'v.',    '分开；区分',       'separate waste into different groups（垃圾分类）'],
            ['protect',    'v.',    '保护',             'protect endangered animals from hunting（自然）'],
        ],
        col_widths=[3, 2, 3.5, 8]
    )

    section_head(doc, '④', '词汇语境练习', tag='换新题 · 26一模风格')
    body_text(doc, '用括号内所给词的适当形式填空：')
    fills = [
        '1. Smoking should be _______ in all public places to protect people\'s health.  (prevent)',
        '2. The new hospital _______ with the most advanced medical equipment last month.  (provide)',
        '3. The documentary _______ young people to learn more about ocean protection.  (educate)',
        '4. Waste must be _______ into recyclable and non-recyclable categories.  (separate)',
        '5. Scientists have _______ that regular exercise can greatly improve brain function.  (discover)',
    ]
    for f in fills:
        body_text(doc, f, indent=0.6)
    red_note(doc, '语境填空务必替换为26一模真题原句。词根第3-4组共10词，下讲继续第5组。')

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L5_v3.docx')
    print('L5 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L6: 宾语从句 + 阅读技巧（文章出处判断&主旨归纳）
# ═══════════════════════════════════════════════════════════════

def gen_L6():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Six    语法（宾语从句）+ 阅读技巧专项（文章出处判断 & 主旨归纳）', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 宾语从句 ─────────────────────────
    part_head(doc, 'PART ONE   语法——宾语从句（Object Clause）')

    section_head(doc, '①', '宾语从句三要素', tag='重点')
    gt(doc,
        ['要素', '规则', '例句'],
        [
            ['连接词',   'that(陈述句) / if·whether(一般疑问) / 特殊疑问词(特殊疑问)',
             'I know that he is a teacher. / I wonder if he will come.'],
            ['语序',     '宾语从句用陈述语序（主+谓），不用疑问语序',
             'Can you tell me where the library is? (✗ where is the library)'],
            ['时态呼应', '主句一般现在→从句任意 / 主句一般过去→从句过去范畴 / 客观事实→一般现在',
             'He said the earth is round. (客观事实不变)'],
        ],
        col_widths=[2.5, 7, 7]
    )

    section_head(doc, '②', '连接词选择详解', tag='考频极高')
    gt(doc,
        ['原句类型', '连接词', '例句转换'],
        [
            ['陈述句',     'that（口语可省）',  '"He is a doctor." → She said (that) he was a doctor.'],
            ['一般疑问句', 'if / whether',       '"Is he at home?" → I don\'t know if/whether he is at home.'],
            ['特殊疑问句', '原疑问词（what/who/when/where/how/why）',
             '"Where does he live?" → Do you know where he lives?'],
        ],
        col_widths=[3, 5, 8.5]
    )

    section_head(doc, '③', 'if vs whether 辨析', tag='易错点')
    gt(doc,
        ['情况', '用 whether', '用 if', '说明'],
        [
            ['or not 紧跟', '✓ whether or not', '✗', 'if or not 不规范'],
            ['介词后',       '✓ about whether',  '✗', '介词后只能用whether'],
            ['不定式前',     '✓ whether to do',   '✗', 'if to do 不正确'],
            ['句首作主语',   '✓ Whether he comes is uncertain.', '✗', '主语从句用whether'],
            ['一般宾语从句', '✓', '✓', '两者均可，日常互换'],
        ],
        col_widths=[3, 4.5, 3, 6]
    )

    section_head(doc, '④', '时态呼应规则', tag='重点')
    gt(doc,
        ['主句时态', '从句时态', '例句'],
        [
            ['一般现在时', '根据实际情况选时态',   'I think he will pass the exam.'],
            ['一般过去时', '对应的过去时态',       'She said she was studying then.'],
            ['一般过去时', '客观真理→一般现在时', 'The teacher told us the sun rises in the east.'],
        ],
        col_widths=[4, 5, 7.5]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    ex_label(doc, '例题精讲（真题选编·宾语从句中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    oc_qs = [
        (1,
         '—Could you please tell me _______?\n—By searching the Internet.',
         ['how you got the information', 'how did you get the information',
          'why you got the information', 'why did you get the information']),
        (2,
         'I wonder _______ the Dragon Boat Festival this year.',
         ['what will they do for', 'what they will do for',
          'how will they celebrate', 'how do they celebrate']),
        (3,
         '—Do you know _______ the meeting will start?\n—At 2:00 p.m.',
         ['where', 'when', 'why', 'whether']),
        (4,
         'The teacher asked the students _______ they had finished the AI project.',
         ['that', 'what', 'whether', 'which']),
        (5,
         'He said that he _______ the Great Wall the year before.',
         ['visited', 'has visited', 'had visited', 'would visit']),
        (6,
         'Can you tell me _______? I want to post a letter.',
         ['where is the post office', 'where the post office is',
          'where was the post office', 'where the post office was']),
        (7,
         'The scientist explained that light _______ faster than sound.',
         ['traveled', 'had traveled', 'travels', 'would travel']),
        (8,
         'I don\'t know _______ to go or stay. What\'s your suggestion?',
         ['if', 'whether', 'that', 'which']),
    ]
    for num, stem, opts in oc_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 阅读技巧 ─────────────────────────
    part_head(doc, 'PART TWO   阅读技巧专项——文章出处判断 & 主旨归纳')

    section_head(doc, '①', '文章出处判断技巧')
    gt(doc,
        ['文章来源', '文体特征', '常见标志词'],
        [
            ['报纸/新闻网站', '时效性强，有5W1H（who/what/when/where/why/how）', 'reported, according to, journalist'],
            ['杂志/期刊',     '专题深度分析，有数据图表',                       'research shows, survey, study'],
            ['广告/宣传单',   '有价格/联系方式/促销信息',                       'discount, call, visit, for more info'],
            ['百科/教材',     '客观描述，无个人观点',                           'is defined as, refers to, is known as'],
            ['个人博客/日记', '第一人称，有情感表达',                           'I think, in my opinion, I felt'],
        ],
        col_widths=[3.5, 6.5, 6.5]
    )

    section_head(doc, '②', '主旨归纳四步法', tag='大招')
    tip_text(doc, 'Step 1 — 读首尾：首段末句+尾段首句通常包含中心思想')
    tip_text(doc, 'Step 2 — 找复现词：全文反复出现的核心词即为主题关键词')
    tip_text(doc, 'Step 3 — 概括段意：每段一句话总结，串联起来即为主旨')
    tip_text(doc, 'Step 4 — 排除法：选项太宽（超出文章范围）或太窄（只涉及一段）均排除')
    doc.add_paragraph()
    dazao_box(doc, stars=3)

    section_head(doc, '③', '传统文化类核心词汇', tag='必背')
    gt(doc,
        ['词汇', '释义', '典型语境'],
        [
            ['heritage',    '遗产',     'cultural heritage / intangible heritage（非物质文化遗产）'],
            ['tradition',   '传统',     'follow/keep the tradition（遵循传统）'],
            ['ancestor',    '祖先',     'learn from our ancestors（向先辈学习）'],
            ['dynasty',     '朝代',     'the Tang Dynasty（唐朝）'],
            ['ceremony',    '仪式；典礼','a tea ceremony（茶艺）'],
            ['folk',        '民间的',   'folk music / folk art（民间音乐/艺术）'],
            ['preserve',    '保护；保存','preserve traditional skills（保护传统技艺）'],
            ['calligraphy', '书法',     'Chinese calligraphy（中国书法）'],
        ],
        col_widths=[3, 3, 10.5]
    )

    section_head(doc, '④', '阅读理解精练（2篇）')
    doc.add_paragraph()

    body_text(doc, '[来源：2025年中考英语传统文化阅读专练·传统技艺，21世纪教育网]')
    passage1 = [
        'Oil paper umbrellas have a history of over 1,000 years in China. Yuhang in Hangzhou is famous for making oil paper umbrellas for over 230 years.',
        'Liu Weixue learned the art of making oil paper umbrellas from his grandfather. "I want to bring the art to life," Liu said. "When it rains, some people use the umbrellas — that\'s what I want to see." After learning the art, Liu wanted to win the hearts of young people. So Liu opened an online shop to sell his umbrellas. It now has more than 80,000 fans. But this doesn\'t make Liu speed up. To keep the high quality, he makes only 1,000 umbrellas each year.',
        'Liu hopes more people can know the art. He goes to schools in Hangzhou to teach students. He also goes to other different places to show more people the art. "Liu told us the history of the umbrellas and the ways to make them. Everything was interesting to me," Ding Anqi, a 14-year-old student in Hangzhou, said.',
    ]
    for para in passage1:
        passage_text(doc, para)
    doc.add_paragraph()

    ex_label(doc, '阅读理解题 — Passage 1')
    reading_qs = [
        (1, 'This passage is most probably from _______.',
            ['a science textbook', 'a travel guidebook',
             'a culture magazine', 'a personal diary']),
        (2, 'Yuhang oil paper umbrellas are well known for _______.',
            ['almost 230 years', 'more than 230 years',
             'almost 1,000 years', 'more than 1,000 years']),
        (3, 'Why did Liu Weixue learn to make oil paper umbrellas?',
            ['To make more money', 'To follow his family\'s dream',
             'To bring this old art to life', 'To win the hearts of young people']),
        (4, 'What does the underlined word "quality" mean in Chinese?',
            ['成本', '质量', '效率', '价格']),
        (5, 'What does Liu do to pass on the art of making oil paper umbrellas?\n'
            '① He sells the umbrellas at school  ② He goes to different places to show the art\n'
            '③ He tells students the history  ④ He teaches students how to make them',
            ['①②③', '①②④', '①③④', '②③④']),
    ]
    for num, stem, opts in reading_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    source_note(doc, '来源：2025年中考英语传统文化阅读专练·传统技艺，21世纪教育网')
    passage2 = [
        'In 2024, the Dragon Boat Festival was on June 10th. When it came, many Chinese watched or took part in dragon boat races. However, in a small village in Fuzhou, people were busy making dragon boats during the festival.',
        'Fangzhuang village, known as the "Dragon Boat Village", has a history of more than 700 years of boat making. It is the biggest dragon boat-making place in China. Fang Shaohuang, a 73-year-old man, has a boat-making factory in Fangzhuang village. He began to learn the skill 59 years ago.',
        'It is not easy to build a dragon boat because workers need to use many tools and finish all the work by hand. In Fang\'s factory, an 18-meter-long boat can carry as many as 35 people. It usually takes one worker over 20 days to finish a boat.',
        '"The materials for making the boat are important," Fang said. "We often choose trees over 40 years old. And the wood must be hard and waterproof, or the boat can\'t be used for a long time."',
        'Besides making boats, the factory also welcomes visitors. Many primary school students and college students come to learn about the traditional boat-making skill. "My son hopes to build a hall for showing our boats so that more young people can know about the dragon boat culture," Fang said.',
    ]
    for para in passage2:
        passage_text(doc, para)
    doc.add_paragraph()

    ex_label(doc, '阅读理解题 — Passage 2')
    reading_qs2 = [
        (6, 'How old was Fang when he started to learn to make dragon boats?',
            ['10 years old', '14 years old', '40 years old', '73 years old']),
        (7, 'Which of the following is TRUE about the boats in Fang\'s factory?',
            ['Each can carry up to 20 people', 'Easy for workers to make',
             'Workers make them by hand', 'Small and cheap']),
        (8, 'Why do they choose trees over 40 years old?',
            ['Because old trees are cheaper',
             'Because the wood must be hard and waterproof',
             'Because young trees are not available',
             'Because Fang\'s grandfather told him to']),
        (9, 'What can we learn from the last paragraph?',
            ['Many students visit the factory', 'Fang built a hall to show boats',
             'The factory holds races yearly', 'Fang\'s son hopes to teach young people']),
        (10, 'What is the best title for the passage?',
            ['A dragon boat race', 'The Dragon Boat Festival',
             'Fang Shaohuang and his son', 'The dragon boat making in Fuzhou']),
    ]
    for num, stem, opts in reading_qs2:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①宾语从句题目覆盖三大连接词+语序+时态呼应。'
        '②阅读Passage 1（油纸伞）测试出处判断，Passage 2（龙舟）测试主旨归纳。'
        '来源：2025年中考英语传统文化阅读与写作专练（第二期），21世纪教育网。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L6_v3.docx')
    print('L6 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L7: 定语从句 + 中考词汇精讲3 + 听口2
# ═══════════════════════════════════════════════════════════════

def gen_L7():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Seven    语法（定语从句）+ 中考词汇精讲 3 + 听口 2', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 定语从句 ─────────────────────────
    part_head(doc, 'PART ONE   语法——定语从句（Attributive Clause）')

    section_head(doc, '①', '定语从句基本概念')
    body_text(doc, '定语从句：修饰名词或代词的从句，被修饰的词叫先行词，引导从句的词叫关系词。')
    body_text(doc, '位置：定语从句紧跟先行词之后。')
    doc.add_paragraph()

    section_head(doc, '②', '关系代词/副词选择', tag='重点')
    gt(doc,
        ['关系词', '先行词', '在从句中作', '例句'],
        [
            ['who',   '人',          '主语/宾语',  'The boy who is standing there is my brother.'],
            ['which', '物',          '主语/宾语',  'The book which you lent me is very interesting.'],
            ['that',  '人/物',       '主语/宾语',  'She is the best teacher that I have ever met.'],
            ['where', '地点名词',    '地点状语',   'This is the school where I studied for six years.'],
            ['when',  '时间名词',    '时间状语',   'I\'ll never forget the day when I first came here.'],
        ],
        col_widths=[2.5, 3, 3, 8]
    )

    section_head(doc, '③', '只能用 that 不能用 which 的情况', tag='考频极高')
    gt(doc,
        ['情况', '例句'],
        [
            ['先行词被最高级修饰',         'This is the most exciting film that I\'ve ever seen.'],
            ['先行词被序数词修饰',          'The first thing that we should do is to study hard.'],
            ['先行词是all/everything/nothing等不定代词', 'Is there anything that I can do for you?'],
            ['先行词被the only/the very修饰','This is the very book that I\'m looking for.'],
            ['先行词既有人又有物',           'We talked about the people and things that we remembered.'],
        ],
        col_widths=[6.5, 10]
    )

    section_head(doc, '④', '关系代词省略规则', tag='易错点')
    body_text(doc, '关系代词在从句中作宾语时可以省略：')
    body_text(doc, '  ✓  The man (who/that) I met yesterday is a doctor.')
    body_text(doc, '  ✓  The book (which/that) she bought is very useful.')
    body_text(doc, '关系代词在从句中作主语时不能省略：')
    body_text(doc, '  ✗  The boy (who) is running is my friend. → who不可省')
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    ex_label(doc, '例题精讲（真题选编·定语从句中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    ac_qs = [
        (1,
         'This is the most beautiful painting _______ I have ever seen in the AI art exhibition.',
         ['which', 'that', 'who', 'where']),
        (2,
         'We should help those _______ are in trouble.',
         ['which', 'whom', 'what', 'who']),
        (3,
         'I still remember the day _______ I first went to primary school.',
         ['that', 'which', 'when', 'where']),
        (4,
         'The city _______ we visited last summer has become a famous tourist attraction.',
         ['where', 'which', 'when', 'what']),
        (5,
         'Is there anything _______ you want to say at the graduation ceremony?',
         ['which', 'what', 'that', 'who']),
        (6,
         'The school _______ my father once studied is going to be rebuilt.',
         ['that', 'which', 'where', 'when']),
        (7,
         'Do you know the girl _______ is talking with our English teacher?',
         ['which', 'whom', 'who', 'whose']),
        (8,
         'This is the only dictionary _______ can be used during the exam.',
         ['which', 'who', 'that', 'what']),
    ]
    for num, stem, opts in ac_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 词汇精讲3 ─────────────────────────
    part_head(doc, 'PART TWO   中考词汇精讲 3——词根记忆法')

    section_head(doc, '①', '词根 spec / spect（看）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['expect',    'v.',   '期望；预期',       'ex(向外) + spect(看) → 向外看→期望'],
            ['respect',   'v./n.','尊重；尊敬',       're(回) + spect(看) → 回头再看→值得敬重'],
            ['inspect',   'v.',   '检查；视察',       'in(里面) + spect(看) → 往里看→检查'],
            ['suspect',   'v.',   '怀疑；猜想',       'su(=sub下) + spect(看) → 从下面偷偷看→怀疑'],
            ['special',   'adj.', '特别的；专门的',   'spec(看) + ial → 值得看的→特别的'],
            ['spectator', 'n.',   '观众；旁观者',     'spect(看) + ator(人) → 看的人→观众'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '②', '词根 ject（投掷）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['project',  'n./v.', '项目；计划；投射', 'pro(向前) + ject(投掷) → 抛到前面→项目/投射'],
            ['reject',   'v.',    '拒绝；排斥',       're(回) + ject(投掷) → 扔回去→拒绝'],
            ['object',   'n./v.', '物体；反对',       'ob(反) + ject(投掷) → 扔回来→反对'],
            ['inject',   'v.',    '注射；注入',       'in(进入) + ject(投掷) → 投入里面→注射'],
            ['subject',  'n.',    '科目；主语；主题', 'sub(下) + ject(投掷) → 被放在下面→从属→科目'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '③', '25-26年考频 TOP 10 高频词', tag='重点')
    gt(doc,
        ['单词', '词性', '释义', '25-26年中考典型语境'],
        [
            ['communicate', 'v.',    '交流；沟通',       'communicate with others effectively（教育）'],
            ['experience',  'v./n.', '经历；体验',       'a valuable learning experience（成长）'],
            ['volunteer',   'v./n.', '志愿者；自愿做',   'volunteer to help the elderly（公益）'],
            ['graduate',    'v./n.', '毕业；毕业生',     'graduate from middle school（未来规划）'],
            ['organize',    'v.',    '组织；安排',       'organize a science competition（学校活动）'],
            ['challenge',   'n./v.', '挑战',             'face challenges bravely（成长励志）'],
            ['opportunity', 'n.',    '机会',             'provide more opportunities for students（教育）'],
            ['responsible', 'adj.',  '负责任的',         'be responsible for the project（团队合作）'],
            ['independent', 'adj.',  '独立的',           'become more independent（成长）'],
            ['purpose',     'n.',    '目的；意图',       'the purpose of education is to ...（教育）'],
        ],
        col_widths=[3, 2, 3.5, 8]
    )

    section_head(doc, '④', '词汇语境练习', tag='换新题 · 26一模风格')
    body_text(doc, '用括号内所给词的适当形式填空：')
    fills = [
        '1. The students who _______ from this school last year have all found good jobs.  (graduate)',
        '2. It\'s important for teenagers to learn how to _______ with people from different backgrounds.  (communicate)',
        '3. He gained a lot of _______ during his summer volunteer work at the hospital.  (experience)',
        '4. We should be _______ for our own actions and decisions.  (responsibility)',
        '5. The school _______ a reading week to encourage students to read more books.  (organize)',
    ]
    for f in fills:
        body_text(doc, f, indent=0.6)
    red_note(doc, '语境填空务必替换为26一模真题原句。词根第5-6组共11词，下讲继续第7组。')
    doc.add_paragraph()

    # ── PART THREE 听口2 ─────────────────────────
    part_head(doc, 'PART THREE   听口训练 2——听后转述 & 短文朗读')

    section_head(doc, '①', '听后转述 — 信息提取要点')
    tip_text(doc, '要点1【抓框架】：听第一遍建立人物-事件-结果框架')
    tip_text(doc, '要点2【记关键词】：第二遍补充数字/时间/地点等细节关键词')
    tip_text(doc, '要点3【转述用第三人称】：原文I→He/She，时态可能需要后退一格')
    doc.add_paragraph()

    section_head(doc, '②', '短文朗读 — 评分四维度')
    gt(doc,
        ['维度', '占比', '要点'],
        [
            ['准确性', '40%', '单词发音准确，无漏读/错读/增读'],
            ['流利度', '30%', '语速适中，无过多停顿或回读'],
            ['语调',   '20%', '陈述句降调、疑问句升调、感叹句强调'],
            ['节奏',   '10%', '意群划分合理，重读/弱读自然'],
        ],
        col_widths=[3, 2, 11.5]
    )

    section_head(doc, '③', '听口练习题（换新题·话题：学校教育/未来规划）', tag='P0换新')
    red_note(doc,
        '听后转述及朗读材料须从26年一模真题中选取。'
        '话题建议：学校特色课程/职业规划/毕业感言。音频二维码由组长在讲义系统审核后插入。'
    )
    doc.add_paragraph()

    body_text(doc, 'A. 听后转述（录音播放两遍，根据所听内容完成转述）')
    body_text(doc, '提示词：school / competition / prize / future / plan')
    tbl = doc.add_table(rows=6, cols=1); tbl.style = 'Table Grid'
    for row in tbl.rows:
        row.height = Cm(0.7)
    doc.add_paragraph()

    body_text(doc, 'B. 短文朗读（朗读以下短文，注意语音语调）')
    red_note(doc, '朗读短文须替换为26一模原题短文。')
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①定语从句覆盖who/which/that/where/when，含"只能用that"专题。'
        '②词根第5-6组(spec/spect + ject)共11词。'
        '③听口材料须替换26一模真题。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L7_v3.docx')
    print('L7 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L8: 一轮复习·代词 + 阅读补全短文
# ═══════════════════════════════════════════════════════════════

def gen_L8():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Eight    一轮复习·语法（代词）+ 阅读补全短文精讲', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 代词 ─────────────────────────
    part_head(doc, 'PART ONE   语法系统梳理——代词（Pronoun）')

    section_head(doc, '①', '人称代词 & 物主代词总表', tag='重点')
    gt(doc,
        ['人称', '主格', '宾格', '形容词性物主', '名词性物主', '反身代词'],
        [
            ['第一人称单数', 'I',    'me',   'my',    'mine',   'myself'],
            ['第二人称单数', 'you',  'you',  'your',  'yours',  'yourself'],
            ['第三人称单数', 'he/she/it', 'him/her/it', 'his/her/its', 'his/hers/its', 'himself/herself/itself'],
            ['第一人称复数', 'we',   'us',   'our',   'ours',   'ourselves'],
            ['第二人称复数', 'you',  'you',  'your',  'yours',  'yourselves'],
            ['第三人称复数', 'they', 'them', 'their', 'theirs', 'themselves'],
        ],
        col_widths=[3.5, 2, 2.5, 3, 2.5, 3]
    )

    section_head(doc, '②', '不定代词辨析', tag='考频极高')
    gt(doc,
        ['对比组', '区别', '例句'],
        [
            ['some vs any',
             'some用于肯定句/希望肯定回答的疑问句；any用于否定句/疑问句',
             'Would you like some tea? / I don\'t have any money.'],
            ['one vs it',
             'one指同类中的另一个；it指同一个',
             'I lost my pen. I must buy one. / I found it under the desk.'],
            ['other vs another',
             'other+复数（其余的）；another+单数（另一个）',
             'I have two books. One is new, the other is old.'],
            ['both vs all',
             'both指两者都；all指三者或以上都',
             'Both of them are students. / All of us passed the exam.'],
            ['either vs neither',
             'either两者之一；neither两者都不',
             'Either answer is OK. / Neither of them came.'],
            ['everyone vs every one',
             'everyone只指人（不+of）；every one指人/物（可+of）',
             'Everyone is here. / Every one of the books is useful.'],
        ],
        col_widths=[3, 6, 7.5]
    )

    section_head(doc, '③', 'it 的特殊用法', tag='易错点')
    gt(doc,
        ['用法', '结构', '例句'],
        [
            ['形式主语', 'It is + adj + to do / that ...', 'It is important to learn English well.'],
            ['形式宾语', 'find/think/make + it + adj + to do', 'I find it difficult to get up early.'],
            ['强调句型', 'It is/was ... that/who ...', 'It was Tom who broke the window.'],
            ['时间/天气', 'It is + 时间/天气', 'It is 8 o\'clock now. / It is rainy today.'],
        ],
        col_widths=[3, 5.5, 8]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=3)

    ex_label(doc, '例题精讲（真题选编·代词中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    pronoun_qs = [
        (1,
         '—Whose notebook is this?\n—It must belong to _______. Her name is on the cover.',
         ['her', 'hers', 'she', 'herself']),
        (2,
         'There are many tall buildings on _______ side of the street.',
         ['both', 'every', 'each', 'all']),
        (3,
         '—Would you like tea or coffee?\n—_______. I\'d prefer a glass of water.',
         ['Both', 'Either', 'Neither', 'All']),
        (4,
         'The twins look so much alike that _______ of their classmates can tell them apart at first.',
         ['all', 'few', 'many', 'most']),
        (5,
         'I don\'t like this hat. Could you show me _______ one?',
         ['other', 'the other', 'another', 'others']),
        (6,
         'The AI robot can teach _______ how to play chess step by step.',
         ['it', 'its', 'itself', 'it\'s']),
        (7,
         '_______ is important for us to protect the environment.',
         ['This', 'That', 'It', 'What']),
        (8,
         '—Have you got _______ ready for the science fair?\n—Not yet. I still need some materials.',
         ['something', 'anything', 'everything', 'nothing']),
    ]
    for num, stem, opts in pronoun_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 阅读补全短文 ─────────────────────────
    part_head(doc, 'PART TWO   阅读补全短文精讲')

    section_head(doc, '①', '题型特征')
    gt(doc,
        ['维度', '说明'],
        [
            ['形式',     '一篇300词左右短文挖去4-5个句子，从6-7个选项中选出正确句子补全'],
            ['核心能力', '语篇逻辑 + 上下文衔接 + 指代关系识别'],
            ['难度',     '中考阅读题中难度最高的题型之一'],
        ],
        col_widths=[3, 13.5]
    )

    section_head(doc, '②', '五步解题法', tag='大招')
    tip_text(doc, 'Step 1 — 通读文章：跳过空格先读完全文，把握主旨和段落逻辑')
    tip_text(doc, 'Step 2 — 读选项：标记每个选项的关键词和功能（过渡/举例/总结/转折）')
    tip_text(doc, 'Step 3 — 看上下文：重点关注空格前后的代词指代、逻辑连接词、复现词')
    tip_text(doc, 'Step 4 — 试填验证：将选项代入，检查语义连贯、代词一致、逻辑通顺')
    tip_text(doc, 'Step 5 — 排除检查：最后用剩余选项验证已填答案，确保全文逻辑完整')
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    section_head(doc, '③', '科普类高频词', tag='必背')
    gt(doc,
        ['词汇', '释义', '典型语境'],
        [
            ['experiment',  '实验',     'do/carry out an experiment（做实验）'],
            ['research',    '研究',     'do research on/into ...（研究……）'],
            ['technology',  '技术',     'advanced technology / information technology'],
            ['explore',     '探索',     'explore outer space / explore the ocean（探索太空/海洋）'],
            ['material',    '材料；物质','new materials / raw materials（新材料/原材料）'],
            ['measure',     '测量；措施','take measures to ...（采取措施……）'],
            ['observe',     '观察',     'observe the stars / observe carefully（观察星星/仔细观察）'],
            ['quantity',    '数量',     'a large quantity of data（大量数据）'],
        ],
        col_widths=[3, 3, 10.5]
    )

    section_head(doc, '④', '阅读补全精练')
    source_note(doc, '来源：2025年江苏省连云港市中考英语真题，阅读还原')
    doc.add_paragraph()

    restore_paras = [
        'Cooking, the art and science of preparing food, is far more than daily housework. '
        'It\'s like a fun classroom where you can learn different subjects all at once!   (41)  ',
        'Science is everywhere in cooking. When you make bread, the yeast eats sugar and makes '
        'bubbles of gas.   (42)   When you cook meat and it turns brown and smells good, that\'s '
        'cool chemistry. Also, think about heating water, which shows how heat moves. It is all about physics.',
        '  (43)   You need to follow the instructions when you mix materials, like "adding 1/4 spoon '
        'of salt". Sometimes you have to make some changes to feed more or fewer people. That means '
        'doing math. Even buying the required materials with low cost teaches you math.',
        'Food has a lot of history. Pasta didn\'t start in Italy — it came from Asia long ago! Spices '
        'were once traded like gold, and this changed how countries talked to each other.   (44)   '
        'In hot places, people grow more different foods than in cold places. Learning about these things '
        'helps you understand how the world works.',
        'Cooking is like a secret key that opens the door to learning many subjects. It\'s not just about '
        'making yummy food.   (45)   Schools should teach cooking because it helps students see how '
        'everything in learning is connected!',
    ]
    for para in restore_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    body_text(doc, '从下列选项中选出能填入文中空白处的最佳选项（7选5）：')
    opts_text = [
        'A. Geography also matters.',
        'B. This is biology in action.',
        'C. Math helps a lot in the kitchen.',
        'D. Cooking avoids wasting money.',
        'E. It\'s also about becoming smarter in lots of ways.',
        'F. Some students often spend money buying ready-made food.',
        'G. The following will explain why cooking should be taught at school.',
    ]
    for o in opts_text:
        body_text(doc, o, indent=0.6)
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①代词真题覆盖人称/物主/反身/不定/it用法，答案：1-B 2-C 3-C 4-A 5-C 6-B 7-C 8-C。'
        '②阅读还原来源2025连云港中考（烹饪中的科学），答案：41-G 42-B 43-C 44-A 45-E。'
        '话题覆盖生物/化学/物理/数学/地理跨学科科普，难度适合A+班。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L8_v3.docx')
    print('L8 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L9: 一轮复习·介词 + 中考词汇精讲4 + 完形填空
# ═══════════════════════════════════════════════════════════════

def gen_L9():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Nine    一轮复习·语法（介词）+ 中考词汇精讲 4 + 完形填空', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 介词 ─────────────────────────
    part_head(doc, 'PART ONE   语法考点精讲——介词（Preposition）')

    section_head(doc, '①', '时间介词', tag='重点')
    gt(doc,
        ['介词', '用法', '例句'],
        [
            ['at',     '具体时刻/节日（无day）',    'at 8:00 / at noon / at night / at Christmas'],
            ['on',     '具体某天/星期/有day的节日', 'on Monday / on May 1st / on Children\'s Day / on a rainy morning'],
            ['in',     '年/月/季节/一天中的段',     'in 2026 / in May / in spring / in the morning'],
            ['for',    '持续时间段',                'for two hours / for a long time'],
            ['since',  '从过去某时起（完成时）',    'since 2020 / since I was born'],
            ['by',     '到……为止',                  'by next Friday / by the end of this term'],
            ['during', '在……期间',                  'during the summer vacation'],
        ],
        col_widths=[2.5, 5, 9]
    )

    section_head(doc, '②', '地点/方向介词', tag='重点')
    gt(doc,
        ['介词', '用法', '例句'],
        [
            ['at',      '小地点/门牌号',            'at the bus stop / at No. 12 Park Street'],
            ['in',      '大地点/内部空间',          'in Beijing / in the classroom'],
            ['on',      '表面/楼层',                'on the desk / on the second floor'],
            ['between', '两者之间',                 'between A and B'],
            ['among',   '三者或以上之间',           'among all the students'],
            ['across',  '穿过（表面）',             'walk across the street'],
            ['through', '穿过（内部空间）',         'go through the forest'],
            ['over',    '从上方越过',               'fly over the bridge'],
        ],
        col_widths=[2.5, 5, 9]
    )

    section_head(doc, '③', '高频介词搭配', tag='考频极高')
    gt(doc,
        ['搭配类型', '常见搭配'],
        [
            ['be + adj + 介词', 'be good at, be interested in, be proud of, be afraid of, be worried about, be famous for'],
            ['动词 + 介词',     'look at, listen to, wait for, depend on, belong to, agree with, look forward to'],
            ['介词短语',        'in time, on time, at once, by accident, in fact, in the end, at first, on purpose'],
            ['特殊搭配',        'on foot, by bus/bike/train, in the way, on the way, in one\'s opinion'],
        ],
        col_widths=[3.5, 13]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=3)

    ex_label(doc, '例题精讲（真题选编·介词中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    prep_qs = [
        (1,
         'The high-speed train from Beijing _______ Shanghai takes only about 4 hours.',
         ['at', 'to', 'in', 'on']),
        (2,
         'Our school usually holds a sports meeting _______ October every year.',
         ['at', 'on', 'in', 'for']),
        (3,
         'The old man has lived in this city _______ more than thirty years.',
         ['since', 'for', 'in', 'during']),
        (4,
         'It\'s very kind _______ you to help me carry the heavy box.',
         ['for', 'to', 'of', 'with']),
        (5,
         'The little river runs _______ the village and provides water for the farmers.',
         ['across', 'over', 'through', 'along']),
        (6,
         '—What do you usually do _______ weekends?\n—I often go cycling with my friends.',
         ['in', 'on', 'at', 'to']),
        (7,
         'The students arrived _______ the train station just _______ time to catch the 8:00 train.',
         ['at; in', 'in; on', 'at; on', 'in; in']),
        (8,
         'Nobody _______ Tom knew the answer, because he was the only one who had read the book.',
         ['beside', 'besides', 'except', 'among']),
    ]
    for num, stem, opts in prep_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 词汇精讲4 ─────────────────────────
    part_head(doc, 'PART TWO   中考词汇精讲 4——词根记忆法')

    section_head(doc, '①', '词根 rupt（破裂）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['interrupt', 'v.',   '打断；中断',       'inter(之间) + rupt(破) → 从中间断开→打断'],
            ['corrupt',   'adj.', '腐败的；堕落的',   'cor(全部) + rupt(破) → 全部破坏→腐败'],
            ['erupt',     'v.',   '爆发；喷发',       'e(出) + rupt(破) → 破裂而出→爆发'],
            ['disrupt',   'v.',   '扰乱；使中断',     'dis(分开) + rupt(破) → 使分裂→扰乱'],
            ['bankrupt',  'adj.', '破产的',           'bank(银行) + rupt(破) → 银行破了→破产'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '②', '词根 press（压）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['express',    'v.',   '表达；快递',       'ex(出) + press(压) → 压出来→表达'],
            ['impress',    'v.',   '给……留下印象',     'im(进入) + press(压) → 压入心中→留下印象'],
            ['pressure',   'n.',   '压力',             'press(压) + ure → 压的状态→压力'],
            ['depress',    'v.',   '使沮丧；使萧条',   'de(向下) + press(压) → 往下压→使沮丧'],
            ['compress',   'v.',   '压缩；压紧',       'com(一起) + press(压) → 压到一起→压缩'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '③', '25-26年考频 TOP 10 高频词', tag='重点')
    gt(doc,
        ['单词', '词性', '释义', '25-26年中考典型语境'],
        [
            ['passenger',  'n.',    '乘客',             'bus/subway passengers（城市交通）'],
            ['direction',  'n.',    '方向；指示',       'follow the direction on the map（出行）'],
            ['crowd',      'n./v.', '人群；拥挤',       'a crowd of tourists / crowded streets（城市）'],
            ['connect',    'v.',    '连接；联系',       'connect different cities by high-speed rail（交通）'],
            ['escape',     'v.',    '逃跑；逃避',       'escape from the burning building（安全）'],
            ['warn',       'v.',    '警告；提醒',       'warn drivers of the icy road（天气）'],
            ['deliver',    'v.',    '递送；发表',       'deliver packages by drone（科技）'],
            ['block',      'v./n.', '堵塞；街区',       'block the traffic / walk two blocks（交通）'],
            ['surround',   'v.',    '围绕；包围',       'be surrounded by mountains（地理）'],
            ['signal',     'n./v.', '信号；标志',       'a traffic signal / send a signal（交通/通信）'],
        ],
        col_widths=[3, 2, 3.5, 8]
    )

    section_head(doc, '④', '词汇语境练习', tag='换新题 · 26一模风格')
    body_text(doc, '用括号内所给词的适当形式填空：')
    fills = [
        '1. All the _______ on the flight were asked to fasten their seatbelts during landing.  (passenger)',
        '2. The new highway _______ the small village to the nearest city last year.  (connect)',
        '3. Police _______ local residents about the coming typhoon through loudspeakers.  (warn)',
        '4. The river is _______ by green hills on both sides, making it a beautiful place.  (surround)',
        '5. A new system has been developed to _______ food to people\'s doors using robots.  (deliver)',
    ]
    for f in fills:
        body_text(doc, f, indent=0.6)
    red_note(doc, '语境填空务必替换为26一模真题原句。词根第7-8组共10词，下讲继续第9组。')
    doc.add_paragraph()

    # ── PART THREE 完形填空 ─────────────────────────
    part_head(doc, 'PART THREE   完形填空精练（2024浙江中考真题）')

    section_head(doc, '⑤', '完形填空精练')
    source_note(doc, '来源：2024年浙江省中考英语真题，第三部分语言运用第一节')
    doc.add_paragraph()

    cloze_paras = [
        'Bruce loves music because it makes him feel easy. He is shy, and making friends has been '
        '  (1)   for him. That\'s why his mother asked him to take the   (2)   this year. '
        '"You\'d better not sit there with your headphones on," she told him.',
        'Bruce nodded. But the moment he sat on the bus, he   (3)   his headphones and closed his eyes. '
        'He was soon lost in a song of his favorite band, the Blue-Bob.   (4)  , he sensed someone around him. '
        'He opened his eyes and saw a boy from his   (5)  .',
        '"Hey, Bruce? I\'m Mike," the boy said. "I think we both take the music class. Did I hear '
        '  (6)   singing a song by the Blue-Bob just now?"',
        '"Uh, no," Bruce\'s face turned red, "That wasn\'t me." Mike looked a little   (7)  . '
        'Without saying anything, he went away.',
        'For the next few days, Bruce kept asking himself   (8)   he had lied to Mike. '
        '"Maybe he also loves the Blue-Bob. Maybe he was just being   (9)  ," he thought. '
        'So when they saw each other on the bus the next day, Bruce managed to   (10)   a smile on his face.',
        '"Hi, Bruce," Mike said. "I can hear your   (11)   every day, and you have such a good taste."',
        'Bruce\'s eyes lit up. "Don\'t you think the songs are too   (12)  ?" he asked. '
        '"Not at all," Mike answered. "I love songs of the past too. My grandpa is really into music, '
        'and he has   (13)   me to sing many of them."',
        'Bruce listened as Mike talked about his favorite songs. And they shared the headphones, '
        'singing softly along with the music   (14)   Bruce got off at his stop.',
        '"See you tomorrow!" Mike said.',
        'Bruce waved goodbye. "I must tell Mom how I   (15)   Mike and we had the best time...'
        'with headphones on," he smiled to himself.',
    ]
    for para in cloze_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    cloze_qs = [
        (1,  'Making friends has been _______ for him.',
             ['funny', 'hard', 'boring', 'relaxing']),
        (2,  'His mother asked him to take the _______.',
             ['car', 'taxi', 'bus', 'train']),
        (3,  'He _______ his headphones and closed his eyes.',
             ['put on', 'took away', 'gave up', 'paid for']),
        (4,  '_______, he sensed someone around him.',
             ['Actually', 'Luckily', 'Naturally', 'Suddenly']),
        (5,  'He saw a boy from his _______.',
             ['story', 'grade', 'lab', 'dream']),
        (6,  'Did I hear _______ singing a song?',
             ['him', 'her', 'you', 'them']),
        (7,  'Mike looked a little _______.',
             ['worried', 'excited', 'tired', 'surprised']),
        (8,  'Bruce kept asking himself _______ he had lied.',
             ['why', 'how', 'when', 'whether']),
        (9,  'Maybe he was just being _______.',
             ['patient', 'honest', 'proud', 'friendly']),
        (10, 'Bruce managed to _______ a smile on his face.',
             ['wear', 'hide', 'win', 'find']),
        (11, 'I can hear your _______ every day.',
             ['news', 'music', 'name', 'ideas']),
        (12, 'Don\'t you think the songs are too _______?',
             ['sad', 'old', 'slow', 'serious']),
        (13, 'He has _______ me to sing many of them.',
             ['chosen', 'allowed', 'taught', 'ordered']),
        (14, 'Singing softly along with the music _______ Bruce got off.',
             ['if', 'after', 'until', 'although']),
        (15, 'I must tell Mom how I _______ Mike.',
             ['met', 'invited', 'caught', 'followed']),
    ]
    for num, stem, opts in cloze_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①介词专题覆盖时间/地点/方向/固定搭配，共8道MC。'
        '②词根第7-8组(rupt + press)共10词。'
        '③完形填空已替换为2024浙江中考真题（Bruce & Mike校车交友故事）。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L9_v3.docx')
    print('L9 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L10: 一轮复习·形容词副词 + 说明文阅读
# ═══════════════════════════════════════════════════════════════

def gen_L10():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Ten    一轮复习·语法（形容词 & 副词）+ 说明文阅读', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 形容词副词 ─────────────────────────
    part_head(doc, 'PART ONE   语法专项——形容词 & 副词（比较级 / 最高级 / 修饰语序）')

    section_head(doc, '①', '比较级 & 最高级变化规则', tag='重点')
    gt(doc,
        ['类型', '变化规则', '原级', '比较级', '最高级'],
        [
            ['单音节词',        '+ er / est',              'tall',   'taller',   'tallest'],
            ['以e结尾',         '+ r / st',                'nice',   'nicer',    'nicest'],
            ['辅+元+辅结尾',    '双写尾字母 + er / est',   'big',    'bigger',   'biggest'],
            ['辅音+y结尾',      '改y为i + er / est',       'happy',  'happier',  'happiest'],
            ['多音节词/部分双音','more / most + 原级',      'beautiful','more beautiful','most beautiful'],
        ],
        col_widths=[3.5, 4.5, 2.5, 3, 3]
    )

    section_head(doc, '②', '不规则变化（高频考查）', tag='必背')
    gt(doc,
        ['原级', '比较级', '最高级', '原级', '比较级', '最高级'],
        [
            ['good/well', 'better',    'best',    'bad/badly/ill', 'worse',  'worst'],
            ['many/much', 'more',      'most',    'little',        'less',   'least'],
            ['far',       'farther/further', 'farthest/furthest', 'old', 'older/elder', 'oldest/eldest'],
        ],
        col_widths=[2.5, 3, 2.5, 3, 3, 2.5]
    )

    section_head(doc, '③', '比较级常考句型', tag='考频极高')
    gt(doc,
        ['句型', '含义', '例句'],
        [
            ['A + 比较级 + than + B',          'A比B更……',       'Tom is taller than Jack.'],
            ['the + 比较级, the + 比较级',     '越……越……',       'The more you read, the more you learn.'],
            ['比较级 + and + 比较级',           '越来越……',       'The weather is getting warmer and warmer.'],
            ['one of the + 最高级 + 复数名词', '最……之一',       'Beijing is one of the largest cities in China.'],
            ['A + as + 原级 + as + B',          'A和B一样……',    'He runs as fast as his brother.'],
            ['not as/so + 原级 + as',           'A不如B……',       'This book is not as interesting as that one.'],
        ],
        col_widths=[5.5, 3, 8]
    )

    section_head(doc, '④', '形容词修饰语序', tag='易错点')
    body_text(doc, '多个形容词修饰同一名词时，遵循以下语序（口诀：限观形龄色国材）：')
    gt(doc,
        ['位置', '类型', '例词'],
        [
            ['1', '限定词（冠词/指示/物主）', 'a, the, this, my'],
            ['2', '观点/评价',                'beautiful, useful, nice'],
            ['3', '大小/形状',                'big, small, round, long'],
            ['4', '年龄/新旧',                'old, new, young'],
            ['5', '颜色',                     'red, blue, white'],
            ['6', '国籍/来源',                'Chinese, American'],
            ['7', '材料',                     'wooden, cotton, glass'],
        ],
        col_widths=[1.5, 5, 10]
    )
    body_text(doc, '例：a beautiful small old red Chinese wooden box')
    doc.add_paragraph()
    dazao_box(doc, stars=3)

    ex_label(doc, '例题精讲（真题选编·形容词副词中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    adj_qs = [
        (1,
         'Of all the subjects, science is _______ for me.',
         ['interesting', 'more interesting', 'most interesting', 'the most interesting']),
        (2,
         'The Yangtze River is _______ than any other river in China.',
         ['long', 'longer', 'longest', 'the longest']),
        (3,
         'The _______ you practice speaking English, the _______ your pronunciation will be.',
         ['more; better', 'much; good', 'more; best', 'most; better']),
        (4,
         'Our environment is becoming _______ because people are planting more trees.',
         ['bad and bad', 'worse and worse', 'better and better', 'good and good']),
        (5,
         'Which is _______, the sun, the moon or the earth?',
         ['big', 'bigger', 'biggest', 'the biggest']),
        (6,
         'This maths problem is not _______ difficult _______ that one.',
         ['so; as', 'as; so', 'as; as', 'so; so']),
        (7,
         'Li Hua speaks English _______ in our class. Everyone admires her.',
         ['good', 'well', 'better', 'best']),
        (8,
         'It\'s getting _______ outside. You\'d better take an umbrella with you.',
         ['more and more heavily', 'heavier and heavier',
          'more and more heavy', 'heavily and heavily']),
    ]
    for num, stem, opts in adj_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 说明文阅读 ─────────────────────────
    part_head(doc, 'PART TWO   阅读精讲——说明文（含数据图表）')

    section_head(doc, '①', '说明文阅读策略')
    tip_text(doc, '1. 抓主题句：说明文每段首句通常是中心句，快速串联即可把握全文结构')
    tip_text(doc, '2. 读数据：图表/数字是命题重点，注意单位、趋势（increase/decrease/remain）')
    tip_text(doc, '3. 因果链：说明文常考"原因→结果"推理，关注because/as a result/lead to/cause')
    tip_text(doc, '4. 词义猜测：利用上下文定义句（...means.../...is called...）或同位语推断')
    doc.add_paragraph()

    section_head(doc, '②', '自然环境类核心词汇', tag='必背')
    gt(doc,
        ['词汇', '释义', '典型语境'],
        [
            ['climate',     '气候',     'climate change / a mild climate（气候变化/温和气候）'],
            ['pollution',   '污染',     'air/water/noise pollution（空气/水/噪音污染）'],
            ['ecosystem',   '生态系统', 'protect the ecosystem（保护生态系统）'],
            ['species',     '物种',     'endangered species / a new species（濒危物种/新物种）'],
            ['carbon',      '碳',       'carbon dioxide / low-carbon life（二氧化碳/低碳生活）'],
            ['recycle',     '回收利用', 'recycle paper and plastic（回收纸和塑料）'],
            ['sustainable', '可持续的', 'sustainable development（可持续发展）'],
            ['disaster',    '灾难',     'natural disaster（自然灾害）'],
        ],
        col_widths=[3, 3, 10.5]
    )

    section_head(doc, '③', '阅读理解精练（2篇）')

    body_text(doc, 'Passage 1')
    source_note(doc, '来源：2024年北京市中考英语真题，阅读理解B篇')
    doc.add_paragraph()

    p1_paras = [
        'The school year began. As president of the recycling club, Scott was thinking about new activities '
        'to encourage other students to become more enthusiastic about recycling. His club had helped to '
        'recycle a lot of waste for the past five years and he hoped that this year they would do even better.',
        'Scott went home one day and looked online, hoping to find ideas for events the club could organize '
        'at school. During his research, he learned that the amount of electronic waste, or e-waste, is '
        'increasing rapidly. Scott was upset with himself for not noticing this problem sooner. At that moment, '
        'he began to develop a plan.',
        '"There is a special project I want us to work on this term," Scott announced at the recycling club '
        'meeting the next day. "We have all heard about e-waste, but recently I learned about the bad effects '
        'it\'s having on our environment." He went on to tell the club members that he wanted them to organize '
        'an e-waste drive — a day when students and their families could drop off unwanted electronics to be '
        'recycled. "Now let\'s get to work!"',
        'For several weeks, the recycling club was busy preparing for the upcoming e-waste drive, or e-drive '
        'as it was soon called.',
        'The big day finally arrived, and Scott was nervous. He and the other members arrived at school early '
        'to make sure that everything would go smoothly. Everyone waited eagerly, hoping that students of the '
        'school would take part in the event.',
        'It wasn\'t long before the first cars pulled into the school parking lot. After a while, more people '
        'arrived. Phones, TV sets, computers and keyboards soon began piling up.',
        'At the end of the event, Scott heard someone suggest that the school should have an e-drive event '
        'every term. Scott smiled, realizing that a simple action could truly have a lasting influence.',
    ]
    for para in p1_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    ex_label(doc, '阅读理解题 — Passage 1')
    p1_qs = [
        (1, 'What did Scott learn during his research online?',
            ['Students are enthusiastic about recycling',
             'There is a lot of waste recycled every year',
             'The amount of e-waste is increasing rapidly',
             'There are many recycling clubs in the country']),
        (2, 'How did Scott feel at first on that big day?',
            ['Proud', 'Upset', 'Excited', 'Nervous']),
        (3, 'What did Scott realize at the end of the e-drive event?',
            ['Environmental problems could be solved',
             'A simple action could bring a long-term effect',
             'Teamwork is valuable for the success of events',
             'Family members are interested in school activities']),
    ]
    for num, stem, opts in p1_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    body_text(doc, 'Passage 2')
    source_note(doc, '来源：2024年浙江省中考英语真题，阅读理解D篇')
    doc.add_paragraph()

    p2_paras = [
        'Do you get in trouble for feeding your dog on food from dinner under the table? A new study '
        'suggests that by sharing "people food", you might actually be helping to keep your pet healthy.',
        'Many dogs are fed dry processed food. Science has shown that a diet of processed food is not healthy '
        'for humans. Scientists from the University of Helsinki, in Finland, say this diet is not good for '
        'our pet dogs, either.',
        'The scientists discovered that the food a dog might find in an owner\'s home such as eggs, fish, '
        'meat and vegetables may be better than a diet of just dog food. This is because having different '
        'kinds of food leads to more variety in microbes, which are little living things, too small to be '
        'seen with our eyes. There are millions of them in human and animal bodies. Some are bad for us, '
        'but many are good for us, and they play important roles in helping us digest food.',
        'A total of 8,500 young pet dogs took part in the study. 4,500 of them were under 6 months old, '
        'and the rest were between 6 and 18 months old. The owners were asked about what food they had fed '
        'their pets and how often. When the dogs reached adulthood, their owners reported that about 20% of '
        'the dogs showed CE symptoms (chronic enteropathy) such as weight loss. However, if dogs had been '
        'fed a mixed diet while young, it was 22% less possible for them to experience CE symptoms later in life.',
        'Dr. Anna Hielm-Bjorkman told The Times of London, "Dog food on sale are described as providing '
        'a complete diet...But what we show is that variety is important. Nobody would give 12 years of '
        'the same food to a child — why should a dog be different?"',
    ]
    for para in p2_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    ex_label(doc, '阅读理解题 — Passage 2')
    p2_qs = [
        (4, 'According to the new study, dog owners _______.',
            ['must offer dry food to their dogs',
             'can feed dogs on "people food"',
             'had better buy more processed food',
             'shouldn\'t make food by themselves']),
        (5, 'Paragraph 3 mainly explains why _______.',
            ['dogs need different kinds of food',
             'little living things stay in human bodies',
             'scientists study dry processed food',
             'people should keep dogs at home']),
        (6, 'How does the writer describe the study in Paragraph 4 to make it believable?',
            ['By using pictures', 'By giving examples',
             'By telling stories', 'By listing numbers']),
        (7, 'Which part of a magazine is the passage probably taken from?',
            ['History', 'Business', 'Science', 'Culture']),
    ]
    for num, stem, opts in p2_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①形容词副词覆盖比较级/最高级变化+句型+语序，共8道MC。'
        '②Passage 1（电子废弃物回收）来源2024北京中考，答案：1-C 2-D 3-B。'
        '③Passage 2（宠物饮食科学研究）来源2024浙江中考，答案：4-B 5-A 6-D 7-C。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L10_v3.docx')
    print('L10 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L11: 一轮复习·短文填空+主题写作2（青少年成长·问题建议）
# ═══════════════════════════════════════════════════════════════

def gen_L11():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Eleven    一轮复习·短文填空精讲 + 主题写作精讲 2（问题建议类）', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 短文填空精讲 ─────────────────────────
    part_head(doc, 'PART ONE   短文填空精讲——青少年成长话题')

    section_head(doc, '①', '短文填空高频考点回顾', tag='重点')
    gt(doc,
        ['考点', '占比', '典型考法'],
        [
            ['动词时态语态', '30%', '根据时间状语/上下文判断时态，被动语态识别'],
            ['词形变换',     '25%', '名词单复数、形副互换、动词三单/过去式/过去分词'],
            ['代词/冠词',    '15%', '指代关系判断、a/an/the选择'],
            ['连词/介词',    '15%', '并列/转折/因果连词、固定搭配介词'],
            ['首字母填空',   '15%', '根据语境和首字母推断词汇'],
        ],
        col_widths=[3.5, 2, 11]
    )

    section_head(doc, '②', '易错陷阱提醒', tag='易错点')
    gt(doc,
        ['陷阱类型', '错误示例', '正确做法'],
        [
            ['忘变词形', 'He enjoy reading. → enjoys', '注意主谓一致，三单加s/es'],
            ['时态混乱', '混用has done和had done',      '找准时间参照点再判断'],
            ['被动遗漏', '误用主动代替被动',             '主语是动作承受者→用被动'],
            ['冠词遗漏', '误用 play piano → the piano', '乐器前加the，球类/三餐前不加'],
        ],
        col_widths=[3, 5.5, 8]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=3)

    section_head(doc, '③', '短文填空精练（青少年成长话题）')
    red_note(doc,
        '本讲短文须替换为26一模真题，话题：青少年成长/心理健康/同伴关系。'
        '五维评估：话题✓ 考频✓ 难度A+✓ 考点覆盖✓ 语篇连贯✓'
    )
    doc.add_paragraph()

    source_note(doc, '来源：2024年河北省中考英语真题，语法填空')
    doc.add_paragraph()

    fill_paras = [
        'The Lin family has a son and a daughter, Lin Ming and Lin Fang. Lin Ming is seven and '
        'Lin Fang is two years   (1)   (young) than her brother. They study in different '
        '  (2)   (school). One day, Mrs. Lin picked up Lin Ming and then they went to Lin Fang\'s '
        'school together.',
        'Some boys and girls were playing   (3)   (happy) on the playground. Lin Ming stood at the '
        'school gate and waited   (4)   his sister. Just then, the bell   (5)   (ring).',
        'Several minutes later, many   (6)   (child) walked out in a line. Lin Fang was the '
        '  (7)   (five) one in it. When Lin Ming saw her, he couldn\'t wait to run to her.',
        '"Fangfang," Lin Ming   (8)   (shout). Jumping with joy, Lin Fang rushed to meet '
        '  (9)   (he). They hugged each other and walked to their mother hand in hand. When '
        'Mrs. Lin saw this,   (10)   big smile appeared on her face.',
    ]
    for para in fill_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    # ── PART TWO 主题写作精讲2 ─────────────────────────
    part_head(doc, 'PART TWO   主题写作精讲 2——问题建议类')

    section_head(doc, '①', '问题建议类写作特征')
    body_text(doc, '话题方向：学业压力/同伴关系/沉迷手机/亲子沟通/时间管理等')
    body_text(doc, '写作框架：描述问题 → 分析原因 → 提出建议 → 总结展望')
    doc.add_paragraph()

    section_head(doc, '②', '问题建议类高频表达', tag='必背')
    gt(doc,
        ['功能', '高频句式'],
        [
            ['描述问题', 'Many students have difficulty (in) doing ... / ... is becoming a serious problem among teenagers.'],
            ['分析原因', 'The reason is that ... / One possible reason is ... / This is mainly because ...'],
            ['提出建议', 'I think you should ... / Why not ...? / You\'d better ... / It would be helpful to ...'],
            ['回信开头', 'I\'m sorry to hear that ... / I\'m writing to give you some advice about ...'],
            ['鼓励结尾', 'I believe you can ... / I hope things will get better. / Don\'t give up!'],
            ['求助表达', 'Could you please give me some advice on ...? / I don\'t know what to do about ...'],
        ],
        col_widths=[3, 13.5]
    )

    section_head(doc, '③', '建议类满分结构', tag='大招')
    gt(doc,
        ['段落', '内容', '句数', '示例开头'],
        [
            ['开头段', '表示理解+引出问题', '1-2句', 'I\'m sorry to hear that you\'re having trouble with ...'],
            ['主体段', '2-3条具体建议',      '4-6句', 'First, you should ... Besides, it\'s a good idea to ... What\'s more, ...'],
            ['结尾段', '鼓励+祝福',          '1-2句', 'I believe you can work it out. / I hope my advice will be helpful.'],
        ],
        col_widths=[2.5, 4, 2, 8]
    )
    doc.add_paragraph()

    section_head(doc, '④', '写作实战训练')
    source_note(doc, '来源：2024年各省中考英语书面表达·压力与建议类，综合改编')
    doc.add_paragraph()
    body_text(doc, '假设你是李华，你的好朋友 Tom 最近因为学业压力大而感到焦虑，经常失眠，上课疲倦。'
              '请你给他写一封建议信，帮助他走出困境。')
    body_text(doc, '内容要求：', indent=0.3)
    body_text(doc, '1. 表示理解和关心', indent=0.6)
    body_text(doc, '2. 至少提出两条具体建议（如：合理安排作息、运动放松、与人沟通等）', indent=0.6)
    body_text(doc, '3. 表达期望和鼓励', indent=0.6)
    body_text(doc, '要求：80-100词，格式正确，语句通顺。开头和结尾已给出，不计入总词数。')
    doc.add_paragraph()
    body_text(doc, 'Dear Tom,')
    body_text(doc, "I'm sorry to hear that you are having trouble these days. ___________________")
    body_text(doc, '___________________________________________________________________________')
    body_text(doc, '___________________________________________________________________________')
    body_text(doc, '___________________________________________________________________________')
    body_text(doc, 'Yours,')
    body_text(doc, 'Li Hua')
    doc.add_paragraph()

    section_head(doc, '⑤', '建议类/问题解决类写作高频表达补充', tag='拓展')
    gt(doc,
        ['表达', '用法'],
        [
            ['have difficulty/trouble (in) doing sth', '做某事有困难'],
            ['be under great pressure',                '承受巨大压力'],
            ['communicate with sb',                    '与某人沟通'],
            ['make a plan for ...',                    '为……制定计划'],
            ['balance ... and ...',                    '平衡……和……'],
            ['take a break / have a rest',             '休息一下'],
            ['set a goal',                             '设定目标'],
            ['ask sb for help',                        '向某人求助'],
        ],
        col_widths=[7, 9.5]
    )

    red_note(doc,
        '【教师版备注】'
        '①短文填空来源2024河北中考，答案：1-younger 2-schools 3-happily 4-for 5-rang 6-children 7-fifth 8-shouted 9-him 10-a。'
        '②写作题已替换为2024年中考真题改编（压力与建议类），课堂先讲框架+表达，再限时15分钟写作。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L11_v3.docx')
    print('L11 v3 done.')


# ═══════════════════════════════════════════════════════════════
#  L12: 一轮复习·形容词副词2 + 中考词汇精讲5 + 综合检测
# ═══════════════════════════════════════════════════════════════

def gen_L12():
    doc = new_doc()

    title_line(doc, '26秋上  初三英语讲义  A+班（全国通用版）', TITLE_SIZE)
    title_line(doc, 'Lesson Twelve    一轮复习·语法（形容词副词进阶）+ 中考词汇精讲 5 + 综合模拟', Pt(12))
    doc.add_paragraph()

    # ── PART ONE 形容词副词进阶 ─────────────────────────
    part_head(doc, 'PART ONE   语法精讲——形容词 & 副词进阶用法')

    section_head(doc, '①', 'easy / difficult / hard + to do 结构', tag='重点')
    body_text(doc, '结构：主语(事/物) + be + adj + to do')
    body_text(doc, '特点：不定式的动词与主语构成逻辑上的动宾关系，不定式用主动形式表被动含义。')
    gt(doc,
        ['句型', '例句', '注意点'],
        [
            ['sth + is easy to do',      'The question is easy to answer.',     'answer后不加it，主语即宾语'],
            ['sth + is difficult to do', 'The problem is difficult to solve.',  'solve后不加it'],
            ['sth + is hard to do',      'The book is hard to understand.',     'understand后不加it'],
            ['sth + is pleasant to do',  'The song is pleasant to listen to.',  '介词to不能省'],
            ['sth + is impossible to do','The task is impossible to finish in one day.', ''],
        ],
        col_widths=[4.5, 5.5, 6.5]
    )
    red_note(doc, '易错：The question is easy to answer it. (✗) → to answer (✓)，不可重复宾语。')

    section_head(doc, '②', 'enough 的用法', tag='考频极高')
    gt(doc,
        ['结构', '位置', '例句'],
        [
            ['adj/adv + enough', 'enough放在形容词/副词后面',  'He is old enough to go to school.'],
            ['enough + n.',      'enough放在名词前面',         'We have enough time to prepare.'],
            ['not + adj + enough','不够……',                    'He is not tall enough to reach the shelf.'],
        ],
        col_widths=[4.5, 5, 7]
    )

    section_head(doc, '③', 'too ... to ... 结构', tag='重点')
    gt(doc,
        ['结构', '含义', '例句'],
        [
            ['too + adj/adv + to do',        '太……以至于不能……',  'He is too young to go to school.'],
            ['too + adj + for sb + to do',   '对某人来说太……',     'The box is too heavy for me to carry.'],
            ['= so ... that ... (not)',      '同义转换',           'He is so young that he can\'t go to school.'],
            ['= not ... enough to do',       '同义转换',           'He is not old enough to go to school.'],
        ],
        col_widths=[5, 4, 7.5]
    )

    section_head(doc, '④', '三组结构同义转换', tag='大招')
    body_text(doc, '掌握以下三组同义句转换是中考考点精粹：')
    gt(doc,
        ['原句', '转换1', '转换2'],
        [
            ['He is too tired to walk.',
             'He is so tired that he can\'t walk.',
             'He is not energetic enough to walk.'],
            ['The ice is thick enough to skate on.',
             'The ice is so thick that we can skate on it.',
             'The ice is not too thin to skate on.'],
        ],
        col_widths=[5.5, 5.5, 5.5]
    )
    doc.add_paragraph()
    dazao_box(doc, stars=4)

    ex_label(doc, '例题精讲（真题选编·形容词副词进阶中考高频考法）')
    source_note(doc, '选自各省2023-2025年中考英语真题及模拟题')

    adj2_qs = [
        (1,
         'The little boy is not _______ to dress himself.',
         ['old enough', 'enough old', 'too old', 'old too']),
        (2,
         'The water in the river is too dirty _______.',
         ['to drink', 'to drink it', 'to be drunk', 'drinking']),
        (3,
         'The math problem was _______ difficult _______ most students couldn\'t work it out.',
         ['too; to', 'so; that', 'such; that', 'enough; to']),
        (4,
         'English is not difficult _______. We just need more practice.',
         ['learn', 'to learn', 'learning', 'learned']),
        (5,
         'The box was _______ heavy for the child _______ carry.',
         ['too; to', 'so; that', 'very; to', 'enough; to']),
        (6,
         'Do we have _______ volunteers for the school sports meeting?',
         ['enough', 'too', 'much', 'many enough']),
        (7,
         'The song sounds _______ pleasant _______ everyone loves it.',
         ['too; to', 'enough; to', 'so; that', 'such; that']),
        (8,
         'She speaks English well _______ to communicate with foreign visitors.',
         ['too', 'so', 'enough', 'very']),
    ]
    for num, stem, opts in adj2_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    # ── PART TWO 词汇精讲5 ─────────────────────────
    part_head(doc, 'PART TWO   中考词汇精讲 5——词根记忆法')

    section_head(doc, '①', '词根 scrib / script（写）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['describe',    'v.',   '描述；描写',       'de(向下) + scrib(写) + e → 写下来→描述'],
            ['subscribe',   'v.',   '订阅；签署',       'sub(下面) + scrib(写) + e → 在下面签名→订阅'],
            ['prescribe',   'v.',   '开处方；规定',     'pre(预先) + scrib(写) + e → 预先写好→开处方'],
            ['description', 'n.',   '描述；说明',       'describe(描述) + tion → 描述的内容'],
            ['script',      'n.',   '剧本；手稿；脚本', 'script(写) → 写出来的东西→剧本'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '②', '词根 form（形式）')
    gt(doc,
        ['单词', '词性', '核心含义', '词根拆解'],
        [
            ['form',      'n./v.', '形式；形成',       'form(形式) → 形成某种形式'],
            ['reform',    'v./n.', '改革',             're(再) + form(形式) → 再次成形→改革'],
            ['inform',    'v.',    '通知；告知',       'in(进入) + form(形式) → 把信息放入形式→通知'],
            ['perform',   'v.',    '表演；执行',       'per(完全) + form(形式) → 完全呈现→表演'],
            ['uniform',   'n.',    '制服',             'uni(统一) + form(形式) → 统一的形式→制服'],
            ['transform', 'v.',    '转变；改造',       'trans(跨越) + form(形式) → 跨越形式→转变'],
        ],
        col_widths=[3, 2, 4, 7.5]
    )

    section_head(doc, '③', '25-26年考频 TOP 10 高频词', tag='重点')
    gt(doc,
        ['单词', '词性', '释义', '25-26年中考典型语境'],
        [
            ['review',     'v./n.', '复习；回顾',       'review for the final exam（复习备考）'],
            ['standard',   'n.',    '标准；水平',       'meet the standard / high standard（达到标准）'],
            ['conclusion', 'n.',    '结论',             'draw/reach a conclusion（得出结论）'],
            ['method',     'n.',    '方法',             'learning methods / scientific method（学习/科学方法）'],
            ['various',    'adj.',  '各种各样的',       'various activities / for various reasons（各种活动/原因）'],
            ['contain',    'v.',    '包含；容纳',       'contain important information（包含重要信息）'],
            ['mention',    'v.',    '提到；提及',       'as mentioned above / mention the topic（如上所述）'],
            ['compare',    'v.',    '比较',             'compare A with B / compared with ...（与……相比）'],
            ['improve',    'v.',    '改善；提高',       'improve our learning efficiency（提高学习效率）'],
            ['complete',   'v./adj.','完成；完整的',    'complete the task on time（按时完成任务）'],
        ],
        col_widths=[3, 2, 3.5, 8]
    )

    section_head(doc, '④', '词汇语境练习', tag='换新题 · 26一模风格')
    body_text(doc, '用括号内所给词的适当形式填空：')
    fills = [
        '1. Could you give a brief _______ of what happened at the meeting yesterday?  (describe)',
        '2. The students are required to wear school _______ on weekdays.  (uniform)',
        '3. _______ with last year, the air quality in our city has greatly improved.  (compare)',
        '4. The teacher asked us to _______ what we had learned before the exam.  (review)',
        '5. The box _______ several old letters and some photos from the 1990s.  (contain)',
    ]
    for f in fills:
        body_text(doc, f, indent=0.6)
    red_note(doc, '语境填空务必替换为26一模真题原句。词根第9-10组(scrib/script + form)共11词，全部词根完结。')
    doc.add_paragraph()

    # ── PART THREE 综合模拟 ─────────────────────────
    part_head(doc, 'PART THREE   阶段综合检测（L1-L12 全覆盖）')

    section_head(doc, '①', '综合模拟说明')
    gt(doc,
        ['维度', '说明'],
        [
            ['检测范围', 'L1-L12 全部语法考点 + 5组词根词汇 + 阅读写作技巧'],
            ['题型配置', '单选15题 + 完形1篇 + 阅读2篇 + 短文填空1篇 + 书面表达1篇'],
            ['时长建议', '90分钟（课堂检测60分钟 + 课后补充30分钟）'],
            ['难度定位', '对标26年中考A+水平，含10%拔高题'],
        ],
        col_widths=[3, 13.5]
    )
    doc.add_paragraph()

    section_head(doc, '②', '语法综合单选', tag='L1-L12 全覆盖')
    red_note(doc,
        '本讲综合检测须用26一模真题或改编题组卷。以下为格式示例。'
        '覆盖考点：过去完成时、时态辨析、感叹句、被动语态、宾语从句、定语从句、代词、介词、形容词副词。'
    )
    doc.add_paragraph()

    comp_qs = [
        (1,
         'By the time I arrived at the airport, the plane _______.',
         ['has taken off', 'took off', 'had taken off', 'was taking off']),
        (2,
         '_______ wonderful speech Emma gave at the graduation ceremony!',
         ['What', 'What a', 'How', 'How a']),
        (3,
         'The 2026 Asian Games _______ in Aichi-Nagoya, Japan next year.',
         ['is held', 'will be held', 'was held', 'has been held']),
        (4,
         'Could you tell me _______ the nearest bookstore?',
         ['how can I get to', 'how I can get to', 'how to get', 'how I get to']),
        (5,
         'The girl _______ won the science competition is from our school.',
         ['which', 'whom', 'who', 'whose']),
        (6,
         '—I forgot to bring my dictionary. May I use _______?\n—Sure, here you are.',
         ['you', 'your', 'yours', 'yourself']),
        (7,
         'We\'ve been friends _______ we were in primary school.',
         ['for', 'since', 'when', 'while']),
        (8,
         'The soup is _______ hot _______ drink right now. Let it cool down.',
         ['too; to', 'so; that', 'enough; to', 'very; to']),
    ]
    for num, stem, opts in comp_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    section_head(doc, '②', '完形填空')
    source_note(doc, '来源：2024年广东省广州市中考英语真题，完形填空（完整版）')
    doc.add_paragraph()

    l12_cloze_paras = [
        '"Hello, Mr Li," I said sweetly. "I came to ask if I could help paint the posters for '
        'the Reading Day. I don\'t mind   (16)   for the whole day at the library tomorrow."',
        '"Aren\'t you going to the old people\'s home?" Mr Li asked. "My class is going. But I '
        'could have a/an   (17)   if you need my help."',
        '"Thanks, Janet. But I think you could come after school tomorrow."   (18)  , there was '
        'no way out. I had to go.',
        'The next day, when my classmates were talking with the old people, I really felt nervous '
        'to   (19)   any of them. I sat there on my own holding the card I made the night before. '
        'Across the room there was an old lady in a wheelchair. She was   (20)  , too. Maybe I '
        'could hand my card to her. It might brighten her day.',
        'So carefully, I went to her... "Thank you, Janet." The old lady reached for the card but '
        'her hand was nowhere near it. I suddenly realized that she was   (21)  .',
        '"What does your card look like?" the old lady asked. I   (22)   my card patiently. '
        'Her fingers touched every inch of it. She couldn\'t enjoy more...',
        'Too soon, it was time to return. But I didn\'t want to   (23)  . "Thanks for your coming, '
        'Janet! I\'ll keep this card to remember you." "I\'m sorry you can\'t see it. I wished I '
        'had brought you a better gift," I said.',
        '"The best gift," the old lady said, "was your   (24)  , Janet."',
        'With tears, I felt light and warm. I couldn\'t wait to come back to see my new '
        '  (25)   again.',
    ]
    for para in l12_cloze_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    l12_cloze_qs = [
        (16, 'I don\'t mind _______ for the whole day.',
             ['reading', 'discussing', 'relaxing', 'working']),
        (17, 'I could have a/an _______.',
             ['idea', 'lesson', 'excuse', 'gift']),
        (18, '_______, there was no way out.',
             ['Clearly', 'Strangely', 'Luckily', 'Suddenly']),
        (19, 'I felt nervous to _______ any of them.',
             ['accept', 'join', 'invite', 'guide']),
        (20, 'She was _______, too.',
             ['alone', 'active', 'ready', 'busy']),
        (21, 'I suddenly realized that she was _______.',
             ['tired', 'blind', 'serious', 'worried']),
        (22, 'I _______ my card patiently.',
             ['changed', 'compared', 'described', 'prepared']),
        (23, 'I didn\'t want to _______.',
             ['forget', 'leave', 'promise', 'imagine']),
        (24, 'The best gift was your _______.',
             ['return', 'plan', 'class', 'visit']),
        (25, 'I couldn\'t wait to see my new _______ again.',
             ['friend', 'teacher', 'classmate', 'neighbour']),
    ]
    for num, stem, opts in l12_cloze_qs:
        mc(doc, num, stem, opts)

    section_head(doc, '③', '阅读理解 + 短文填空 + 书面表达')
    red_note(doc,
        '以上三大题型须从26一模真题中组卷。\n'
        '阅读2篇（1篇记叙文+1篇说明文）+ 短文填空1篇 + 书面表达1篇（建议/环保/成长任选）。\n'
        '总分值建议：单选30分 + 完形15分 + 阅读30分 + 短文填空10分 + 写作15分 = 100分。'
    )
    doc.add_paragraph()

    red_note(doc,
        '【教师版备注】'
        '①形容词副词进阶覆盖easy/difficult to do、enough、too...to...及同义转换。'
        '②词根第9-10组(scrib/script + form)共11词，至此全部35个词根完结。'
        '②完形填空来源2024广州中考（Janet养老院，完整版10空），答案：16-D 17-C 18-A 19-B 20-A 21-B 22-C 23-B 24-D 25-A。'
        '③综合检测其余题型须用26一模真题组卷。'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L12_v3.docx')
    print('L12 v3 done.')


if __name__ == '__main__':
    gen_L1()
    gen_L2()
    gen_L3()
    gen_L4()
    gen_L5()
    gen_L6()
    gen_L7()
    gen_L8()
    gen_L9()
    gen_L10()
    gen_L11()
    gen_L12()
    print('全部完成。')
