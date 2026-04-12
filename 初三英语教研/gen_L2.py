import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    doc.sections[0].page_width  = Cm(21)
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].left_margin = doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin  = doc.sections[0].bottom_margin = Cm(2)
    return doc

def cell_bg(cell, hex6):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    pr.append(shd)

def banner(doc, label, zh, en, bg='E8672A'):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; cell_bg(c, bg)
    r = c.paragraphs[0].add_run(f'  {label}   {zh}   {en}  ')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    doc.add_paragraph()

def badge(doc, text, bg='F5A623'):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; cell_bg(c, bg)
    r = c.paragraphs[0].add_run(f'  {text}  ')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    doc.add_paragraph()

def sh(doc, num, title, tag=None):
    p = doc.add_paragraph()
    rn = p.add_run(f'{num} '); rn.bold = True; rn.font.color.rgb = RGBColor(0xE8,0x67,0x2A)
    p.add_run(title).bold = True
    if tag:
        rx = p.add_run(f'  [{tag}]'); rx.bold = True; rx.font.color.rgb = RGBColor(0xFF,0x4D,0x4D)

def gt(doc, headers, rows, hbg='C0A0FF'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers)); tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; cell_bg(c, hbg)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            c.paragraphs[0].add_run(str(v)).font.size = Pt(10)
    doc.add_paragraph()

def dazao(doc, n=3):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; cell_bg(c, 'FFF3E0')
    c.paragraphs[0].add_run(
        f'学习大招   难度系数：{"★"*n}{"☆"*(5-n)}\n\n动手，写下【大招】吧！\n\n\n'
    ).font.size = Pt(10)
    doc.add_paragraph()

def mc(doc, num, stem, opts):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    p.add_run(f'{num}. {stem}').font.size = Pt(10.5)
    op = doc.add_paragraph(); op.paragraph_format.left_indent = Cm(0.8)
    op.add_run('    '.join([f'{lb}. {o}' for lb, o in zip('ABCD', opts)])).font.size = Pt(10.5)

def body(doc, text):
    p = doc.add_paragraph(text); p.paragraph_format.left_indent = Cm(0.5)
    if p.runs: p.runs[0].font.size = Pt(10.5)

def exlabel(doc):
    p = doc.add_paragraph()
    r = p.add_run('▶ 例题精讲'); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xE8,0x67,0x2A)

def rednote(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0xCC,0x44,0x44)

# ─── L2 正文 ──────────────────────────────────────────────────
doc = new_doc()

tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('26秋上 初三英语讲义 A+班（全国通用版）')
tr.bold = True; tr.font.size = Pt(18)

sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run('Lesson Two   语法（易混时态辨析）+ 中考词汇精讲 1').font.size = Pt(12)
doc.add_paragraph()

# PART ONE
banner(doc, 'PART ONE', '语法 —— 易混时态辨析', 'Grammar', 'E8672A')
badge(doc, '知识精讲', 'F5A623')
doc.add_paragraph()

sh(doc, '①', '四大时态核心对比', tag='重点')
gt(doc,
    ['时态', '结构', '核心含义', '标志词（高频）'],
    [
        ['一般过去时',   'did',            '过去某时发生，已结束',           'yesterday, ago, last, in+过去年份'],
        ['过去进行时',   'was/were doing', '过去某时正在进行',               'at that time, when引导的从句'],
        ['现在完成时',   'have/has done',  '过去动作对现在有影响/持续至今',  'already, yet, ever, just, for, since'],
        ['过去完成时',   'had done',       '"过去的过去"，先于过去某时完成', 'by, before, when, after+过去时从句'],
    ]
)

sh(doc, '②', '辨析1：一般过去时 vs 现在完成时', tag='考频极高')
gt(doc,
    ['对比维度', '一般过去时', '现在完成时'],
    [
        ['时间焦点', '过去某一具体时间点',            '过去动作→现在结果/影响'],
        ['时间状语', 'yesterday / last year / in 2020','already / yet / ever / just / for / since'],
        ['典型句',   'I saw the film last week.',      'I have seen the film.（现在了解剧情）'],
        ['陷阱提示', '含具体过去时间→必用过去时',     'for/since+时间段→必用现在完成时'],
    ],
    'A0D0FF'
)

sh(doc, '③', '辨析2：过去进行时 vs 一般过去时', tag='常考')
gt(doc,
    ['对比维度', '过去进行时', '一般过去时'],
    [
        ['强调重点', '过去某时段正在持续（未结束）',      '过去某时点动作（已完成）'],
        ['结构',     'was/were + doing',                 'did（动词过去式）'],
        ['when从句', '主句用进行时（背景动作）',          '主句用过去时（先后顺序）'],
        ['例句',     'I was reading when he called.',    'I read the book last night.'],
    ],
    'A0D0FF'
)

sh(doc, '④', '辨析3：现在完成时 vs 过去完成时', tag='重点')
gt(doc,
    ['对比维度', '现在完成时', '过去完成时'],
    [
        ['参照时间', '以"现在"为参照',           '以"过去某时"为参照（过去的过去）'],
        ['结构',     'have/has + done',           'had + done'],
        ['例句',     'She has finished her work.','She had finished her work before dinner.'],
        ['信号词',   '强调现在状态/结果',         '主从句均在过去，需分出先后顺序'],
    ],
    'A0D0FF'
)

sh(doc, '⑤', '时态选择三步口诀', tag='大招')
body(doc, '第一步：找时间状语 → 锁定时间范围（过去/现在/过去的过去）')
body(doc, '第二步：判完成状态 → 动作是否对现在有影响 / 是否先于另一过去动作')
body(doc, '第三步：看主从关系 → 两个过去动作谁先谁后，先发生的用过去完成时')
doc.add_paragraph()

dazao(doc, n=4)
exlabel(doc)

questions_L2 = [
    (1,
     'By the time the rescue team arrived, the trapped villagers _______ to a safer place.',
     ['moved', 'have moved', 'had moved', 'were moving']),
    (2,
     '—Did you watch the AI documentary last night?\n—No, I _______ it already before it was broadcast on TV.',
     ['watched', 'have watched', 'had watched', 'was watching']),
    (3,
     'The young scientist _______ on the climate project for three years when she finally got the key result.',
     ['has worked', 'had been working', 'worked', 'was working']),
    (4,
     'When I got to the community centre, the volunteers _______ rubbish along the river for over two hours.',
     ['collected', 'have collected', 'had collected', 'had been collecting']),
    (5,
     '—I cannot find my digital student ID. Have you seen it?\n—Yes, you _______ it on the desk when you came in just now.',
     ['put', 'have put', 'had put', 'were putting']),
    (6,
     'The intangible heritage master said he _______ this traditional craft from his grandfather when he was only eight.',
     ['learned', 'has learned', 'had learned', 'was learning']),
    (7,
     'It was already midnight. The smart city streets _______ quiet, and only a few self-driving cars _______ past.',
     ['grew / drove', 'had grown / drove', 'grew / had driven', 'had grown / had driven']),
    (8,
     'Since the new green energy policy _______ into effect, carbon emissions in the city _______ by nearly 30%.',
     ['came / dropped', 'had come / dropped', 'came / have dropped', 'has come / has dropped']),
]
for num, stem, opts in questions_L2:
    mc(doc, num, stem, opts)
doc.add_paragraph()

# PART TWO 词汇精讲1
banner(doc, 'PART TWO', '中考词汇精讲 1', 'Vocabulary', '2E7D9E')
badge(doc, '词根记忆法', '2E7D9E')
doc.add_paragraph()

sh(doc, '①', '词根 vis / vid（看）')
gt(doc,
    ['单词', '词性', '核心含义', '词根拆解记忆'],
    [
        ['visible',   'adj.', '看得见的；明显的', 'vis(看) + ible → 能被看到'],
        ['vision',    'n.',   '视力；视野；愿景', 'vis(看) + ion → 看的能力'],
        ['revise',    'v.',   '修改；复习',       're(再) + vis(看) → 再看一遍'],
        ['evidence',  'n.',   '证据；迹象',       'e + vid(看) + ence → 被看到的东西'],
        ['supervise', 'v.',   '监督；管理',       'super(上面) + vis(看) → 从上面盯着看'],
    ],
    '2E7D9E'
)

sh(doc, '②', '词根 port（运；携带）')
gt(doc,
    ['单词', '词性', '核心含义', '词根拆解记忆'],
    [
        ['transport', 'v./n.', '运输；交通工具', 'trans(跨越) + port(运) → 运过去'],
        ['export',    'v./n.', '出口；输出',     'ex(出) + port(运) → 运出去'],
        ['import',    'v./n.', '进口；输入',     'im(进入) + port(运) → 运进来'],
        ['report',    'v./n.', '报告；汇报',     're(回) + port(运) → 把消息运回来'],
        ['support',   'v./n.', '支持；支撑',     'sup(在下) + port(运) → 在下面撑着'],
        ['portable',  'adj.',  '便携的；手提的', 'port(运) + able → 可以带着走的'],
    ],
    '2E7D9E'
)

sh(doc, '③', '25-26年考频 TOP 10 高频词', tag='重点')
gt(doc,
    ['单词', '词性', '释义', '25-26年中考典型语境话题'],
    [
        ['achieve',     'v.',    '实现；达到',         '科技突破：achieve a breakthrough in AI research'],
        ['attempt',     'v./n.', '尝试；努力',         '青少年挑战自我：attempt to climb the mountain'],
        ['contribute',  'v.',    '贡献；有助于',       '环保：contribute to reducing plastic pollution'],
        ['influence',   'v./n.', '影响；作用',         '榜样人物：have a great influence on young people'],
        ['recognize',   'v.',    '识别；认出；承认',   'AI识别：recognize faces with smart technology'],
        ['appreciate',  'v.',    '感激；欣赏；理解',   '传统文化：appreciate the beauty of Chinese opera'],
        ['determine',   'v.',    '决定；使下定决心',   '人物传记：be determined to make a difference'],
        ['encourage',   'v.',    '鼓励；促进',         '教育：encourage students to think independently'],
        ['struggle',    'v./n.', '奋斗；挣扎',         '励志：struggle against all kinds of difficulties'],
        ['transform',   'v.',    '转变；彻底改变',     '城市发展：transform the old community into a park'],
    ],
    'A0C4FF'
)

sh(doc, '④', '词汇语境练习', tag='换新题 · 26一模风格')
body(doc, '用括号内所给词的适当形式填空：')
vocab_fills = [
    '1. The project team finally _______ their goal of building a zero-carbon school.  (achieve)',
    '2. With AI tools, the system can now _______ handwritten Chinese characters accurately.  (recognition)',
    '3. She was so _______ to win that she trained six hours a day for months.  (determine)',
    '4. Young people should learn to _______ the traditional art forms passed down from ancestors.  (appreciation)',
    '5. The volunteers _______ greatly to improving the water quality of the local river last year.  (contribute)',
]
for f in vocab_fills:
    body(doc, f)

rednote(doc, '【教师版备注】语境填空务必换26一模真题原句。词根第1-2组共11词，下讲继续第3组。')

doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L2.docx')
print('L2 done.')
