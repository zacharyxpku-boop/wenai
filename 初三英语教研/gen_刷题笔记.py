"""
刷题笔记批量生成：3份Word文档
1. 2024浙江省中考英语真题精析.docx
2. 2024北京中考英语阅读B-C篇精读.docx
3. 2024各省感叹句+过去完成时真题20道分类归档.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TNR = 'Times New Roman'
TITLE_SIZE   = Pt(16)
PART_SIZE    = Pt(13)
SECTION_SIZE = Pt(11)
BODY_SIZE    = Pt(11)
NOTE_SIZE    = Pt(9)

# ── helper functions (same pattern as gen_all.py) ──

def new_doc():
    doc = Document()
    doc.sections[0].page_width   = Cm(21)
    doc.sections[0].page_height  = Cm(29.7)
    doc.sections[0].left_margin  = Cm(2.5)
    doc.sections[0].right_margin = Cm(2.5)
    doc.sections[0].top_margin   = Cm(2.5)
    doc.sections[0].bottom_margin= Cm(2)
    doc.styles['Normal'].font.name = TNR
    doc.styles['Normal'].font.size = BODY_SIZE
    return doc

def set_font(run, bold=False, size=BODY_SIZE, color=None, italic=False):
    run.font.name = TNR
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def title_line(doc, text, size=TITLE_SIZE, center=True):
    p = doc.add_paragraph()
    if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_font(r, bold=True, size=size)
    return p

def part_head(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    set_font(r, bold=True, size=PART_SIZE, color='1F4E79')
    p2 = doc.add_paragraph()
    p2.add_run('\u2500' * 55).font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def section_head(doc, num_text, title, tag=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f'{num_text} {title}')
    set_font(r1, bold=True, size=SECTION_SIZE)
    if tag:
        r2 = p.add_run(f'  \u2605{tag}')
        set_font(r2, bold=True, size=Pt(9), color='C00000')

def body_text(doc, text, indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_font(r, size=BODY_SIZE)
    return p

def tip_text(doc, text, indent=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    set_font(r, size=BODY_SIZE, italic=False)

def red_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'\u3010\u6ce8\u3011{text}')
    set_font(r, size=NOTE_SIZE, color='C00000', italic=True)

def source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'[\u6765\u6e90] {text}')
    set_font(r, size=NOTE_SIZE, color='595959', italic=True)

def _set_cell_bg(cell, hex6):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    pr.append(shd)

def gt(doc, headers, rows, col_widths=None):
    ncols = len(headers)
    tbl = doc.add_table(rows=1+len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _set_cell_bg(c, 'D6E4F0')
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, bold=True, size=Pt(10))
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            p = c.paragraphs[0]
            r = p.add_run(str(v)); set_font(r, size=Pt(10))
    if col_widths:
        for ri in range(len(tbl.rows)):
            for ci, w in enumerate(col_widths):
                tbl.rows[ri].cells[ci].width = Cm(w)
    doc.add_paragraph()

def passage_text(doc, text, indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(0.7)
    r = p.add_run(text)
    set_font(r, size=BODY_SIZE)

def mc(doc, num, stem, opts, labels='ABCD', indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(f'{num}. {stem}')
    set_font(r, size=BODY_SIZE)
    op = doc.add_paragraph()
    op.paragraph_format.left_indent = Cm(indent + 0.4)
    op.paragraph_format.space_before = Pt(0)
    op.add_run('    '.join([f'{lb}. {o}' for lb, o in zip(labels, opts)])).font.size = BODY_SIZE

def highlight_box(doc, text, bg_color='FFFBF0', border_color='7F3F00'):
    """Single-cell highlight box for summary content"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.rows[0].cells[0]
    _set_cell_bg(c, bg_color)
    p = c.paragraphs[0]
    r = p.add_run(text)
    set_font(r, size=Pt(10), bold=True, color=border_color)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
#  DOC 1: 2024浙江省中考英语真题精析
# ═══════════════════════════════════════════════════════════════

def gen_doc1():
    doc = new_doc()

    title_line(doc, '2024\u6d59\u6c5f\u7701\u4e2d\u8003\u82f1\u8bed\u771f\u9898\u00b7\u5237\u9898\u7b14\u8bb0', TITLE_SIZE)
    doc.add_paragraph()

    # ── Section 1: 完形填空精析 ──
    part_head(doc, 'Section 1   \u5b8c\u5f62\u586b\u7a7a\u7cbe\u6790\uff08\u60c5\u611f\u6001\u5ea6\u8bcd\u8003\u6cd5\u6807\u6ce8\uff09')
    source_note(doc, '\u6765\u6e90\uff1a2024\u5e74\u6d59\u6c5f\u7701\u4e2d\u8003\u82f1\u8bed\u771f\u9898\uff0c\u7b2c\u4e09\u90e8\u5206\u8bed\u8a00\u8fd0\u7528\u7b2c\u4e00\u8282')
    doc.add_paragraph()

    section_head(doc, '\u2460', '\u539f\u6587\u901a\u8bfb\uff08Bruce & Mike \u6821\u8f66\u4ea4\u53cb\u6545\u4e8b\uff09')

    cloze_paras = [
        'Bruce loves music because it makes him feel easy. He is shy, and making friends has been '
        '  (1)   for him. That\'s why his mother asked him to take the   (2)   this year. '
        '"You\'d better not sit there with your headphones on," she told him.',
        'Bruce nodded. But the moment he sat on the bus, he   (3)   his headphones and closed his eyes. '
        'He was soon lost in a song of his favorite band, the Blue-Bob.   (4)  , he sensed someone around him. '
        'He opened his eyes and saw a boy from his   (5)  .',
        '"Hey, Bruce? I\'m Mike," the boy said. "I think we both take the music class. Did I hear '
        '  (6)   singing a song by the Blue-Bob just now?"',
        '"Uh, no," Bruce\'s face turned red, "That wasn\'t me." Mike looked a little   (7)  . '
        'Without saying anything, he went away.',
        'For the next few days, Bruce kept asking himself   (8)   he had lied to Mike. '
        '"Maybe he also loves the Blue-Bob. Maybe he was just being   (9)  ," he thought. '
        'So when they saw each other on the bus the next day, Bruce managed to   (10)   a smile on his face.',
        '"Hi, Bruce," Mike said. "I can hear your   (11)   every day, and you have such a good taste."',
        'Bruce\'s eyes lit up. "Don\'t you think the songs are too   (12)  ?" he asked. '
        '"Not at all," Mike answered. "I love songs of the past too. My grandpa is really into music, '
        'and he has   (13)   me to sing many of them."',
        'Bruce listened as Mike talked about his favorite songs. And they shared the headphones, '
        'singing softly along with the music   (14)   Bruce got off at his stop.',
        '"See you tomorrow!" Mike said.',
        'Bruce waved goodbye. "I must tell Mom how I   (15)   Mike and we had the best time...'
        'with headphones on," he smiled to himself.',
    ]
    for para in cloze_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    section_head(doc, '\u2461', '\u901015\u9898\u7cbe\u6790\u8868', tag='\u91cd\u70b9')
    gt(doc,
        ['\u9898\u53f7', '\u6b63\u786e\u7b54\u6848', '\u8003\u67e5\u7c7b\u578b', '\u89e3\u9898\u5173\u952e'],
        [
            ['1',  'B. hard',      '\u60c5\u611f\u6001\u5ea6\u8bcd\u00b7\u63cf\u8ff0\u6027\u683c\u56f0\u5883',
             'shy\u2192making friends has been hard\uff0c\u4e0a\u4e0b\u6587\u8bed\u4e49\u987a\u627f'],
            ['2',  'C. bus',       '\u573a\u666f\u63a8\u65ad\u00b7\u4e0a\u4e0b\u6587\u7ebf\u7d22',
             '\u540e\u6587sat on the bus\u590d\u73b0'],
            ['3',  'A. put on',    '\u52a8\u8bcd\u77ed\u8bed\u8fa8\u6790',
             'put on headphones\u56fa\u5b9a\u642d\u914d\uff0c\u4e0eclosed his eyes\u5e76\u5217'],
            ['4',  'D. Suddenly',  '\u526f\u8bcd\u00b7\u53d9\u4e8b\u8f6c\u6298\u4fe1\u53f7',
             '\u4ece\u95ed\u773c\u201c\u6c89\u6d78\u201d\u2192\u7a81\u7136\u611f\u77e5\u6709\u4eba\uff0c\u53d9\u4e8b\u7a81\u8f6c'],
            ['5',  'B. grade',     '\u540d\u8bcd\u00b7\u573a\u666f\u63a8\u65ad',
             '\u540c\u5e74\u7ea7\u540c\u5b66\uff0c\u4e0emusic class\u547c\u5e94'],
            ['6',  'C. you',       '\u4ee3\u8bcd\u00b7\u5bf9\u8bdd\u903b\u8f91',
             'Mike\u5f53\u9762\u95eeBruce\uff0c\u7b2c\u4e8c\u4eba\u79f0you'],
            ['7',  'D. surprised', '\u2605\u60c5\u611f\u6001\u5ea6\u8bcd\u00b7\u5373\u65f6\u53cd\u5e94',
             'Bruce\u5426\u8ba4\u2192Mike\u610f\u5916\uff0csurprised\u7b26\u5408\u201c\u672a\u9884\u6599\u201d\u8bed\u5883'],
            ['8',  'A. why',       '\u8fde\u63a5\u8bcd\u00b7\u5fc3\u7406\u6d3b\u52a8',
             'asking himself why he had lied\uff0c\u8ffd\u95ee\u64af\u8c0e\u52a8\u673a'],
            ['9',  'D. friendly',  '\u2605\u60c5\u611f\u6001\u5ea6\u8bcd\u00b7\u4eba\u7269\u6001\u5ea6\u63a8\u6d4b',
             'Maybe he was just being friendly\uff0c\u53cd\u601d\u540e\u8ba4\u4e3aMike\u53ea\u662f\u53cb\u597d'],
            ['10', 'A. wear',      '\u52a8\u8bcd\u642d\u914d\u00b7wear a smile',
             '\u2605\u60c5\u611f\u52a8\u4f5c\uff1awear a smile\u56fa\u5b9a\u8868\u8fbe'],
            ['11', 'B. music',     '\u540d\u8bcd\u00b7\u4e0a\u4e0b\u6587\u590d\u73b0',
             '\u5168\u6587\u4e3b\u9898\u8bcdmusic\u590d\u73b0\uff0c\u4e0eheadphones/singing\u547c\u5e94'],
            ['12', 'B. old',       '\u5f62\u5bb9\u8bcd\u00b7\u8bed\u5883\u63a8\u65ad',
             'songs of the past\u2192old\uff0c\u540e\u6587grandpa\u9a8c\u8bc1'],
            ['13', 'C. taught',    '\u52a8\u8bcd\u00b7\u7956\u5b59\u5173\u7cfb\u903b\u8f91',
             'grandpa taught me to sing\uff0c\u7956\u5b59\u4f20\u627f\u903b\u8f91'],
            ['14', 'C. until',     '\u8fde\u8bcd\u00b7\u65f6\u95f4\u903b\u8f91',
             'singing...until Bruce got off\uff0c\u4e00\u76f4\u5230\u4e0b\u8f66'],
            ['15', 'A. met',       '\u52a8\u8bcd\u00b7\u6545\u4e8b\u4e3b\u7ebf\u547c\u5e94',
             'how I met Mike\uff0c\u56de\u5e94\u5168\u6587\u6821\u8f66\u521d\u9047\u4e3b\u7ebf'],
        ],
        col_widths=[1.5, 3, 5, 7]
    )

    highlight_box(doc,
        '\u60c5\u611f\u6001\u5ea6\u8bcd\u8003\u67e5\u70b9\uff1a'
        '\u7b2c1\u9898(hard\u63cf\u8ff0\u56f0\u5883)\u3001\u7b2c7\u9898(surprised\u5373\u65f6\u53cd\u5e94)\u3001'
        '\u7b2c9\u9898(friendly\u6001\u5ea6\u63a8\u6d4b)\u3001\u7b2c10\u9898(wear a smile\u60c5\u611f\u52a8\u4f5c)\uff0c'
        '\u517127%\uff084\u5904/15\u9898\uff09'
    )

    # ── Section 2: 阅读还原精析 ──
    part_head(doc, 'Section 2   \u9605\u8bfb\u8fd8\u539f\u7cbe\u6790\uff08\u89e3\u9898\u903b\u8f91\u6807\u6ce8\uff09')
    source_note(doc, '\u6765\u6e90\uff1a2024\u5e74\u6d59\u6c5f\u7701\u4e2d\u8003\u82f1\u8bed\u771f\u9898\uff0c\u9605\u8bfb\u8fd8\u539f')
    doc.add_paragraph()

    section_head(doc, '\u2460', '\u539f\u6587\u901a\u8bfb\uff08People in the UK Say Sorry\uff09')

    restore_paras = [
        '  (31)   If a person burps at the dinner table or runs into someone, '
        'you\'ll probably hear "Oops, I\'m so sorry!"',
        '  (32)   US scientist Robert Trivers gave a possible answer in his studies. '
        'He said that we are polite because we need to belong and be included in society to live. '
        'So we behave well, help each other and say sorry, hoping that we can be treated well and be safe.',
        'In the UK, politeness could be, as the Great British Mag put it, '
        '"a throwback to the British class system."',
        '  (33)   Does this mean that when British people say sorry, they don\'t mean it?',
        '  (34)   But saying sorry has its uses. If you\'re ever in the UK, remember a little '
        '"sorry" can go a long way.',
    ]
    for para in restore_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    body_text(doc, '\u9009\u9879\uff1a')
    body_text(doc, 'A. Perhaps it does sometimes.', indent=0.6)
    body_text(doc, 'B. We say sorry to show people of a higher class, or level, respect, and to stay out of trouble.', indent=0.6)
    body_text(doc, 'C. But why do we say sorry so much?', indent=0.6)
    body_text(doc, 'D. I wonder when people in the UK say sorry.', indent=0.6)
    body_text(doc, 'E. We can\'t help saying sorry in the UK.', indent=0.6)
    doc.add_paragraph()

    section_head(doc, '\u2461', '\u8fd8\u539f\u903b\u8f91\u7cbe\u6790\u8868', tag='\u91cd\u70b9')
    gt(doc,
        ['\u7a7a\u683c', '\u7b54\u6848', '\u8fd8\u539f\u903b\u8f91'],
        [
            ['31', 'E', '\u9996\u6bb5\u603b\u8d77\uff0c"can\'t help saying sorry"\u547c\u5e94\u6807\u9898\u4e3b\u65e8'],
            ['32', 'C', '\u627f\u63a5\u4e0a\u6587\u73b0\u8c61\uff0c\u7528\u7591\u95ee\u53e5"why"\u5f15\u51fa\u4e0b\u6587\u79d1\u5b66\u5bb6\u89e3\u91ca'],
            ['33', 'B', '\u89e3\u91ca\u82f1\u56fd\u4eba\u8bf4sorry\u7684\u9636\u7ea7\u539f\u56e0\uff0c\u4e0e\u524d\u53e5"class system"\u884d\u63a5'],
            ['34', 'A', '"Perhaps it does sometimes"\u56de\u5e94\u524d\u53e5"don\'t mean it?"\u7684\u8d28\u7591'],
            ['35', '\u5f00\u653e\u9898', 'Saying sorry politely is very useful/helpful in the UK.'],
        ],
        col_widths=[2, 2, 12.5]
    )

    section_head(doc, '\u2462', '\u8fd8\u539f\u9898\u89e3\u9898\u6280\u5de7\u603b\u7ed3', tag='\u5927\u62db')
    highlight_box(doc,
        '\u8fd8\u539f\u9898\u4e09\u770b\uff1a'
        '\u2460\u770b\u4ee3\u8bcd\u6307\u4ee3\uff08it/this/they\u6307\u5411\u8c01\uff09  '
        '\u2461\u770b\u903b\u8f91\u8fde\u63a5\u8bcd\uff08but/so/however\u524d\u540e\u65b9\u5411\uff09  '
        '\u2462\u770b\u590d\u73b0\u8bcd\uff08sorry/polite/class\u5728\u524d\u540e\u53e5\u662f\u5426\u51fa\u73b0\uff09',
        bg_color='E8F0FE', border_color='1F4E79'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\2024浙江省中考英语真题精析.docx')
    print('Doc 1 done: 2024浙江省中考英语真题精析.docx')


# ═══════════════════════════════════════════════════════════════
#  DOC 2: 2024北京中考英语阅读B-C篇精读
# ═══════════════════════════════════════════════════════════════

def gen_doc2():
    doc = new_doc()

    title_line(doc, '2024\u5317\u4eac\u5e02\u4e2d\u8003\u82f1\u8bed\u00b7\u9605\u8bfb\u7406\u89e3\u7cbe\u8bfb\u7b14\u8bb0', TITLE_SIZE)
    doc.add_paragraph()

    # ── Passage B: E-Waste Recycling ──
    part_head(doc, 'Passage B   E-Waste Recycling\uff08\u7535\u5b50\u5e9f\u5f03\u7269\u56de\u6536\uff09')
    source_note(doc, '\u6765\u6e90\uff1a2024\u5e74\u5317\u4eac\u5e02\u4e2d\u8003\u82f1\u8bed\u771f\u9898\uff0c\u9605\u8bfb\u7406\u89e3B\u7bc7')
    doc.add_paragraph()

    section_head(doc, '\u2460', '\u539f\u6587\u901a\u8bfb')

    p1_paras = [
        'The school year began. As president of the recycling club, Scott was thinking about new activities '
        'to encourage other students to become more enthusiastic about recycling. His club had helped to '
        'recycle a lot of waste for the past five years and he hoped that this year they would do even better.',
        'Scott went home one day and looked online, hoping to find ideas for events the club could organize '
        'at school. During his research, he learned that the amount of electronic waste, or e-waste, is '
        'increasing rapidly. Scott was upset with himself for not noticing this problem sooner. At that moment, '
        'he began to develop a plan.',
        '"There is a special project I want us to work on this term," Scott announced at the recycling club '
        'meeting the next day. "We have all heard about e-waste, but recently I learned about the bad effects '
        'it\'s having on our environment." He went on to tell the club members that he wanted them to organize '
        'an e-waste drive \u2014 a day when students and their families could drop off unwanted electronics to be '
        'recycled. "Now let\'s get to work!"',
        'For several weeks, the recycling club was busy preparing for the upcoming e-waste drive, or e-drive '
        'as it was soon called.',
        'The big day finally arrived, and Scott was nervous. He and the other members arrived at school early '
        'to make sure that everything would go smoothly. Everyone waited eagerly, hoping that students of the '
        'school would take part in the event.',
        'It wasn\'t long before the first cars pulled into the school parking lot. After a while, more people '
        'arrived. Phones, TV sets, computers and keyboards soon began piling up.',
        'At the end of the event, Scott heard someone suggest that the school should have an e-drive event '
        'every term. Scott smiled, realizing that a simple action could truly have a lasting influence.',
    ]
    for para in p1_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    section_head(doc, '\u2461', '\u6bb5\u843d\u7ed3\u6784\u5206\u6790\u8868', tag='\u91cd\u70b9')
    gt(doc,
        ['\u6bb5\u843d', '\u6bb5\u610f', '\u529f\u80fd'],
        [
            ['P1', 'Scott\u8eab\u4efd\uff08\u56de\u6536\u793e\u56e2\u4e3b\u5e2d\uff09+\u60f3\u6cd5', '\u80cc\u666f\u94fa\u8bbe'],
            ['P2', '\u7f51\u4e0a\u641c\u7d22\u53d1\u73b0e-waste\u95ee\u9898', '\u51b2\u7a81\u5f15\u5165'],
            ['P3', '\u5728\u793e\u56e2\u4f1a\u8bae\u63d0\u51fae-waste drive\u8ba1\u5212', '\u884c\u52a8\u542f\u52a8'],
            ['P4', '\u51c6\u5907\u8fc7\u7a0b', '\u8fc7\u6e21'],
            ['P5', '\u6d3b\u52a8\u5f53\u5929\u7d27\u5f20\u7b49\u5f85', '\u60c5\u611f\u63cf\u5199'],
            ['P6', '\u4eba\u4eec\u7eb7\u7eb7\u53c2\u4e0e\uff0c\u7535\u5b50\u4ea7\u54c1\u5806\u79ef', '\u9ad8\u6f6e'],
            ['P7', '\u6d3b\u52a8\u6210\u529f+\u611f\u609f\uff08simple action\u2192lasting influence\uff09', '\u4e3b\u65e8\u5347\u534e'],
        ],
        col_widths=[2, 7.5, 7]
    )

    section_head(doc, '\u2462', '\u51fa\u9898\u89c4\u5f8b\u5206\u6790', tag='\u91cd\u70b9')
    gt(doc,
        ['\u9898\u53f7', '\u8003\u67e5\u7c7b\u578b', '\u7b54\u6848', '\u89e3\u9898\u903b\u8f91'],
        [
            ['Q1', '\u7ec6\u8282\u5b9a\u4f4d', 'C. The amount of e-waste is increasing rapidly',
             'research online\u2192\u7b2c2\u6bb5\u539f\u6587\u590d\u73b0'],
            ['Q2', '\u60c5\u611f\u6001\u5ea6', 'D. Nervous',
             '\u7b2c5\u6bb5"Scott was nervous"\u76f4\u63a5\u51fa\u73b0'],
            ['Q3', '\u4e3b\u65e8\u63a8\u65ad', 'B. A simple action could bring a long-term effect',
             '\u5c3e\u6bb5\u5173\u952e\u53e5"a simple action could truly have a lasting influence"'],
        ],
        col_widths=[2, 3, 5, 6.5]
    )

    highlight_box(doc,
        '\u5317\u4eac\u4e2d\u8003\u9605\u8bfbB\u7bc7\u89c4\u5f8b\uff1a\u504f\u53d9\u4e8b\u578b\u8bf4\u660e\u6587\uff0c\u5e38\u8003\u2014\u2014'
        '\u2460\u7ec6\u8282\u5b9a\u4f4d\uff08\u5173\u952e\u8bcd\u56de\u6587\u5b9a\u4f4d\uff09  '
        '\u2461\u4eba\u7269\u60c5\u611f\u53d8\u5316\uff08\u6807\u6ce8\u60c5\u7eea\u8bcd\uff09  '
        '\u2462\u4e3b\u65e8\u6982\u62ec\uff08\u5c3e\u6bb5\u5173\u952e\u53e5\uff09',
        bg_color='E8F0FE', border_color='1F4E79'
    )

    # ── Passage C: Navigation Science ──
    part_head(doc, 'Passage C   Navigation Science\uff08\u65b9\u5411\u611f\u7814\u7a76\uff09')
    source_note(doc, '\u6765\u6e90\uff1a2024\u5e74\u5317\u4eac\u5e02\u4e2d\u8003\u82f1\u8bed\u771f\u9898\uff0c\u9605\u8bfb\u7406\u89e3C\u7bc7')
    doc.add_paragraph()

    section_head(doc, '\u2460', '\u539f\u6587\u901a\u8bfb')

    nav_paras = [
        'Sam Hill is really bad at finding his way from place to place. The world is full of people like '
        'Hill \u2014 and their opposites, who always seem to know exactly where they are and how to get where '
        'they want to go. It has proved hard to explain why. However, with the development of technology, '
        'there\'s new excitement happening in the research world.',

        'An experiment was carried out in 2022 to find out what might influence wayfinding ability. '
        'Researchers developed an online game in which players travel by boat to find where a lot of '
        'checkpoints lie. The game asked players to provide basic background information, and nearly four '
        'million people worldwide did so. Through the game, the researchers were able to judge navigational '
        'ability by looking at how far each person traveled to reach all the checkpoints. Then they compared '
        'players\' performance with their background information.',

        'The researchers found that Northern Europeans seemed to be better navigators, perhaps because they '
        'love orienteering, a sport which involves cross-country running and navigation. And those from cities '
        'with more disorganized street networks did better than those from cities with orderly ones. Perhaps '
        'people of planned cities don\'t need to build complex maps in their minds.',

        'Research results like these suggest that people\'s life experience decides how well they find their '
        'way. In fact, experience may even explain a popular belief that men are more likely to perform better '
        'than women. It turns out that this difference is more a question of culture and experience than of '
        'inborn ability. Northern Europeans, for example, show almost no gender difference in navigation. '
        'However, men do much better than women in places where women face cultural limits on exploring their '
        'environment on their own.',

        'That finding is also supported by studies on the Tsimane, a community living in a forest in South '
        'America. Researchers put GPS units on 305 Tsimane people to check their daily movements over a '
        'three-day period, and found no difference between men and women in navigational ability. Even '
        'children performed very well \u2014 a result, researchers think, of growing up in an environment that '
        'encourages children to explore the forest.',
    ]
    for para in nav_paras:
        passage_text(doc, para)
    doc.add_paragraph()

    section_head(doc, '\u2461', '\u9605\u8bfb\u7406\u89e3\u9898\u76ee\u4e0e\u7cbe\u6790', tag='\u91cd\u70b9')
    gt(doc,
        ['\u9898\u53f7', '\u9898\u5e72', '\u7b54\u6848', '\u89e3\u6790'],
        [
            ['Q27', 'Why was the experiment carried out?',
             'D. To find out why people differ in wayfinding ability',
             '\u7b2c2\u6bb5\u9996\u53e5"to find out what might influence wayfinding ability"'],
            ['Q28', 'Who is probably best at finding their way?',
             'A. A woman who often explores nature',
             '\u5168\u6587\u8bba\u70b9\uff1a\u7ecf\u9a8c\u51b3\u5b9a\u80fd\u529b\uff0c\u7ecf\u5e38\u63a2\u7d22\u81ea\u7136\u7684\u4eba\u5bfc\u822a\u66f4\u5f3a'],
            ['Q29', 'What can we learn from the passage?',
             'A. Good navigators are mostly made, not born',
             '\u4e3b\u65e8\u53e5"life experience decides"\uff0c\u800c\u975e\u5929\u751f\u80fd\u529b'],
        ],
        col_widths=[2, 5, 4.5, 5]
    )

    section_head(doc, '\u2462', '\u6570\u636e\u8bba\u8bc1\u624b\u6cd5\u5206\u6790')
    gt(doc,
        ['\u6570\u636e\u70b9', '\u51fa\u5904', '\u8bba\u8bc1\u4f5c\u7528'],
        [
            ['400\u4e07\u73a9\u5bb6', '\u7b2c2\u6bb5\uff08nearly four million people\uff09', '\u6837\u672c\u91cf\u5927\uff0c\u589e\u5f3a\u5b9e\u9a8c\u53ef\u4fe1\u5ea6'],
            ['305\u4ebaTsimane\u4eba', '\u7b2c5\u6bb5\uff08305 Tsimane people\uff09', '\u5177\u4f53\u4eba\u6570\uff0c\u8bc1\u660e\u8de8\u6587\u5316\u7814\u7a76\u4e25\u8c28'],
            ['3\u5929\u89c2\u6d4b\u671f', '\u7b2c5\u6bb5\uff08a three-day period\uff09', '\u660e\u786e\u65f6\u95f4\u8303\u56f4\uff0c\u589e\u5f3a\u65b9\u6cd5\u8bba\u900f\u660e\u5ea6'],
        ],
        col_widths=[3.5, 6, 7]
    )

    highlight_box(doc,
        '\u5317\u4eacC\u7bc7\u89c4\u5f8b\uff1a\u504f\u5b66\u672f\u578b\u79d1\u666e\uff0c\u5e38\u8003\u2014\u2014'
        '\u2460\u5b9e\u9a8c\u76ee\u7684(why)  '
        '\u2461\u6570\u636e\u63a8\u65ad(who/which)  '
        '\u2462\u5168\u6587\u4e3b\u65e8(what can we learn)',
        bg_color='E8F0FE', border_color='1F4E79'
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\2024北京中考英语阅读B-C篇精读.docx')
    print('Doc 2 done: 2024北京中考英语阅读B-C篇精读.docx')


# ═══════════════════════════════════════════════════════════════
#  DOC 3: 感叹句+过去完成时真题20道分类归档
# ═══════════════════════════════════════════════════════════════

def gen_doc3():
    doc = new_doc()

    title_line(doc, '2024\u5e74\u5404\u7701\u4e2d\u8003\u82f1\u8bed\u00b7\u611f\u53f9\u53e5+\u8fc7\u53bb\u5b8c\u6210\u65f6\u00b7\u771f\u9898\u5206\u7c7b\u5f52\u6863', TITLE_SIZE)
    doc.add_paragraph()

    # ── Part 1: 感叹句真题 ──
    part_head(doc, 'Part 1   \u611f\u53f9\u53e5\u771f\u9898\uff0810\u9053\uff09')

    section_head(doc, '\u2460', '\u771f\u9898\u7cbe\u9009\u8868', tag='\u91cd\u70b9')
    gt(doc,
        ['\u9898\u53f7', '\u7701\u4efd/\u5e74\u4efd', '\u9898\u5e72', '\u7b54\u6848', '\u8003\u6cd5\u5206\u7c7b'],
        [
            ['1', '2024 \u4e50\u5c71',
             '_______ amazing it is! The Shenzhou XVIII members raise fish for the first time in Tiangong space station.',
             'How', 'How + adj + S + V'],
            ['2', '2024 \u5170\u5dde',
             'China has set a 6G speed world record. _______ great country it is!',
             'What a', 'What a + adj + \u53ef\u6570\u5355\u6570'],
            ['3', '2024 \u9042\u5b81',
             'Our team won the first place. _______ exciting news it is!',
             'What', 'What + adj + \u4e0d\u53ef\u6570'],
            ['4', '2024 \u5409\u6797',
             '_______ beautiful poem Happy Rain on a Spring Night is!',
             'What a', 'What a + adj + \u53ef\u6570\u5355\u6570'],
            ['5', '2024 \u957f\u6625',
             '_______ helpful the speech is! It tells us to use the Internet safely.',
             'How', 'How + adj + S + V'],
            ['6', '2024 \u4e91\u5357',
             'China won 201 gold medals. _______ encouraging the news is.',
             'How', 'How + adj + S + V'],
            ['7', '2024 \u4e91\u5357',
             '_______ meaningful day! We volunteered to clean up our city park.',
             'What a', 'What a + adj + \u53ef\u6570\u5355\u6570'],
            ['8', '2024 \u4e91\u5357',
             '_______ unforgettable experience I had in Harbin Ice and Snow World!',
             'What an', 'What an + adj + \u5143\u97f3\u5f00\u5934\u5355\u6570'],
            ['9', '2024 \u6cf0\u5b89',
             '_______ hard-working students they are! They never stop trying.',
             'What', 'What + adj + \u590d\u6570\u540d\u8bcd'],
            ['10', '2024 \u6dc4\u535a',
             '_______ rapidly China is developing! More and more achievements have been made.',
             'How', 'How + adv + S + V'],
        ],
        col_widths=[1.2, 2.5, 6.5, 2, 4.3]
    )

    section_head(doc, '\u2461', '\u8003\u6cd5\u5206\u5e03\u7edf\u8ba1', tag='\u5f52\u7eb3')
    gt(doc,
        ['\u8003\u6cd5', '\u9898\u53f7', '\u5360\u6bd4'],
        [
            ['How + adj + S + V',              '1, 5, 6', '30%'],
            ['What a + adj + \u53ef\u6570\u5355\u6570',  '2, 4, 7', '30%'],
            ['What + adj + \u4e0d\u53ef\u6570/\u590d\u6570', '3, 9',    '20%'],
            ['What an + adj + \u5143\u97f3\u5f00\u5934\u5355\u6570', '8',      '10%'],
            ['How + adv + S + V',              '10',      '10%'],
        ],
        col_widths=[6, 3, 2]
    )

    # ── Part 2: 过去完成时真题 ──
    part_head(doc, 'Part 2   \u8fc7\u53bb\u5b8c\u6210\u65f6\u771f\u9898\uff0810\u9053\uff09')

    section_head(doc, '\u2460', '\u771f\u9898\u7cbe\u9009', tag='\u91cd\u70b9')

    ppt_qs = [
        (1,
         'The film ______ for ten minutes when we got to the cinema.',
         ['have already been on', 'had already begun', 'had already been on', 'have already begun'],
         'C', 'for + \u5ef6\u7eed\u6027\u52a8\u8bcd\uff0cwhen\u4ece\u53e5\u8fc7\u53bb\u65f6'),
        (2,
         'We ______ five English songs by the end of last term.',
         ['had learned', 'learned', 'have learned', 'will have learned'],
         'A', 'by + \u8fc7\u53bb\u65f6\u95f4'),
        (3,
         'Han Mei told me she ______ lunch, so she was very hungry.',
         ['has had', "hasn't have", 'have had', "hadn't had"],
         'D', 'when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09'),
        (4,
         'She ______ her keys in the office so she had to wait until her husband ______ home.',
         ['has left; comes', 'had left; would come', 'had left; came', 'left; had come'],
         'C', 'by + \u8fc7\u53bb\u65f6\u95f4\u70b9\uff0c\u524d\u540e\u65f6\u6001\u642d\u914d'),
        (5,
         'The meeting ______ when Mr. Wang ______ to school.',
         ['has begun; get', 'has been on; get', 'had begun; got', 'had been on; got'],
         'D', 'when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09+\u5ef6\u7eed\u6027'),
        (6,
         'By the end of last week, they ______ the bridge.',
         ['has completed', 'completed', 'will complete', 'had completed'],
         'D', 'when\u4ece\u53e5\uff08by the end of last week\uff09'),
        (7,
         'The students ______ their classroom when the visitors arrived.',
         ['have cleaned', 'had cleaned', 'was cleaned', 'have been cleaned'],
         'B', '\u5bbe\u8bed\u4ece\u53e5\u65f6\u6001\u540e\u79fb'),
        (8,
         'What ______ Jane ______ by the time he was seven?',
         ['did, do', 'has, done', 'did, did', 'had, done'],
         'D', 'by the time + \u8fc7\u53bb\u65f6'),
        (9,
         'When I arrived at the cinema, the movie ______ for 15 minutes.',
         ['began', 'had begun', 'had been on', 'has been on'],
         'C', 'by the time\u578b\uff0c\u5ef6\u7eed\u6027\u52a8\u8bcd'),
        (10,
         '\u2014Did you see Mr. Smith when you were in New York?\n\u2014No. When I arrived, he ______ to London.',
         ['went', 'has gone', 'had gone', 'was going'],
         'C', '"\u8fc7\u53bb\u7684\u8fc7\u53bb"\u7ecf\u5178\u8003\u6cd5'),
    ]

    for num, stem, opts, ans, note in ppt_qs:
        mc(doc, num, stem, opts)
    doc.add_paragraph()

    section_head(doc, '\u2461', '\u7b54\u6848\u4e0e\u89e3\u6790', tag='\u91cd\u70b9')
    gt(doc,
        ['\u9898\u53f7', '\u7b54\u6848', '\u8003\u6cd5\u5206\u7c7b', '\u89e3\u6790\u8981\u70b9'],
        [
            ['1', 'C', 'for/since + \u5ef6\u7eed\u6027', 'for ten minutes\u2192\u5ef6\u7eed\u6027\u52a8\u8bcdhad been on'],
            ['2', 'A', 'by + \u8fc7\u53bb\u65f6\u95f4', 'by the end of last term\u2192had learned'],
            ['3', 'D', 'when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09', 'told\u662f\u8fc7\u53bb\u65f6\uff0c\u201c\u6ca1\u5403\u5348\u996d\u201d\u5728told\u4e4b\u524d'],
            ['4', 'C', 'by + \u8fc7\u53bb\u65f6\u95f4\u70b9', 'had left\u5728\u524d\uff0ccame\u5728\u540e\uff0c\u65f6\u6001\u642d\u914d'],
            ['5', 'D', 'when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09', 'had been on\u5ef6\u7eed\u6027+when got\u8fc7\u53bb\u65f6'],
            ['6', 'D', 'when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09', 'by the end of last week\u2192had completed'],
            ['7', 'B', '\u5bbe\u8bed\u4ece\u53e5\u65f6\u6001\u540e\u79fb', '\u8bbf\u5ba2arrived\u65f6\uff0c\u5df2\u6253\u626b\u5b8c\uff0c\u65f6\u6001\u540e\u79fb'],
            ['8', 'D', 'by the time + \u8fc7\u53bb\u65f6', 'by the time he was seven\u2192had done'],
            ['9', 'C', 'for/since + \u5ef6\u7eed\u6027', 'for 15 minutes\u2192\u5ef6\u7eed\u6027had been on'],
            ['10', 'C', 'when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09', 'arrived\u662f\u8fc7\u53bb\u65f6\uff0chad gone\u5728\u5176\u4e4b\u524d'],
        ],
        col_widths=[1.5, 2, 5, 8]
    )

    section_head(doc, '\u2462', '\u8003\u6cd5\u5206\u5e03\u7edf\u8ba1', tag='\u5f52\u7eb3')
    gt(doc,
        ['\u8003\u6cd5', '\u9898\u53f7', '\u5360\u6bd4'],
        [
            ['by + \u8fc7\u53bb\u65f6\u95f4',                    '2, 4',        '20%'],
            ['when\u4ece\u53e5\uff08\u8fc7\u53bb\u7684\u8fc7\u53bb\uff09', '3, 5, 6, 10', '40%'],
            ['for/since + \u5ef6\u7eed\u6027',                   '1, 9',        '20%'],
            ['\u5bbe\u8bed\u4ece\u53e5\u65f6\u6001\u540e\u79fb',              '7',           '10%'],
            ['by the time + \u8fc7\u53bb\u65f6',                  '8',           '10%'],
        ],
        col_widths=[6, 3, 2]
    )

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\2024各省感叹句+过去完成时真题20道分类归档.docx')
    print('Doc 3 done: 2024各省感叹句+过去完成时真题20道分类归档.docx')


# ═══════════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    gen_doc1()
    gen_doc2()
    gen_doc3()
    print('\nAll 3 documents generated successfully.')
