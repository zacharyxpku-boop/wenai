import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def new_doc():
    doc = Document()
    doc.sections[0].page_width  = Cm(21)
    doc.sections[0].page_height = Cm(29.7)
    doc.sections[0].left_margin = doc.sections[0].right_margin = Cm(2)
    doc.sections[0].top_margin  = doc.sections[0].bottom_margin = Cm(2)
    return doc

def cell_bg(cell, hex6):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    pr.append(shd)

def banner(doc, label, zh, en, bg='E8672A'):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; cell_bg(c, bg)
    r = c.paragraphs[0].add_run(f'  {label}   {zh}   {en}  ')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    doc.add_paragraph()

def badge(doc, text, bg='F5A623'):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; cell_bg(c, bg)
    r = c.paragraphs[0].add_run(f'  {text}  ')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    doc.add_paragraph()

def sh(doc, num, title, tag=None):
    p = doc.add_paragraph()
    rn = p.add_run(f'{num} '); rn.bold = True; rn.font.color.rgb = RGBColor(0xE8,0x67,0x2A)
    p.add_run(title).bold = True
    if tag:
        rx = p.add_run(f'  [{tag}]'); rx.bold = True; rx.font.color.rgb = RGBColor(0xFF,0x4D,0x4D)

def gt(doc, headers, rows, hbg='C0A0FF'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers)); tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; cell_bg(c, hbg)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            c.paragraphs[0].add_run(str(v)).font.size = Pt(10)
    doc.add_paragraph()

def dazao(doc, n=3):
    t = doc.add_table(rows=1, cols=1); t.style = 'Table Grid'
    c = t.rows[0].cells[0]; cell_bg(c, 'FFF3E0')
    c.paragraphs[0].add_run(
        f'学习大招   难度系数：{"★"*n}{"☆"*(5-n)}\n\n动手，写下【大招】吧！\n\n\n'
    ).font.size = Pt(10)
    doc.add_paragraph()

def mc(doc, num, stem, opts):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    p.add_run(f'{num}. {stem}').font.size = Pt(10.5)
    op = doc.add_paragraph(); op.paragraph_format.left_indent = Cm(0.8)
    op.add_run('    '.join([f'{lb}. {o}' for lb, o in zip('ABCD', opts)])).font.size = Pt(10.5)

def mc3(doc, num, stem, opts):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.3)
    p.add_run(f'{num}. {stem}').font.size = Pt(10.5)
    op = doc.add_paragraph(); op.paragraph_format.left_indent = Cm(0.8)
    op.add_run('    '.join([f'{lb}. {o}' for lb, o in zip('ABC', opts)])).font.size = Pt(10.5)

def body(doc, text):
    p = doc.add_paragraph(text); p.paragraph_format.left_indent = Cm(0.5)
    if p.runs: p.runs[0].font.size = Pt(10.5)

def exlabel(doc, color='E8672A'):
    p = doc.add_paragraph()
    r = p.add_run('▶ 例题精讲'); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(*bytes.fromhex(color))

def rednote(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0xCC,0x44,0x44)

# ─── L3 正文 ──────────────────────────────────────────────────
doc = new_doc()

tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run('26秋上 初三英语讲义 A+班（全国通用版）')
tr.bold = True; tr.font.size = Pt(18)
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.add_run('Lesson Three   语法（感叹句）+ 完形技巧专项（情感态度推理）+ 听口 1').font.size = Pt(12)
doc.add_paragraph()

# ── PART ONE 感叹句 ──────────────────────────────────────────
banner(doc, 'PART ONE', '语法 —— 感叹句', 'Grammar', 'E8672A')
badge(doc, '知识精讲', 'F5A623')
doc.add_paragraph()

sh(doc, '①', '感叹句两大句型', tag='重点')
gt(doc,
    ['句型', '结构', '适用对象', '例句'],
    [
        ['What感叹句（可数名词单数）',
         'What + a/an + adj + 可数名词单数 + 主谓！',
         '修饰可数名词单数',
         'What an exciting journey it is!'],
        ['What感叹句（复数/不可数）',
         'What + adj + 可数名词复数/不可数名词 + 主谓！',
         '修饰复数/不可数名词',
         'What beautiful photos they took!'],
        ['How感叹句（adj）',
         'How + adj + 主谓（含be动词）！',
         '修饰形容词',
         'How proud her parents are!'],
        ['How感叹句（adv）',
         'How + adv + 主谓（含动作动词）！',
         '修饰副词',
         'How quickly time flies!'],
    ]
)

sh(doc, '②', 'What vs How 辨析口诀', tag='易错点')
body(doc, '口诀：看感叹词后面紧跟的是什么词性——')
body(doc, '✦  感叹词后紧跟【名词/名词短语】→ 用 What')
body(doc, '✦  感叹词后紧跟【形容词/副词】→ 用 How')
body(doc, '✦  a/an 的用法：可数名词单数前用 a/an；复数和不可数名词前不加')
doc.add_paragraph()

sh(doc, '③', '感叹句与陈述句互换', tag='重点 · 常考题型')
gt(doc,
    ['陈述句（原句）', '改为 How 感叹句', '改为 What 感叹句'],
    [
        ['The AI robot is very smart.',
         'How smart the AI robot is!',
         'What a smart AI robot it is!'],
        ['It is a wonderful invention.',
         'How wonderful the invention is!',
         'What a wonderful invention it is!'],
        ['The students are running very fast.',
         'How fast the students are running!',
         '—（fast为副词，无对应What句型）'],
        ['They are great scientists.',
         'How great they are!',
         'What great scientists they are!'],
    ],
    'A0C4FF'
)

sh(doc, '④', '感叹句三大易错点', tag='易错 · A+必看')
gt(doc,
    ['易错点', '错误示例', '正确形式', '原因'],
    [
        ['a/an 漏加',
         'What exciting game it is!',
         'What an exciting game it is!',
         'game为可数名词单数，exciting以元音音素开头，用an'],
        ['What/How混用',
         'How a clever boy he is!',
         'What a clever boy he is!',
         '后接名词boy，应用What'],
        ['语序错误',
         'What smart is he!',
         'How smart he is!',
         'How后接形容词，主谓不倒装'],
    ],
    'FFD0A0'
)
doc.add_paragraph()

dazao(doc, n=2)
exlabel(doc)

questions_L3_gram = [
    (1,
     '_______ creative idea the young programmer came up with at the AI competition!',
     ['What a', 'What', 'How a', 'How']),
    (2,
     '—Look at those giant pandas playing in the snow on the live stream!\n—_______ cute they look!',
     ['What', 'What a', 'How', 'How a']),
    (3,
     '_______ great progress our country has made in deep-sea exploration in recent years!',
     ['What a', 'What', 'How', 'How a']),
    (4,
     'The intangible cultural heritage show was breathtaking. _______ skilled performance the artists gave!',
     ['How', 'What a', 'What', 'How a']),
    (5,
     '_______ fast electric vehicles are developing and spreading across Chinese cities!',
     ['What', 'How', 'What a', 'How a']),
    (6,
     'The documentary about ocean plastic really shocked me. _______ serious problem marine pollution has become!',
     ['What a', 'What', 'How a', 'How']),
    (7,
     '—I just heard that the school robotics team won first place nationally.\n—Really? _______ exciting news!',
     ['What', 'What an', 'How', 'How an']),
]
for num, stem, opts in questions_L3_gram:
    mc(doc, num, stem, opts)
doc.add_paragraph()

# ── PART TWO 完形技巧：情感态度推理 ─────────────────────────
banner(doc, 'PART TWO', '完形技巧专项 —— 情感态度推理', 'Cloze Skills', '9B59B6')
badge(doc, '技巧精讲', '9B59B6')
doc.add_paragraph()

sh(doc, '①', '情感态度题命题规律')
body(doc, '✦ 考查形式：选出最符合语境的情感/态度词（形容词·副词·动词）')
body(doc, '✦ 命题位置：文章转折点、人物即时反应句、结尾总结段')
body(doc, '✦ 干扰项特征：含义相近但情感色彩相反，或程度轻重不符合语境')
body(doc, '✦ 核心逻辑：情感词必须与上下文事件逻辑一致，优先看转折词和递进词')
doc.add_paragraph()

sh(doc, '②', '情感词高频分类', tag='必背')
gt(doc,
    ['情感类别', '中考高频词汇'],
    [
        ['积极 / 正面', 'proud, grateful, delighted, amazed, inspired, confident, hopeful, relieved, touched, moved'],
        ['消极 / 负面', 'ashamed, disappointed, anxious, frightened, embarrassed, desperate, helpless, upset, nervous'],
        ['态度词',      'curious, doubtful, supportive, critical, optimistic, pessimistic, indifferent, enthusiastic'],
        ['程度辨析',    'surprised(惊讶) < amazed(惊叹) < astonished(震惊)；glad < pleased < delighted < overjoyed'],
        ['转折信号',    'but / however / yet / although → 情感方向逆转；even / still / increasingly → 情感递进加强'],
    ],
    'D0A0FF'
)

sh(doc, '③', '三步解题法', tag='大招')
body(doc, 'Step 1 — 定位：圈出空格所在句，标注情感触发事件（好事/坏事/意外）')
body(doc, 'Step 2 — 逻辑：判断上下文是顺承（→同向情感）还是转折（→逆向情感）')
body(doc, 'Step 3 — 排除：删去语义方向错误选项，再从剩余选项中比程度·比搭配·选最优')
doc.add_paragraph()

dazao(doc, n=4)

badge(doc, '完形填空精讲（换新题·25年末-26年一模热点话题）', '9B59B6')
doc.add_paragraph()
rednote(doc,
    '【选篇说明】本讲完形需换25年末-26年一模真题：成长励志类/青少年科技类语篇，'
    '约200-250词，10-15空，其中含3-4处情感态度词考查。'
    '五维评估：话题匹配✓  考频高✓  难度A+✓  情感词丰富✓  语篇有趣可讲✓'
)
doc.add_paragraph()

# 示例完形（话题：青少年面对AI辅助学习的心理成长，25-26高频）
body(doc, '[示例完形·正式版须替换为26一模原题]')
doc.add_paragraph()

passage = (
    '    When Li Wei first heard that his school would replace some traditional classes with AI-assisted '
    'learning, he felt   (1)  . He had always loved listening to teachers explain things face to face. '
    'Would a screen ever be as   (2)   as a real teacher?\n\n'
    '    However, his attitude began to   (3)   after just one week. The AI system could identify exactly '
    'which areas of maths he found difficult and provide   (4)   practice on those specific points. Unlike '
    'a classroom of forty students, the AI never seemed   (5)   when he asked the same question twice.\n\n'
    '    By the end of the month, Li Wei had   (6)   far more than in the whole previous term. He still '
    'believed human teachers were   (7)  , but he could now see that AI was a powerful   (8)  . '
    '"It\'s not about replacing teachers," he told his parents   (9)  . "It\'s about making learning more '
    '(10)   for every student." His mother smiled. She had never seen him so confident about school before.'
)
pp = doc.add_paragraph(passage)
pp.paragraph_format.left_indent = Cm(0.5)
if pp.runs: pp.runs[0].font.size = Pt(10.5)
doc.add_paragraph()

cloze = [
    (1,  'Li Wei felt _______ when he first heard the news.',
         ['excited', 'worried', 'bored', 'proud']),
    (2,  'Would a screen ever be as _______ as a real teacher?',
         ['strict', 'effective', 'popular', 'expensive']),
    (3,  'His attitude began to _______.',
         ['worsen', 'disappear', 'change', 'spread']),
    (4,  'The AI could provide _______ practice on difficult points.',
         ['targeted', 'limited', 'boring', 'random']),
    (5,  'The AI never seemed _______ when he asked the same question twice.',
         ['happy', 'impatient', 'confused', 'interested']),
    (6,  'Li Wei had _______ far more than in the previous term.',
         ['forgotten', 'achieved', 'spent', 'missed']),
    (7,  'He still believed human teachers were _______.',
         ['unnecessary', 'irreplaceable', 'outdated', 'ordinary']),
    (8,  'He could see that AI was a powerful _______.',
         ['problem', 'tool', 'barrier', 'subject']),
    (9,  'He told his parents _______.',
         ['angrily', 'proudly', 'sadly', 'quietly']),
    (10, 'Making learning more _______ for every student.',
         ['expensive', 'personal', 'traditional', 'simple']),
]
for num, stem, opts in cloze:
    mc(doc, num, stem, opts)
doc.add_paragraph()

# ── PART THREE 听口1 ────────────────────────────────────────
banner(doc, 'PART THREE', '听口训练 1 —— 听后选择 & 听后回答', 'Listening & Speaking', '27AE60')
badge(doc, '解题策略精讲', '27AE60')
doc.add_paragraph()

sh(doc, '①', '听后选择 — 预读三要素')
body(doc, '要素1【预读选项】：听前10秒扫描所有选项，预判话题和提问方向（人物/时间/地点/原因）')
body(doc, '要素2【抓关键词】：重点捕捉数字·地点·人名·情感词，遇到生词不要停，继续听')
body(doc, '要素3【注意否定】：not / never / hardly 等否定词后往往是答案关键信息')
doc.add_paragraph()

sh(doc, '②', '听后回答 — 答题四规范')
body(doc, '规范1：用完整句回答，不能只写单词或短语')
body(doc, '规范2：回答必须含录音中的关键信息词（人名·数字·地点·事件）')
body(doc, '规范3：时态与问题保持一致（问过去→回答用过去时；问现在→用现在时）')
body(doc, '规范4：字数不超过要求上限，语法正确比内容丰富更重要')
doc.add_paragraph()

sh(doc, '③', '听力练习题（换新题·26一模音频·话题：青少年·科技·社区）', tag='P0换新')
rednote(doc,
    '【选材说明】听口音频及题目须从26年一模听力原题或改编题中选取。'
    '话题建议：青少年科技体验 / 校园活动 / 社区志愿服务。'
    '音频二维码在讲义系统插入后由年级组长扫码确认。'
)
doc.add_paragraph()

body(doc, 'A. 听后选择（共5小题，每小题1分，共5分）')
listen_mc = [
    (1, 'What will the boy do this weekend?',
        ['Visit a science museum', 'Watch a football match', 'Join a volunteer activity']),
    (2, 'How does the girl feel about the AI project?',
        ['Worried about the result', 'Excited about the challenge', 'Bored with the topic']),
    (3, 'Where does the conversation probably take place?',
        ['In a school library', 'In a community lab', 'In a science museum']),
    (4, 'What does the woman suggest the boy do?',
        ['Read more books about AI', 'Talk to the teacher first', 'Join the robotics club']),
    (5, 'How long has the girl been learning to code?',
        ['Six months', 'One year', 'Two years']),
]
for num, stem, opts in listen_mc:
    mc3(doc, num, stem, opts)

doc.add_paragraph()
body(doc, 'B. 听后回答（共2小题，每小题2分，共4分）')
body(doc, '6. _______________________________________________')
body(doc, '7. _______________________________________________')
doc.add_paragraph()

rednote(doc,
    '【教师版备注】'
    '①感叹句例题话题已全面对标25-26年中考热点（AI/大熊猫/航天/非遗/新能源/海洋环保）。'
    '②完形示例文章务必换为26一模原题，确保情感态度词考查点≥3处。'
    '③听口部分音频二维码插入后由组长审核，题目配套音频须同步上传讲义系统。'
)

doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L3.docx')
print('L3 done.')
