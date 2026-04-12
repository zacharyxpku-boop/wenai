from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 页面设置 A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2)
section.right_margin = Cm(2)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading_banner(doc, part_num, title_cn, title_en, bg_hex='E8672A'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f'  PART {part_num}  {title_cn}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 背景色通过表格模拟
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    # 用table模拟色块
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cp.add_run(f'PART {part_num}   {title_cn}   {title_en}')
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()

def add_section_badge(doc, text, badge_color='F5A623'):
    p = doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, badge_color)
    cp = cell.paragraphs[0]
    r = cp.add_run(f'  {text}  ')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def add_numbered_section(doc, num, title, tag=None):
    p = doc.add_paragraph()
    run_num = p.add_run(f'{num} ')
    run_num.bold = True
    run_num.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)
    run_title = p.add_run(title)
    run_title.bold = True
    if tag:
        run_tag = p.add_run(f'  [{tag}]')
        run_tag.bold = True
        run_tag.font.color.rgb = RGBColor(0xFF, 0x4D, 0x4D)

def add_grammar_table(doc, headers, rows, header_bg='C0A0FF'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri+1].cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    doc.add_paragraph()

def add_fill_blank(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(0.5)
    p.runs[0].font.size = Pt(10.5)

def add_mc_question(doc, num, stem, options):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(f'{num}. {stem}')
    r.font.size = Pt(10.5)
    op = doc.add_paragraph()
    op.paragraph_format.left_indent = Cm(0.8)
    labels = ['A', 'B', 'C', 'D']
    text = '    '.join([f'{labels[i]}. {opt}' for i, opt in enumerate(options)])
    r2 = op.add_run(text)
    r2.font.size = Pt(10.5)

def add_dazao_box(doc, difficulty=3):
    p = doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, 'FFF3E0')
    cp = cell.paragraphs[0]
    stars = '★' * difficulty + '☆' * (5 - difficulty)
    r = cp.add_run(f'学习大招   难度系数：{stars}\n\n动手，写下【大招】吧！\n\n\n')
    r.font.size = Pt(10)
    doc.add_paragraph()

# ─────────────────────────────────────────────
# 封面信息
# ─────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('26秋上 初三英语讲义 A+班（全国通用版）')
tr.bold = True
tr.font.size = Pt(18)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('Lesson One   语法（过去完成时）+ 人物传记类阅读')
sr.font.size = Pt(12)
sr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ─────────────────────────────────────────────
# PART ONE 语法：过去完成时
# ─────────────────────────────────────────────
add_heading_banner(doc, 'ONE', '语法 —— 过去完成时', 'Grammar', 'E8672A')

add_section_badge(doc, '知识精讲', 'F5A623')
doc.add_paragraph()

add_numbered_section(doc, '①', '定义')
add_fill_blank(doc, '过去完成时表示在过去某一时间点之前已经发生或完成的动作，即"过去的过去"。')
add_fill_blank(doc, '结构：主语 + had + 过去分词（done）')
doc.add_paragraph()

add_numbered_section(doc, '②', '构成', tag='重点')
doc.add_paragraph()

add_grammar_table(doc,
    ['类型', '构成规则', '例句'],
    [
        ['肯定句', 'S + had + done', 'She had finished her homework before dinner.'],
        ['否定句', 'S + had not + done', 'He hadn\'t seen that film before.'],
        ['疑问句', 'Had + S + done?', 'Had they left when you arrived?'],
    ],
    'C0A0FF'
)

add_numbered_section(doc, '③', '过去分词构成', tag='易错点')
doc.add_paragraph()
add_grammar_table(doc,
    ['动词类型', '变化规则', '例词'],
    [
        ['一般动词', '直接 + ed', 'work → worked'],
        ['以e结尾', '去e + ed', 'live → lived'],
        ['辅音+元音+辅音结尾', '双写尾字母 + ed', 'stop → stopped'],
        ['辅音字母 + y结尾', '改 y为 i + ed', 'study → studied'],
    ]
)

doc.add_paragraph('不规则变化（高频考查）：').runs[0].bold = True
add_grammar_table(doc,
    ['原形', '过去式', '过去分词'],
    [
        ['go', 'went', 'gone'],
        ['take', 'took', 'taken'],
        ['write', 'wrote', 'written'],
        ['see', 'saw', 'seen'],
        ['come', 'came', 'come'],
        ['do', 'did', 'done'],
        ['give', 'gave', 'given'],
        ['know', 'knew', 'known'],
    ]
)

add_numbered_section(doc, '④', '常见时间状语', tag='重点')
doc.add_paragraph()
add_grammar_table(doc,
    ['时间状语', '含义'],
    [
        ['by + 过去时间', '到……为止（截止过去某时）'],
        ['by the time + 从句（一般过去时）', '到……时候'],
        ['before / when + 过去时/从句', '在……之前 / 当……时'],
        ['after + 从句（一般过去时）', '在……之后（主句用过去完成时）'],
    ],
    'A0C4FF'
)

add_dazao_box(doc, difficulty=3)

# 例题精讲
p_ex = doc.add_paragraph()
r_ex = p_ex.add_run('▶ 例题精讲')
r_ex.bold = True
r_ex.font.color.rgb = RGBColor(0xE8, 0x67, 0x2A)
r_ex.font.size = Pt(11)

add_mc_question(doc, 1,
    'By seven o\'clock this morning, Tom _______ breakfast with his family.',
    ["will have", "had had", "has had", "is having"])

add_mc_question(doc, 2,
    'She _______ the novel before the teacher asked about it in class.',
    ["reads", "has read", "had read", "was reading"])

add_mc_question(doc, 3,
    'When I arrived at the station, the last train _______ already _______.',
    ["has / left", "had / left", "was / leaving", "did / leave"])

add_mc_question(doc, 4,
    '—How was the competition yesterday?\n—We _______ for three hours when the power went out.',
    ["practised", "have practised", "had been practising", "were practising"])

add_mc_question(doc, 5,
    'By the end of last term, the students _______ over 500 English words.',
    ["learn", "learned", "have learned", "had learned"])

add_mc_question(doc, 6,
    'The scientist said that he _______ the experiment many times before he got the result.',
    ["tries", "tried", "had tried", "has tried"])

doc.add_paragraph()

# ─────────────────────────────────────────────
# PART TWO 语篇：人物传记类阅读
# ─────────────────────────────────────────────
add_heading_banner(doc, 'TWO', '语篇阅读 —— 人物传记类', 'Reading', '2E7D9E')

add_section_badge(doc, '语篇精讲·技巧归纳', '2E7D9E')
doc.add_paragraph()

tip_p = doc.add_paragraph()
tip_p.paragraph_format.left_indent = Cm(0.5)
tr2 = tip_p.add_run('人物传记类阅读解题策略：')
tr2.bold = True
tr2.font.size = Pt(11)

tips = [
    '① 关注时间线：传记类文章多以时间顺序叙述，快速定位"when-what"结构',
    '② 识别转折词：but / however / yet 后常是命题重点',
    '③ 主旨题定位：首段尾句或尾段首句往往是中心思想',
    '④ 推断题技巧：答案要有文本依据，不可凭常识判断',
]
for tip in tips:
    p = doc.add_paragraph(tip)
    p.paragraph_format.left_indent = Cm(0.8)
    p.runs[0].font.size = Pt(10.5)

doc.add_paragraph()

add_section_badge(doc, '阅读理解精讲（换新题·25年末-26年一模）', '2E7D9E')
doc.add_paragraph()

notice_p = doc.add_paragraph()
notice_p.paragraph_format.left_indent = Cm(0.5)
nr = notice_p.add_run('【选篇说明】本讲语篇需从25年末-26年一模各省真题中选取人物传记类阅读1篇。\n选篇评估维度：主题匹配大纲 ✓  考频高 ✓  难度符合A+班定位 ✓  题型丰富 ✓  语篇内容有趣可讲 ✓')
nr.font.size = Pt(10)
nr.font.color.rgb = RGBColor(0xCC, 0x44, 0x44)
nr.italic = True

doc.add_paragraph()

# 示例阅读文章框架（待替换）
sample_p = doc.add_paragraph()
sample_p.paragraph_format.left_indent = Cm(0.5)
sr2 = sample_p.add_run('[待选篇·示例结构] 人物传记类阅读（约250-300词）')
sr2.bold = True
sr2.font.size = Pt(11)

passage_placeholder = doc.add_paragraph(
    '    Malala Yousafzai was born in Pakistan in 1997. From an early age, she showed a great passion '
    'for education. Her father ran a school in their hometown, and Malala grew up believing that every '
    'child deserved the right to learn.\n\n'
    '    When the Taliban took control of her region, they banned girls from attending school. Malala '
    'refused to stay silent. She began writing a blog for the BBC, describing life under Taliban rule '
    'and calling for girls\' right to education. Her courage brought her international attention.\n\n'
    '    In October 2012, a Taliban gunman shot Malala on her way home from school. She survived after '
    'months of treatment in the UK. Rather than giving up, she became even more determined. In 2013, '
    'she gave a powerful speech at the United Nations on her 16th birthday.\n\n'
    '    In 2014, Malala became the youngest person ever to win the Nobel Peace Prize. She continues '
    'her work through the Malala Fund, fighting for 12 years of free, quality education for every girl.'
)
passage_placeholder.paragraph_format.left_indent = Cm(0.5)
passage_placeholder.runs[0].font.size = Pt(10.5)

doc.add_paragraph()

# 阅读题
q_label = doc.add_paragraph()
r_ql = q_label.add_run('▶ 阅读理解题')
r_ql.bold = True
r_ql.font.color.rgb = RGBColor(0x2E, 0x7D, 0x9E)
r_ql.font.size = Pt(11)

add_mc_question(doc, 1,
    'What is the main idea of the passage?',
    ["The history of the Taliban in Pakistan",
     "A brave girl\'s fight for education",
     "How to win the Nobel Peace Prize",
     "The importance of writing blogs"])

add_mc_question(doc, 2,
    'Why did Malala start writing a blog for the BBC?',
    ["To practise her writing skills",
     "To report on her school life",
     "To speak out against the ban on girls\' education",
     "To become internationally famous"])

add_mc_question(doc, 3,
    'What happened to Malala in October 2012?',
    ["She won the Nobel Peace Prize",
     "She spoke at the United Nations",
     "She was shot by a Taliban gunman",
     "She started the Malala Fund"])

add_mc_question(doc, 4,
    'The underlined word "determined" in Paragraph 3 probably means _______.',
    ["frightened", "firm and decided", "tired and sad", "surprised"])

add_mc_question(doc, 5,
    'Which of the following best describes Malala?',
    ["Careful and quiet", "Brave and determined", "Selfish and proud", "Shy but talented"])

doc.add_paragraph()

# ─────────────────────────────────────────────
# 词汇精讲板块（本讲无独立词汇，标注备注）
# ─────────────────────────────────────────────
note_p = doc.add_paragraph()
nr2 = note_p.add_run('【教师版备注】本讲语法重点：过去完成时。语篇选篇务必换25末-26一模人物传记类新题，示例文章仅供格式参考，正式版需替换。')
nr2.font.size = Pt(9)
nr2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
nr2.italic = True

# 保存
output_path = r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L1.docx'
doc.save(output_path)
print('Done:', output_path)
