import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 公共工具函数 ───────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2)
    sec.top_margin  = sec.bottom_margin = Cm(2)
    return doc

def banner(doc, part, title_cn, title_en, bg='E8672A'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r = p.add_run(f'PART {part}   {title_cn}   {title_en}')
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    doc.add_paragraph()

def badge(doc, text, bg='F5A623'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r = p.add_run(f'  {text}  ')
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    doc.add_paragraph()

def sec_title(doc, num, title, tag=None):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num} ')
    r1.bold = True; r1.font.color.rgb = RGBColor(0xFF,0x6B,0x35)
    r2 = p.add_run(title); r2.bold = True
    if tag:
        r3 = p.add_run(f'  [{tag}]')
        r3.bold = True; r3.font.color.rgb = RGBColor(0xFF,0x4D,0x4D)

def gtable(doc, headers, rows, hbg='C0A0FF'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, hbg)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold=True; r.font.size=Pt(10)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(val)).font.size = Pt(10)
    doc.add_paragraph()

def dazao(doc, stars=3):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.rows[0].cells[0]
    set_cell_bg(c, 'FFF3E0')
    r = c.paragraphs[0].add_run(
        f'学习大招   难度系数：{"★"*stars}{"☆"*(5-stars)}\n\n动手，写下【大招】吧！\n\n\n')
    r.font.size = Pt(10)
    doc.add_paragraph()

def mc(doc, num, stem, opts, indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.add_run(f'{num}. {stem}').font.size = Pt(10.5)
    op = doc.add_paragraph()
    op.paragraph_format.left_indent = Cm(indent+0.5)
    labels = ['A','B','C','D']
    op.add_run('    '.join(f'{labels[i]}. {o}' for i,o in enumerate(opts))).font.size = Pt(10.5)

def ex_header(doc, color_rgb=(0xE8,0x67,0x2A)):
    p = doc.add_paragraph()
    r = p.add_run('▶ 例题精讲')
    r.bold=True; r.font.size=Pt(11)
    r.font.color.rgb = RGBColor(*color_rgb)

def body(doc, text, indent=0.5):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(indent)
    if p.runs: p.runs[0].font.size = Pt(10.5)

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size=Pt(9); r.italic=True
    r.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)

# ═══════════════════════════════════════════════════════
# L2：易混时态辨析 + 中考词汇精讲1
# ═══════════════════════════════════════════════════════

def make_L2():
    doc = new_doc()

    # 封面
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('26秋上 初三英语讲义 A+班（全国通用版）')
    r.bold=True; r.font.size=Pt(18)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Lesson Two   语法（易混时态辨析）+ 中考词汇精讲1')
    r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(0x88,0x88,0x88)
    doc.add_paragraph()

    # PART ONE
    banner(doc, 'ONE', '语法 —— 易混时态辨析', 'Grammar', 'E8672A')
    badge(doc, '知识精讲', 'F5A623')
    doc.add_paragraph()

    sec_title(doc, '①', '四大时态对比总览', tag='重点')
    doc.add_paragraph()
    gtable(doc,
        ['时态','结构','核心含义','标志词（举例）'],
        [
            ['一般过去时',    'did',            '过去某时发生的动作',      'yesterday / last year / in 2020 / ago'],
            ['过去进行时',    'was/were doing',  '过去某时正在进行',        'at that time / when sb. did…'],
            ['现在完成时',    'have/has done',   '过去发生→现在有影响',     'already / yet / just / ever / never / since / for'],
            ['过去完成时',    'had done',        '过去的过去，更早发生',     'by then / before / when（从句一般过去时）'],
        ],
        'C0A0FF'
    )

    sec_title(doc, '②', '易混点1：现在完成时 vs 一般过去时', tag='易错点')
    doc.add_paragraph()
    gtable(doc,
        ['对比维度','现在完成时','一般过去时'],
        [
            ['时间状语', 'since 2020 / for 3 years / just / already / ever / never', 'last year / in 2020 / yesterday / ago / just now'],
            ['与现在的联系', '有（强调结果或影响仍存在）', '无（单纯陈述过去）'],
            ['例句对比',
             'I have lived here for 5 years. （还住着）',
             'I lived there for 5 years. （已不住了）'],
        ],
        'A0C4FF'
    )

    sec_title(doc, '③', '易混点2：过去进行时 vs 一般过去时', tag='易错点')
    doc.add_paragraph()
    body(doc, '口诀：when 引导时间状语从句时——')
    body(doc, '• 主句一般过去时 + when从句一般过去时 → 两动作先后发生')
    body(doc, '• 主句过去进行时 + when从句一般过去时 → 从句动作发生时，主句动作正在进行')
    body(doc, '• 主句一般过去时 + while从句过去进行时 → 进行中被打断')
    doc.add_paragraph()

    sec_title(doc, '④', '易混点3：过去完成时 vs 一般过去时', tag='易错点')
    doc.add_paragraph()
    gtable(doc,
        ['判断依据','用过去完成时','用一般过去时'],
        [
            ['时间先后', '有明确的"更早"参照点（by / before / when…）', '没有参照点，只陈述过去'],
            ['关键词',   'by the time / had already / no sooner…than', 'when / then / yesterday…'],
            ['例句',
             'He had left when I arrived.',
             'He left and then I arrived.'],
        ],
        'C0A0FF'
    )

    dazao(doc, stars=4)

    ex_header(doc)
    # 6道题——主题涉及AI科技/传统文化/学校生活（26一模热点话题）
    mc(doc, 1,
        'Liu Yang _______ hard at the robotics project when her teacher called her name.',
        ['studied','had studied','was studying','has studied'])
    mc(doc, 2,
        '—Have you ever tried the new AI learning app?\n—Yes, I _______ it last week and found it really helpful.',
        ['tried','have tried','had tried','was trying'])
    mc(doc, 3,
        'By the time the Spring Festival Gala started, Grandma _______ all the dumplings.',
        ['makes','made','has made','had made'])
    mc(doc, 4,
        'The researchers found that the ancient painting _______ in the museum for over 300 years.',
        ['kept','has been kept','had been kept','was kept'])
    mc(doc, 5,
        'I _______ my homework for two hours when the power suddenly went off.',
        ['did','have done','had done','had been doing'])
    mc(doc, 6,
        '—Where is Xiao Ming?\n—He _______ to the library. He _______ there since 9 a.m.',
        ['went / was','has gone / has been','had gone / had been','went / has been'])
    mc(doc, 7,
        'When I got to the station, I realized I _______ my student ID at home.',
        ['leave','left','had left','have left'])
    mc(doc, 8,
        'The athlete said he _______ for this competition for over a year.',
        ['is preparing','prepared','had been preparing','has prepared'])

    doc.add_paragraph()

    # PART TWO 词汇精讲1
    banner(doc, 'TWO', '中考词汇精讲1 —— 词根记忆法（第1-7组）', 'Vocabulary', '2A6AE8')
    badge(doc, '词根精讲·以根串词', '2A6AE8')
    doc.add_paragraph()

    sec_title(doc, '①', '词根 port（运输/携带）')
    doc.add_paragraph()
    gtable(doc,
        ['单词','词性','中文释义','记忆联想/例句'],
        [
            ['export',    'v./n.', '出口，输出',   'ex（出）+ port → 运出去'],
            ['import',    'v./n.', '进口，输入',   'im（入）+ port → 运进来'],
            ['transport', 'v./n.', '运输，交通',   'trans（越过）+ port → 运过去'],
            ['report',    'v./n.', '报道，汇报',   're（回）+ port → 带回消息'],
            ['support',   'v./n.', '支持，支撑',   'sup（下）+ port → 从下面撑起'],
            ['portable',  'adj.', '便携的',        'port + able → 能携带的'],
        ],
        '5BA3F5'
    )

    sec_title(doc, '②', '词根 dict（说/表达）')
    doc.add_paragraph()
    gtable(doc,
        ['单词','词性','中文释义','记忆联想/例句'],
        [
            ['predict',      'v.',   '预言，预测',  'pre（提前）+ dict → 提前说出来'],
            ['dictionary',   'n.',   '词典',        'dict + ionary → 说话的书'],
            ['indicate',     'v.',   '表明，指出',  'in + dic + ate → 向内说明'],
        ],
        '5BA3F5'
    )

    sec_title(doc, '③', '词根 form（形状/构成）')
    doc.add_paragraph()
    gtable(doc,
        ['单词','词性','中文释义','记忆联想/例句'],
        [
            ['inform',    'v.',   '通知，告知',  'in（使）+ form → 使成形→传达信息'],
            ['perform',   'v.',   '表演，执行',  'per（完全）+ form → 完整呈现'],
            ['reform',    'v./n.','改革，改良',  're（再）+ form → 重新塑形'],
            ['uniform',   'n.',   '制服，校服',  'uni（一）+ form → 统一形式'],
            ['platform',  'n.',   '平台，站台',  'plat（平）+ form → 平整的形状'],
        ],
        '5BA3F5'
    )

    doc.add_paragraph()
    badge(doc, '语境练习——用词填空', '2A6AE8')
    doc.add_paragraph()
    body(doc, '用方框内单词的适当形式填空（每词限用一次）：')
    body(doc, '[ export  import  support  predict  perform  reform  uniform  inform ]')
    doc.add_paragraph()
    fills = [
        '1. The government decided to _______ the education system to meet modern needs.',
        '2. Scientists cannot _______ exactly when the next earthquake will happen.',
        '3. All the students wear _______ to school in China.',
        '4. China _______ many high-quality products to countries around the world.',
        '5. The young pianist will _______ at the school concert next Friday.',
        '6. Please _______ us of any changes to your travel plans in advance.',
        '7. The school _______ students in developing good study habits.',
    ]
    for f in fills:
        body(doc, f)

    doc.add_paragraph()
    note(doc, '【教师版备注】词汇精讲选题已换新语境；例题8道涵盖26一模高频话题（AI/传统节日/学校生活）。语篇类内容本讲无独立阅读，词汇练习作为语言输出训练。')

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L2.docx')
    print('L2 Done')

# ═══════════════════════════════════════════════════════
# L3：感叹句 + 完形技巧（情感态度推理）+ 听口1
# ═══════════════════════════════════════════════════════

def make_L3():
    doc = new_doc()

    # 封面
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('26秋上 初三英语讲义 A+班（全国通用版）')
    r.bold=True; r.font.size=Pt(18)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Lesson Three   语法（感叹句）+ 完形技巧（情感态度推理）+ 听口1')
    r2.font.size=Pt(12); r2.font.color.rgb=RGBColor(0x88,0x88,0x88)
    doc.add_paragraph()

    # PART ONE 感叹句
    banner(doc, 'ONE', '语法 —— 感叹句', 'Grammar', 'E8672A')
    badge(doc, '知识精讲', 'F5A623')
    doc.add_paragraph()

    sec_title(doc, '①', '感叹句两大句型', tag='重点')
    doc.add_paragraph()
    gtable(doc,
        ['句型','结构公式','适用情况','例句'],
        [
            ['What感叹句',
             'What + (a/an) + adj + 可数名词单数 + 主谓！\nWhat + adj + 不可数名词/复数名词 + 主谓！',
             '感叹名词',
             'What a clever boy he is!\nWhat beautiful flowers they are!'],
            ['How感叹句',
             'How + adj/adv + 主谓！',
             '感叹形容词或副词',
             'How fast she runs!\nHow exciting the game is!'],
        ],
        'C0A0FF'
    )

    sec_title(doc, '②', 'What vs How 辨析口诀', tag='易错点')
    doc.add_paragraph()
    body(doc, '判断步骤：感叹的核心词是什么？')
    body(doc, '• 核心词是 名词 → 用 What（名词前可加形容词）')
    body(doc, '• 核心词是 形容词/副词 → 用 How')
    body(doc, '')
    body(doc, '特殊情况：')
    body(doc, '• What + a/an 后接 形+单数可数名词 → What an amazing idea it is!')
    body(doc, '• 不可数名词/复数名词前无 a/an → What great progress she has made!')
    body(doc, '• How + 形容词 后有名词，但名词是主语，不影响句型 → How heavy the bag is!')
    doc.add_paragraph()

    sec_title(doc, '③', '感叹句与陈述句互换', tag='重点')
    doc.add_paragraph()
    gtable(doc,
        ['感叹句','互换陈述句','规律'],
        [
            ['What a nice day it is!',       'The day is very nice.',      'What a/an + adj + n → very + adj'],
            ['How hard he works!',           'He works very hard.',        'How + adv → very + adv'],
            ['What clever students they are!','The students are very clever.','What + adj + n(复数) → very + adj'],
        ],
        'A0C4FF'
    )

    dazao(doc, stars=3)
    ex_header(doc)

    mc(doc, 1,
        '_______ interesting story the writer has told us!',
        ['What an','What a','How','How an'])
    mc(doc, 2,
        '_______ quickly time passes when you are having fun!',
        ['What','What a','How','How a'])
    mc(doc, 3,
        '_______ progress Chinese technology has made in recent years!',
        ['What a','What','How','How an'])
    mc(doc, 4,
        '—I heard Lin Tao won first prize at the science fair.\n—_______ proud his parents must be!',
        ['What','What a','How','How a'])
    mc(doc, 5,
        '_______ wonderful experience it was to visit the Palace Museum!',
        ['What a','What','How','How a'])
    mc(doc, 6,
        '_______ hard the volunteers worked to help the community!',
        ['What','What a','How','How a'])
    mc(doc, 7,
        'The sentence "It is such a touching film!" can be rewritten as "_______".',
        ['How touching film it is!',
         'What a touching film it is!',
         'How a touching film it is!',
         'What touching film it is!'])

    doc.add_paragraph()

    # PART TWO 完形技巧
    banner(doc, 'TWO', '完形技巧 —— 情感态度推理', 'Cloze Technique', '2E9E50')
    badge(doc, '技巧精讲', '2E9E50')
    doc.add_paragraph()

    sec_title(doc, '①', '什么是情感态度推理题？')
    body(doc, '完形填空中，部分空格要求判断人物的情感状态（happy/sad/proud/scared…）或作者态度（supportive/critical/neutral…），需结合上下文语境推断，不能靠直觉。')
    doc.add_paragraph()

    sec_title(doc, '②', '解题四步法', tag='重点')
    doc.add_paragraph()
    steps = [
        'Step 1  锁定情感词所在句，找触发情感的"事件"',
        'Step 2  判断事件性质（正面事件→积极情感，负面事件→消极情感）',
        'Step 3  排除语义矛盾选项（语义相近但强度不符的要特别小心）',
        'Step 4  回代验证——将选项填入后通读，确认前后逻辑通顺',
    ]
    for s in steps:
        body(doc, s)
    doc.add_paragraph()

    sec_title(doc, '③', '高频情感词分类', tag='重点')
    doc.add_paragraph()
    gtable(doc,
        ['情感类别','高频词汇'],
        [
            ['积极情感', 'proud / excited / delighted / grateful / relieved / hopeful / moved / inspired'],
            ['消极情感', 'worried / frightened / embarrassed / disappointed / ashamed / upset / guilty'],
            ['中性/转折', 'surprised / confused / curious / determined / calm / nervous'],
        ],
        '2E9E50'
    )

    doc.add_paragraph()
    badge(doc, '完形填空精讲（换新题·26一模励志类语篇）', '2E9E50')
    doc.add_paragraph()

    p_note = doc.add_paragraph()
    nr = p_note.add_run('【选篇说明】本讲完形填空需从25年末-26年一模各省真题中选取励志/成长类语篇1篇（约250词，15空）。情感词类空格不少于3处，便于技巧应用。')
    nr.font.size=Pt(10); nr.italic=True; nr.font.color.rgb=RGBColor(0xCC,0x44,0x44)
    doc.add_paragraph()

    # 示例完形（可换）
    body(doc, '[示例完形·可替换为26一模真题]', indent=0.3)
    passage = (
        '    Last summer, I joined a volunteer programme to help elderly people in my community. '
        'At first, I felt a little  1  because I didn\'t know what to say to them.\n\n'
        '    My first visit was to Mr. Chen, an 80-year-old man who lived alone. When I knocked on his door, '
        'he looked  2  to see a young visitor. We talked for hours about his life as an engineer. '
        'I was deeply  3  by his stories of building bridges in the mountains.\n\n'
        '    Before I left, Mr. Chen said, "You\'ve made my day, young friend." I felt so  4  '
        'that a simple visit could mean so much. From then on, I visited him every week.\n\n'
        '    Looking back, I\'m  5  I signed up for the programme. It taught me that giving time '
        'is sometimes more valuable than giving money.'
    )
    body(doc, passage)
    doc.add_paragraph()

    blanks = [
        (1, '_______ when meeting Mr. Chen for the first time.', ['nervous','excited','proud','bored']),
        (2, 'Mr. Chen looked _______ when he saw the visitor.',  ['angry','surprised','sad','tired']),
        (3, 'The narrator was _______ by Mr. Chen\'s stories.',  ['bored','frightened','moved','confused']),
        (4, 'The narrator felt _______ that a simple visit meant so much.',['guilty','amazed','ashamed','hopeless']),
        (5, 'The narrator is _______ she joined the programme.', ['sorry','regretful','grateful','worried']),
    ]
    for num, stem, opts in blanks:
        mc(doc, num, stem, opts)

    doc.add_paragraph()

    # PART THREE 听口
    banner(doc, 'THREE', '听口训练1 —— 听后选择 & 听后回答', 'Listening & Speaking', '9E2E7A')
    badge(doc, '听力策略精讲', '9E2E7A')
    doc.add_paragraph()

    sec_title(doc, '①', '听后选择——中考常见题型及解题策略')
    doc.add_paragraph()
    gtable(doc,
        ['题型','预读重点','答题技巧'],
        [
            ['对话判断题',    '关键名词/数字/地点',       '听第一遍抓主题，第二遍锁定答案词'],
            ['图片选择题',    '图片差异（动作/地点/数量）','预读图片差异，听时对号入座'],
            ['短文选择题',    '题目选项的逻辑关系',        '关注首句主题句，数字/时间务必记录'],
        ],
        '9E2E7A'
    )

    sec_title(doc, '②', '听后回答——回答规范与常见句式', tag='重点')
    doc.add_paragraph()
    body(doc, '核心原则：答案完整，不可只答 Yes/No，必须补充信息。')
    doc.add_paragraph()
    gtable(doc,
        ['问题类型','规范回答句式','示例'],
        [
            ['Yes/No 问题',   'Yes, + 完整句 / No, + 完整句',      'Yes, I do. I usually study for two hours a day.'],
            ['What 问题',     '直接给出内容，主语+谓语完整',          'My favourite subject is science because…'],
            ['Why 问题',      'Because + 原因句，可加 Also…',       'Because it helps me think creatively. Also…'],
            ['How 问题',      '描述方式/程度，给出具体细节',           'I usually relax by reading or taking a walk.'],
        ],
        '9E2E7A'
    )

    doc.add_paragraph()
    p_lo = doc.add_paragraph()
    r_lo = p_lo.add_run('【听口练习材料】')
    r_lo.bold=True; r_lo.font.color.rgb=RGBColor(0x9E,0x2E,0x7A)
    body(doc, '本讲听口训练材料（听后选择3题 + 听后回答2题）需对应26一模听力原题或改编题，由教研组统一提供音频文件，教师根据班级实际情况调整播放节奏。')

    doc.add_paragraph()
    note(doc, '【教师版备注】L3语法感叹句7道例题，完形5道情感题，听口策略表格。完形正式版需换26一模真题，示例文章可供格式/难度参考。')

    doc.save(r'C:\Users\86136\Desktop\初三英语教研\【26秋上】初三英语讲义A+班_L3.docx')
    print('L3 Done')

make_L2()
make_L3()
